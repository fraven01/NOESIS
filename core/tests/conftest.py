"""Gemeinsame Testkonfiguration für das core-Modul."""
from unittest.mock import patch

import pytest
from django.core.management import call_command
from io import BytesIO
from docx import Document
import fitz
from pathlib import Path

pytest_plugins = ["core.tests.factories"]


@pytest.fixture(autouse=True, scope="session")
def disable_langfuse_for_tests():
    """Deaktiviert Langfuse global während der Tests.

    Verhindert, dass beim (Neu-)Import von ``core.llm_utils``
    die Langfuse-Instrumentierung aktiviert wird und Events sendet.
    """
    # Erzwinge das Feature-Flag über die Umgebung
    mp = pytest.MonkeyPatch()
    mp.setenv("LANGFUSE_ENABLED", "False")
    # Falls Einstellungen bereits geladen sind, explizit abschalten
    try:
        from django.conf import settings as dj_settings  # type: ignore
        # raising=False, damit es auch funktioniert, wenn das Attribut fehlt
        mp.setattr(dj_settings, "LANGFUSE_ENABLED", False, raising=False)
    except Exception:
        pass
    try:
        yield
    finally:
        mp.undo()


@pytest.fixture(scope="module")
def seed_db(django_db_setup, django_db_blocker) -> None:
    """Initialisiert die Testdatenbank mit Seed-Daten."""
    with django_db_blocker.unblock():
        from django.contrib.auth.models import User
        from core.models import LLMConfig, Anlage4Config, Anlage4ParserConfig

        call_command("seed_initial_data", verbosity=0)
        LLMConfig.objects.get_or_create()
        Anlage4Config.objects.get_or_create()
        Anlage4ParserConfig.objects.get_or_create()
        if not User.objects.filter(username="baseuser").exists():
            User.objects.create_user("baseuser", password="pass")
        # Sicherstellen, dass der Superuser aus dem Seed-Skript
        # ein bekanntes Passwort für die Tests besitzt.
        superuser = User.objects.get(username="frank")
        superuser.set_password("pass")
        superuser.save()


@pytest.fixture
def user(seed_db, db) -> "User":
    """Gibt den Basisbenutzer zurück."""
    from django.contrib.auth.models import User

    return User.objects.get(username="baseuser")


@pytest.fixture
def superuser(seed_db, db) -> "User":
    """Gibt den Basis-Superuser zurück."""
    from django.contrib.auth.models import User

    return User.objects.get(username="frank")


@pytest.fixture(autouse=True)
def mock_llm_api_calls():
    """Ersetzt externe LLM-Aufrufe durch statische Antworten."""
    with (
        patch("core.llm_utils.query_llm", return_value="Mock-Antwort"),
        patch("core.llm_utils.query_llm_with_images", return_value="Mock-Antwort"),
        patch("core.llm_tasks.query_llm", return_value="Mock-Antwort"),
        patch("google.generativeai.configure"),
        patch("google.generativeai.list_models", return_value=[]),
    ):
        yield


@pytest.fixture
def docx_content_path(tmp_path: "Path") -> "Path":
    """Erzeugt ein einfaches DOCX-Dokument."""
    doc = Document()
    doc.add_paragraph("Docx Inhalt")
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    path = tmp_path / "content.docx"
    path.write_bytes(bio.getvalue())
    return path


@pytest.fixture
def docx_two_page_path(tmp_path: "Path") -> "Path":
    """DOCX-Dokument mit zwei Seiten."""
    doc = Document()
    doc.add_paragraph("Seite 1")
    doc.add_page_break()
    doc.add_paragraph("Seite 2")
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    path = tmp_path / "two_pages.docx"
    path.write_bytes(bio.getvalue())
    return path


@pytest.fixture
def pdf_one_page_path(tmp_path: "Path") -> "Path":
    """Einseitiges PDF-Dokument."""
    pdf = fitz.open()
    pdf.new_page()
    bio = BytesIO()
    pdf.save(bio)
    path = tmp_path / "one_page.pdf"
    path.write_bytes(bio.getvalue())
    return path


@pytest.fixture
def pdf_two_page_path(tmp_path: "Path") -> "Path":
    """PDF-Dokument mit zwei Seiten."""
    pdf = fitz.open()
    pdf.new_page()
    pdf.new_page()
    bio = BytesIO()
    pdf.save(bio)
    path = tmp_path / "two_pages.pdf"
    path.write_bytes(bio.getvalue())
    return path


@pytest.fixture
def anlage2_table_docx_path(tmp_path: "Path") -> "Path":
    """DOCX mit einfacher Anlage-2-Tabelle."""
    doc = Document()
    table = doc.add_table(rows=2, cols=5)
    table.cell(0, 0).text = "Funktion"
    table.cell(0, 1).text = "Technisch vorhanden"
    table.cell(0, 2).text = "Einsatz bei Telefónica"
    table.cell(0, 3).text = "Zur LV-Kontrolle"
    table.cell(0, 4).text = "KI-Beteiligung"
    table.cell(1, 0).text = "Anmelden"
    table.cell(1, 1).text = "Ja"
    table.cell(1, 2).text = "Nein"
    table.cell(1, 3).text = "Nein"
    table.cell(1, 4).text = "Ja"
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    path = tmp_path / "anlage2_table.docx"
    path.write_bytes(bio.getvalue())
    return path


@pytest.fixture
def prepared_files(
    request,
    docx_content_path,
    docx_two_page_path,
    pdf_one_page_path,
    pdf_two_page_path,
    anlage2_table_docx_path,
) -> None:
    """Stellt vorbereitete Dateien als Klassenattribute bereit."""
    request.cls.docx_content_path = docx_content_path
    request.cls.docx_two_page_path = docx_two_page_path
    request.cls.pdf_one_page_path = pdf_one_page_path
    request.cls.pdf_two_page_path = pdf_two_page_path
    request.cls.anlage2_table_docx_path = anlage2_table_docx_path
