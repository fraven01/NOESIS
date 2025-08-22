from django.contrib.auth.models import User, Group
from django.urls import reverse
from django.test import TestCase
from django.conf import settings
from django.http import QueryDict
from django.db import IntegrityError
from types import SimpleNamespace
import os


from django.apps import apps
from ..models import (
    BVProject,
    BVProjectFile,
    Recording,
    Prompt,
    LLMConfig,
    Tile,
    UserTileAccess,
    GroupTileAccess,
    Anlage1Question,
    Anlage1Config,
    Area,
    Anlage2Function,
    Anlage2Config,
    Anlage2ColumnHeading,
    Anlage2SubQuestion,
    AnlagenFunktionsMetadaten,
    FunktionsErgebnis,
    SoftwareKnowledge,
    BVSoftware,
    Gutachten,
    AntwortErkennungsRegel,
    Anlage4Config,
    Anlage4ParserConfig,
    ZweckKategorieA,
    Anlage5Review,
)
from ..docx_utils import (
    extract_text,
    get_docx_page_count,
    get_pdf_page_count,
    extract_images,
    parse_anlage2_table,
    _normalize_header_text,
)

from ..utils import start_analysis_for_file
from .. import text_parser

from core.text_parser import parse_anlage2_text, PHRASE_TYPE_CHOICES

from ..anlage4_parser import parse_anlage4

from ..parser_manager import parser_manager
from ..parsers import AbstractParser

from pathlib import Path
from tempfile import NamedTemporaryFile, TemporaryDirectory
from io import BytesIO
from docx import Document
import shutil
from PIL import Image
import fitz
from django.conf import settings

from django.core.files.uploadedfile import SimpleUploadedFile
from ..forms import (
    BVProjectForm,
    BVProjectUploadForm,
    BVProjectFileJSONForm,
    BVProjectFileForm,
    Anlage2ConfigForm,
    Anlage2ReviewForm,
)
from ..workflow import set_project_status
from ..models import ProjectStatus
from ..llm_tasks import (
    classify_system,
    check_anlage1,
    check_anlage2,
    analyse_anlage3,
    analyse_anlage4,
    analyse_anlage4_async,
    check_anlage3_vision,
    check_anlage5,
    run_conditional_anlage2_check,
    worker_verify_feature,
    worker_generate_gutachten,
    worker_run_initial_check,
    worker_run_anlage3_vision,
    worker_anlage4_evaluate,
    worker_generate_gap_summary,
    summarize_anlage1_gaps,
    summarize_anlage2_gaps,
    get_prompt,
    generate_gutachten,
    run_anlage2_analysis,
    parse_anlage1_questions,
    _parse_anlage2,
)
from ..views import (
    _verification_to_initial,
    _build_row_data,
    _build_supervision_row,
    _save_project_file,
    extract_anlage_nr,
    get_user_tiles,
    _has_manual_gap,
    _build_supervision_groups,
    _resolve_value,
)
from ..reporting import generate_gap_analysis, generate_management_summary
from unittest.mock import patch, ANY, Mock, call
from django.core.management import call_command
from django.test import override_settings
from django.conf import settings
import json


def create_statuses() -> None:
    ProjectStatus.objects.all().delete()
    data = [
        ("NEW", "Neu"),
        ("CLASSIFIED", "Klassifiziert"),
        ("GUTACHTEN_OK", "Gutachten OK"),
        ("GUTACHTEN_FREIGEGEBEN", "Gutachten freigegeben"),
        ("IN_PRUEFUNG_ANLAGE_X", "In Prüfung Anlage X"),
        ("FB_IN_PRUEFUNG", "FB in Prüfung"),
        ("ENDGEPRUEFT", "Endgeprüft"),
    ]
    for idx, (key, name) in enumerate(data, start=1):
        ProjectStatus.objects.create(
            name=name,
            key=key,
            ordering=idx,
            is_default=key == "NEW",
            is_done_status=key == "ENDGEPRUEFT",
        )


def setUpModule():
    pass


def create_project(software: list[str] | None = None, **kwargs) -> BVProject:
    projekt = BVProject.objects.create(**kwargs)
    for name in software or []:
        BVSoftware.objects.create(project=projekt, name=name)
    return projekt


def seed_test_data(*, skip_prompts: bool = False) -> None:
    """Befüllt die Test-Datenbank mit Initialdaten.

    Bestehende Einträge werden zuvor gelöscht. Optional können die
    Prompt-Definitionen übersprungen werden.
    """
    BVProject.objects.all().delete()
    Prompt.objects.all().delete()
    Tile.objects.all().delete()
    Area.objects.all().delete()
    Anlage2Function.objects.all().delete()
    Anlage2SubQuestion.objects.all().delete()
    Anlage2ColumnHeading.objects.all().delete()
    Anlage2Config.objects.all().delete()
    Anlage1Question.objects.all().delete()
    apps.get_model("core", "Anlage1QuestionVariant").objects.all().delete()
    apps.get_model("core", "LLMRole").objects.all().delete()
    AntwortErkennungsRegel.objects.all().delete()
    Anlage4Config.objects.all().delete()
    Anlage4ParserConfig.objects.all().delete()
    LLMConfig.objects.all().delete()
    apps.get_model("core", "ZweckKategorieA").objects.all().delete()
    apps.get_model("core", "Anlage3ParserRule").objects.all().delete()
    apps.get_model("core", "SupervisionStandardNote").objects.all().delete()
    from django.apps import apps as django_apps
    from core.management.commands.seed_initial_data import (
        create_initial_data,
    )
    from ..llm_tasks import ANLAGE1_QUESTIONS

    try:
        create_initial_data(django_apps)
    except LookupError:
        # Falls die Migrationsfunktion wegen entfernter Modelle
        # fehlschlägt, legen wir die benötigten Objekte manuell an.
        Anlage1QuestionModel = apps.get_model("core", "Anlage1Question")
        Anlage1QuestionVariant = apps.get_model("core", "Anlage1QuestionVariant")
        for idx, text in enumerate(ANLAGE1_QUESTIONS, start=1):
            question, _ = Anlage1QuestionModel.objects.update_or_create(
                num=idx,
                defaults={
                    "text": text,
                    "enabled": True,
                    "parser_enabled": True,
                    "llm_enabled": True,
                },
            )
            Anlage1QuestionVariant.objects.get_or_create(question=question, text=text)
    create_statuses()

    # Erforderliche Konfigurationen bereitstellen
    LLMConfig.objects.get_or_create()
    Anlage4Config.objects.get_or_create()
    Anlage4ParserConfig.objects.get_or_create()

    # Anlage1 Fragen aktualisieren
    Anlage1QuestionModel = apps.get_model("core", "Anlage1Question")
    Anlage1QuestionVariant = apps.get_model("core", "Anlage1QuestionVariant")
    for idx, text in enumerate(ANLAGE1_QUESTIONS, start=1):
        try:
            question = Anlage1QuestionModel.objects.get(num=idx)
            question.text = text
            question.parser_enabled = True
            question.llm_enabled = True
            question.save()
        except Anlage1QuestionModel.DoesNotExist:
            question = Anlage1QuestionModel.objects.create(
                num=idx,
                text=text,
                parser_enabled=True,
                llm_enabled=True,
            )
        Anlage1QuestionVariant.objects.get_or_create(question=question, text=text)

    if skip_prompts:
        return

    for idx, text in enumerate(ANLAGE1_QUESTIONS, start=1):
        Prompt.objects.update_or_create(name=f"anlage1_q{idx}", defaults={"text": text})

    Prompt.objects.update_or_create(
        name="check_anlage3_vision",
        defaults={
            "text": "Prüfe die folgenden Bilder der Anlage. Gib ein JSON mit 'ok' und 'hinweis' zurück:\n\n"
        },
    )
    # Weitere Prompts für Tests bereitstellen
    roles = {r.name: r for r in apps.get_model("core", "LLMRole").objects.all()}
    prompt_data = {
        "anlage1_email": {
            "text": (
                "Formuliere eine freundliche E-Mail an den Fachbereich. "
                "Wir haben die Anlage 1 geprüft und noch folgende Vorschläge, "
                "bevor der Mitbestimmungsprozess weiter gehen kann. "
                "Bitte fasse die folgenden Vorschläge zusammen:\r\n\r\n"
            )
        },
        "anlage2_ai_involvement_check": {
            "text": (
                "Antworte ausschließlich mit 'Ja' oder 'Nein'. Frage: "
                "Beinhaltet die Funktion '{function_name}' der Software "
                "'{software_name}' typischerweise eine KI-Komponente? "
                "Eine KI-Komponente liegt vor, wenn die Funktion "
                "unstrukturierte Daten (Text, Bild, Ton) verarbeitet, "
                "Sentiment-Analysen durchführt oder nicht-deterministische, "
                "probabilistische Ergebnisse liefert."
            )
        },
        "anlage2_feature_justification": {
            "text": (
                "Warum besitzt die Software '{software_name}' typischerweise die "
                "Funktion oder Eigenschaft '{function_name}'?   Ist es möglich "
                "mit der {function_name} eine Leistungskontrolle oder eine "
                "Verhaltenskontrolle  Ist damit eine Leistungskontrolle oder "
                "eine Verhaltenskontrolle im Sinne des §87 Abs. 1 Nr. 6 möglich? "
                "Wenn ja, wie?"
            )
        },
        "anlage2_subquestion_possibility_check": {
            "text": (
                "Im Kontext der Funktion '{function_name}' der Software "
                "'{software_name}': Ist die spezifische Anforderung "
                "'{subquestion_text}' technisch möglich? Antworte nur mit 'Ja', "
                "'Nein' oder 'Unsicher'."
            )
        },
        "anlage2_subquestion_justification_check": {
            "text": (
                " [SYSTEM]\nDu bist Fachautor*in für IT-Mitbestimmung (§87 Abs. 1 Nr. 6 BetrVG).\n"
                "Antworte Unterfrage prägnant in **maximal zwei Sätzen** (insgesamt ≤ 65 Wörter) und erfülle folgende Regeln :\n\n"
                "1. Starte Teil A mit „Typischer Zweck: …“  \n2. Starte Teil B mit „Kontrolle: Ja, …“ oder „Kontrolle: Nein, …“.  \n"
                "3. Nenne exakt die übergebene Funktion/Eigenschaft, erfinde nichts dazu.  \n"
                "4. Erkläre knapp *warum* mit der Funktion die Unterfrage (oder warum nicht) eine Leistungs- oder Verhaltenskontrolle möglich ist.  \n"
                "5. Verwende Alltagssprache, keine Marketing-Floskeln.\n\n"
                ' [USER]\nSoftware: {{software_name}}  \nFunktion/Eigenschaft: {{function_name}}  \nUnterfrage: "{{subquestion_text}}"'
            )
        },
        "anlage2_ai_verification_prompt": {
            "text": (
                "Gib eine kurze Begründung, warum die Funktion '{function_name}' "
                "(oder die Unterfrage '{subquestion_text}') der Software "
                "'{software_name}' eine KI-Komponente beinhaltet oder beinhalten kann, "
                "insbesondere im Hinblick auf die Verarbeitung unstrukturierter Daten "
                "oder nicht-deterministischer Ergebnisse."
            )
        },
        "anlage2_feature_verification": {
            "text": (
                "Deine einzige Aufgabe ist es, die folgende Frage mit einem einzigen "
                'Wort zu beantworten. Deine Antwort darf AUSSCHLIESSLICH "Ja", '
                '"Nein", oder "Unsicher" sein. Gib keine Einleitung, keine '
                "Begründung und keine weiteren Erklärungen ab.\r\n\r\nFrage: "
                "Besitzt die Software '{software_name}' basierend auf allgemeinem "
                "Wissen typischerweise die Funktion oder Eigenschaft "
                "'{function_name}'?\n\n{gutachten}"
            ),
            "role": roles.get("Standard"),
            "use_system_role": False,
        },
        "check_anlage2_function": {
            "text": (
                "Prüfe anhand des folgenden Textes, ob die genannte Funktion "
                'vorhanden ist. Gib ein JSON mit den Schlüsseln "technisch_verfuegbar", '
                '"einsatz_telefonica", "zur_lv_kontrolle" und "ki_beteiligung" '
                "zurück.\n\n"
            )
        },
        "check_anlage4": {
            "text": "Prüfe die folgende Anlage auf Vollständigkeit. Gib ein JSON mit 'ok' und 'hinweis' zurück:\n\n",
        },
        "check_anlage5": {
            "text": "Prüfe die folgende Anlage auf Vollständigkeit. Gib ein JSON mit 'ok' und 'hinweis' zurück:\n\n"
        },
        "classify_system": {
            "text": (
                "Bitte klassifiziere das folgende Softwaresystem. "
                "Gib ein JSON mit den Schlüsseln 'kategorie' und 'begruendung' zurück.\n\n"
            )
        },
        "generate_gutachten": {
            "text": (
                "Erstelle ein tiefgehendes Gutachten zu der Software im Sinne des § 87 Abs. 1 Nr. 6 BetrVG. "
                "Richte das Gutachten ausschließlich an Betriebsräte und überspringe allgemeine Erläuterungen "
                "zu DSGVO oder Datenschutzrecht ebenso musst du nicht erläutern, wann Mitbestimmungsrechte "
                "nach §87 (1) Nr. 6 gelten.\n\nDein Gutachten soll folgende Punkte abdecken: \n\n1. **Mitbestimmungspflichtige Funktionen**   \n- Liste alleFeatures auf, die der Leistungs- oder Verhaltenskontrolle dienen (z. B. Analyse von Nutzungshistorien, App- und Kommunikationsauswertung, Dateizugriffsprotokolle).\n- Erläutere kurz, warum jede einzelne Funktion unter § 87 1 Nr. 6 BetrVG fällt. \n\n2. **Eingriffsintensität aus Mitarbeitersicht**   \n   - Beschreibe, wie stark jede dieser Funktionen in den Arbeitsablauf eingreift und welche Verhaltensaspekte sie überwacht (z. B. Häufigkeit von App-Nutzung, Kommunikationsverhalten, Standortdaten).   \n   - Nutze eine Skala (gering – mittel – hoch) und begründe die Einstufung anhand typischer Betriebsabläufe in einem Telekommunikationsunternehmen. \n\n3. **Betroffene Leistungs- und Verhaltensaspekte**   \n   - Identifiziere konkret, welche Leistungskennzahlen (z. B. Aktivitätszeiten, App-Nutzungsdauer) und Verhaltensmuster (z. B. Kommunikationshäufigkeit, Datenübertragung) erfasst und ausgewertet werden.   \n   - Schätze ab, wie umfassend und detailliert die Auswertung jeweils ausfällt. \n\n4. **Handlungsbedarf für den Betriebsrat**   \n   - Fasse zusammen, bei welchen Funktionen und Einsatzszenarien eine Betriebsvereinbarung zwingend erforderlich ist. \n\n5. **Weitere Mitbestimmungsrechte (Kurzhinweise)**   \n   - Wenn offensichtlich erkennbar ist, dass andere relevante Mitbestimmungsrechte nach BetrVG (z. B. §§ 80 ff. zur Informationspflicht) berührt sind, bewerte kurz, warum diese Software dieses Recht des Betriebsrats berühren könnte. \n\nArbeite strukturiert mit klaren Überschriften und Bullet-Points. Wo sinnvoll, nutze kurze Tabellen oder Zusammenfassungen zur Übersicht. \n. Antworte auf deutsch.\nSoftware: \n"
            ),
            "role": roles.get("Gutachten"),
        },
        "initial_check_knowledge": {
            "text": "Kennst du die Software '{name}'? Antworte ausschließlich mit einem einzigen Wort: 'Ja' oder 'Nein'.",
            "use_system_role": False,
        },
        "initial_check_knowledge_with_context": {
            "text": (
                "Kennst du die Software '{name}'? Hier ist zusätzlicher Kontext, um sie zu identifizieren: \"{user_context}\". "
                "Antworte ausschließlich mit einem einigen Wort: 'Ja' oder 'Nein'."
            )
        },
        "initial_llm_check": {
            "text": (
                "Erstelle eine kurze, technisch korrekte Beschreibung für die Software '{name}'. "
                "Nutze Markdown mit Überschriften, Listen oder Fettdruck, um den Text zu strukturieren. "
                "Erläutere, was sie tut und wie sie typischerweise eingesetzt wird."
            ),
            "role": roles.get("Gutachten"),
        },
    }

    for name, data in prompt_data.items():
        Prompt.objects.update_or_create(
            name=name,
            defaults={
                "text": data["text"],
                "role": data.get("role"),
                "use_system_role": data.get("use_system_role", True),
            },
        )


class SeedInitialDataTests(TestCase):
    """Tests für das Seeding der Antwortregeln."""

    def test_answer_rules_seeded(self) -> None:
        """Prüft, ob die Antwortregeln angelegt werden."""
        call_command("seed_initial_data")
        from ..initial_data_constants import INITIAL_ANSWER_RULES

        for rule in INITIAL_ANSWER_RULES:
            obj = AntwortErkennungsRegel.objects.get(
                regel_name=rule["regel_name"]
            )
            self.assertEqual(
                obj.erkennungs_phrase,
                rule["erkennungs_phrase"],
            )
            self.assertEqual(
                obj.actions_json,
                rule["actions"],
            )


class NoesisTestCase(TestCase):
    """Basisklasse für alle Tests mit gefüllter Datenbank."""

    @classmethod
    def setUpTestData(cls):
        seed_test_data()
        cls.user = User.objects.create_user("baseuser", password="pass")
        cls.superuser = User.objects.create_superuser(
            "basesuper", "admin@example.com", password="pass"
        )


class ExtractAnlageNrTests(TestCase):
    """Tests für die Erkennung der Anlagen-Nummer aus Dateinamen."""

    def test_variants(self):
        self.assertEqual(extract_anlage_nr("Anlage 1.docx"), 1)
        self.assertEqual(extract_anlage_nr("Anlage-2.pdf"), 2)
        self.assertEqual(extract_anlage_nr("Anlage3.docx"), 3)


class BVProjectFileTests(NoesisTestCase):
    def test_create_project_with_files(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        for i in range(1, 4):
            f = SimpleUploadedFile(f"f{i}.txt", b"data")
            BVProjectFile.objects.create(
                project=projekt,
                anlage_nr=i,
                upload=f,
                text_content="data",
            )
        self.assertEqual(projekt.anlagen.count(), 3)
        self.assertListEqual(
            list(projekt.anlagen.values_list("anlage_nr", flat=True)), [1, 2, 3]
        )

    def test_default_flags(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content="d",
        )
        self.assertFalse(pf.manual_reviewed)
        self.assertFalse(pf.verhandlungsfaehig)
        pf.manual_reviewed = True
        pf.verhandlungsfaehig = True
        pf.save()
        pf.refresh_from_db()
        self.assertTrue(pf.manual_reviewed)
        self.assertTrue(pf.verhandlungsfaehig)

    def test_project_delete_removes_files(self):
        """Sichert, dass beim Löschen eines Projekts die Dateien entfernt werden."""
        with TemporaryDirectory() as tmpdir, override_settings(MEDIA_ROOT=tmpdir):
            projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
            pf = BVProjectFile.objects.create(
                project=projekt,
                anlage_nr=1,
                upload=SimpleUploadedFile("a.txt", b"data"),
                text_content="d",
            )
            file_path = pf.upload.path
            self.assertTrue(os.path.exists(file_path))
            projekt.delete()
            self.assertFalse(os.path.exists(file_path))

    def test_auto_start_analysis_saves_task_id(self):
        """Die Analyse-ID wird nach dem Upload gespeichert."""
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        with patch(
            "core.signals.start_analysis_for_file", return_value="tid"
        ) as mock_start:
            pf = BVProjectFile.objects.create(
                project=projekt,
                anlage_nr=1,
                upload=SimpleUploadedFile("a.txt", b"data"),
            )
        mock_start.assert_called_with(pf.pk)
        pf.refresh_from_db()
        self.assertEqual(pf.verification_task_id, "tid")

    def test_json_form_shows_analysis_field_for_anlage3(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=3,
            upload=SimpleUploadedFile("a.txt", b"x"),
            text_content="t",
            analysis_json={"ok": True},
        )
        form = BVProjectFileJSONForm(instance=pf)
        self.assertIn("analysis_json", form.fields)

    def test_save_does_not_start_task(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        with patch("core.models.async_task") as mock_task:
            pf = BVProjectFile.objects.create(
                project=projekt,
                anlage_nr=2,
                upload=SimpleUploadedFile("a.txt", b"x"),
                text_content="t",
            )
        self.assertEqual(pf.verification_task_id, "")
        mock_task.assert_not_called()

    def test_is_verification_running(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"x"),
            verification_task_id="tid",
        )
        with patch("core.models.fetch") as mock_fetch:
            mock_fetch.return_value = SimpleNamespace(success=None)
            self.assertTrue(pf.is_verification_running())
            mock_fetch.return_value = SimpleNamespace(success=True)
            self.assertFalse(pf.is_verification_running())

    def test_check_functions_clears_task_id(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        with patch("core.signals.start_analysis_for_file", return_value="tid"):
            pf = BVProjectFile.objects.create(
                project=projekt,
                anlage_nr=2,
                upload=SimpleUploadedFile("a.txt", b"x"),
                verification_task_id="tid",
            )
        Anlage2Function.objects.get(name="Anmelden")
        with (
            patch("core.llm_tasks.query_llm", return_value="{}"),
            patch("core.llm_tasks.async_task") as mock_async,
            patch("core.llm_tasks.result") as mock_result,
        ):
            mock_async.side_effect = lambda name, *a, **k: (
                worker_verify_feature(*a, **k) or "tid"
            )
            mock_result.side_effect = lambda *a, **k: None
            run_conditional_anlage2_check(pf.pk)
        pf.refresh_from_db()
        self.assertEqual(pf.verification_task_id, "")
        self.assertEqual(pf.processing_status, BVProjectFile.COMPLETE)

    def test_template_shows_disabled_state_when_task_running(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"x"),
            verification_task_id="tid",
        )
        self.client.login(username=self.superuser.username, password="pass")
        with patch("core.models.fetch") as mock_fetch:
            mock_fetch.return_value = SimpleNamespace(success=None)
            url = reverse("projekt_detail", args=[projekt.pk])
            resp = self.client.get(url)
        self.assertContains(resp, "Analyse läuft")

    def test_hx_project_file_status_running(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        with patch("core.signals.start_analysis_for_file", return_value="tid"):
            pf = BVProjectFile.objects.create(
                project=projekt,
                anlage_nr=2,
                upload=SimpleUploadedFile("a.txt", b"x"),
                verification_task_id="tid",
            )
        self.client.login(username=self.superuser.username, password="pass")
        with patch("core.models.fetch") as mock_fetch:
            mock_fetch.return_value = SimpleNamespace(success=None)
            url = reverse("hx_anlage_status", args=[pf.pk])
            resp = self.client.get(url)
        self.assertContains(resp, "hx-trigger")
        self.assertContains(resp, "animate-spin")

    def test_hx_project_file_status_ready(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        with patch("core.signals.start_analysis_for_file", return_value="tid"):
            pf = BVProjectFile.objects.create(
                project=projekt,
                anlage_nr=2,
                upload=SimpleUploadedFile("a.txt", b"x"),
                verification_task_id="tid",
            )
        self.client.login(username=self.superuser.username, password="pass")
        url = reverse("hx_anlage_status", args=[pf.pk])
        with patch("core.models.fetch") as mock_fetch:
            mock_fetch.return_value = SimpleNamespace(success=True)
            resp = self.client.get(url)
        self.assertContains(resp, "hx-trigger")
        # Finalen Status simulieren
        pf.processing_status = BVProjectFile.COMPLETE
        pf.verification_task_id = ""
        pf.save()
        resp = self.client.get(url)
        self.assertNotContains(resp, "hx-trigger")
        self.assertContains(resp, "Analyse bearbeiten")

    def test_hx_anlage_status_processing(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"x"),
            processing_status=BVProjectFile.PROCESSING,
        )
        self.client.login(username=self.superuser.username, password="pass")
        url = reverse("hx_anlage_status", args=[pf.pk])
        resp = self.client.get(url)
        self.assertContains(resp, "animate-spin")

    def test_hx_anlage_status_ready(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        # Der Signal-Handler `auto_start_analysis` speichert die
        # R\xFCckgabe von `start_analysis_for_file` als Task-ID. Wenn der
        # Mock kein einfaches String-Ergebnis liefert, w\xFCrde ein
        # `FieldError` auftreten. Daher liefern wir explizit eine
        # Zeichenkette zur\xFCck.
        with patch(
            "core.signals.start_analysis_for_file", return_value="mocked_task_id"
        ):
            pf = BVProjectFile.objects.create(
                project=projekt,
                anlage_nr=2,
                upload=SimpleUploadedFile("a.txt", b"x"),
                processing_status=BVProjectFile.COMPLETE,
                analysis_json={},
            )
        self.client.login(username=self.superuser.username, password="pass")
        url = reverse("hx_anlage_status", args=[pf.pk])
        resp = self.client.get(url)
        self.assertContains(resp, "Analyse bearbeiten")
        self.assertContains(resp, "Erneut analysieren")

    def test_hx_anlage_status_failed(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        with patch("core.signals.start_analysis_for_file", return_value=""):
            pf = BVProjectFile.objects.create(
                project=projekt,
                anlage_nr=2,
                upload=SimpleUploadedFile("a.txt", b"x"),
                processing_status=BVProjectFile.FAILED,
            )
        self.client.login(username=self.superuser.username, password="pass")
        url = reverse("hx_anlage_status", args=[pf.pk])
        resp = self.client.get(url)
        self.assertNotContains(resp, "hx-trigger")
        self.assertContains(resp, "Analyse fehlgeschlagen")
        self.assertContains(resp, "Erneut versuchen")

    def test_hx_anlage_status_pending(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        with patch("core.signals.start_analysis_for_file", return_value=""):
            pf = BVProjectFile.objects.create(
                project=projekt,
                anlage_nr=2,
                upload=SimpleUploadedFile("a.txt", b"x"),
                processing_status=BVProjectFile.PENDING,
            )
        self.client.login(username=self.superuser.username, password="pass")
        url = reverse("hx_anlage_status", args=[pf.pk])
        resp = self.client.get(url)
        self.assertContains(resp, "Analyse starten")
        self.assertContains(resp, "hx-trigger")

    def test_hx_project_anlage_tab(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        with patch("core.signals.start_analysis_for_file", return_value=""):
            pf = BVProjectFile.objects.create(
                project=projekt,
                anlage_nr=1,
                upload=SimpleUploadedFile("a.txt", b"x"),
            )
        self.client.login(username=self.superuser.username, password="pass")
        url = reverse("hx_project_anlage_tab", args=[projekt.pk, 1])
        resp = self.client.get(url)
        self.assertContains(resp, 'href="/media/bv_files/a_')
        self.assertContains(resp, '.txt"')
        self.assertContains(resp, "hx-trigger")

    def test_hx_anlage_row(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        with patch("core.signals.start_analysis_for_file", return_value=""):
            pf = BVProjectFile.objects.create(
                project=projekt,
                anlage_nr=1,
                upload=SimpleUploadedFile("a.txt", b"x"),
            )
        self.client.login(username=self.superuser.username, password="pass")
        url = reverse("hx_anlage_row", args=[pf.pk])
        resp = self.client.get(url)
        self.assertContains(resp, 'href="/media/bv_files/a_')
        self.assertContains(resp, '.txt"')
        self.assertContains(resp, "hx-trigger")

    def test_hx_toggle_project_file_flag(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"x"),
        )
        self.client.login(username=self.superuser.username, password="pass")
        url = reverse("hx_toggle_project_file_flag", args=[pf.pk, "manual_reviewed"])
        resp = self.client.post(
            url,
            {"value": "1"},
            HTTP_HX_REQUEST="true",
        )
        self.assertEqual(resp.status_code, 200)
        pf.refresh_from_db()
        self.assertTrue(pf.manual_reviewed)
        self.assertContains(resp, "<tr")
        self.assertContains(resp, "fa-check")

    def test_hx_project_software_tab(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        SoftwareKnowledge.objects.create(project=projekt, software_name="A")
        self.client.login(username=self.superuser.username, password="pass")
        url = reverse("hx_project_software_tab", args=[projekt.pk, "tech"])
        resp = self.client.get(url)
        self.assertContains(resp, "Prüfung starten")

    def test_trigger_file_analysis_starts_tasks(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"x"),
            processing_status=BVProjectFile.PENDING,
        )
        self.client.login(username=self.superuser.username, password="pass")
        with patch("core.views.start_analysis_for_file", return_value="123") as mock_start:
            url = reverse("trigger_file_analysis", args=[pf.pk])
            resp = self.client.post(url)
        mock_start.assert_called_with(pf.pk)
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.json(), {"task_id": "123"})

    def test_start_analysis_for_file_enqueues_tasks(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"x"),
            processing_status=BVProjectFile.PENDING,
        )
        with patch.object(
            BVProjectFile,
            "get_analysis_tasks",
            return_value=[("core.llm_tasks.check_anlage1", pf.pk)],
        ), patch("core.utils.async_task") as mock_async, patch(
            "core.utils.transaction.on_commit", side_effect=lambda func: func()
        ):
            mock_async.return_value = "t1"
            task_id = start_analysis_for_file(pf.pk)
        mock_async.assert_called_with("core.llm_tasks.check_anlage1", pf.pk)
        self.assertEqual(task_id, "t1")
        pf.refresh_from_db()
        self.assertEqual(pf.processing_status, BVProjectFile.PROCESSING)

    def test_get_analysis_tasks_returns_project_id_for_conditional_check(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"x"),
        )
        tasks = pf.get_analysis_tasks()
        self.assertEqual(
            tasks,
            [
                ("core.llm_tasks.worker_run_anlage2_analysis", pf.pk),
                ("core.llm_tasks.run_conditional_anlage2_check", pf.pk),
            ],
        )


@override_settings(ALLOWED_HOSTS=["testserver"])
class ProjektFileUploadTests(NoesisTestCase):
    def setUp(self):
        self.user = User.objects.create_user("user", password="pass")
        self.client.login(username="user", password="pass")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")

    def test_docx_upload_extracts_text(self):
        doc = Document()
        doc.add_paragraph("Docx Inhalt")
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("Anlage_1.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)

        url = reverse("hx_project_file_upload", args=[self.projekt.pk])
        resp = self.client.post(
            url,
            {"anlage_nr": 1, "upload": upload, "manual_comment": ""},
            format="multipart",
            HTTP_HX_REQUEST="true",
        )
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.content.decode().count("<tr"), 1)
        file_obj = self.projekt.anlagen.first()
        self.assertIsNotNone(file_obj)
        self.assertIn("Docx Inhalt", file_obj.text_content)

    def test_ownerless_project_allows_upload(self):
        """Prüft, dass ein neuer Nutzer Dateien in ein besitzerloses Projekt laden darf."""
        doc = Document()
        doc.add_paragraph("Inhalt")
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("Anlage_1.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)

        url = reverse("hx_project_file_upload", args=[self.projekt.pk])
        resp = self.client.post(
            url,
            {"anlage_nr": 1, "upload": upload, "manual_comment": ""},
            format="multipart",
            HTTP_HX_REQUEST="true",
        )
        self.assertEqual(resp.status_code, 200)

    def test_pdf_upload_stores_bytes(self):
        pdf = fitz.open()
        pdf.new_page()
        tmp = NamedTemporaryFile(delete=False, suffix=".pdf")
        tmp.close()
        pdf.save(tmp.name)
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("Anlage_3.pdf", fh.read())
        Path(tmp.name).unlink(missing_ok=True)

        url = reverse("hx_project_file_upload", args=[self.projekt.pk])
        resp = self.client.post(
            url,
            {"anlage_nr": 3, "upload": upload, "manual_comment": ""},
            format="multipart",
            HTTP_HX_REQUEST="true",
        )
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.content.decode().count("<tr"), 1)
        file_obj = self.projekt.anlagen.get(anlage_nr=3)
        self.assertEqual(file_obj.text_content, "")

    def test_upload_without_anlage_nr_uses_filename(self):
        """Nutzt die Anlagen-Nummer aus dem Dateinamen."""
        doc = Document()
        doc.add_paragraph("Inhalt")
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("Anlage 4 - Entwurf.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)

        url = reverse("hx_project_file_upload", args=[self.projekt.pk])
        resp = self.client.post(
            url,
            {"upload": upload, "manual_comment": ""},
            format="multipart",
            HTTP_HX_REQUEST="true",
        )
        self.assertEqual(resp.status_code, 200)
        self.assertTrue(self.projekt.anlagen.filter(anlage_nr=4).exists())

    def test_anlage2_upload_queues_check(self):
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Funktion"
        table.cell(0, 1).text = "Technisch vorhanden"
        table.cell(1, 0).text = "Anmelden"
        table.cell(1, 1).text = "Ja"
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("Anlage_2.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)

        Anlage2Function.objects.get(name="Anmelden")

        url = reverse("hx_project_file_upload", args=[self.projekt.pk])
        mock_async = Mock(side_effect=["tid1", "tid2"])

        def fake_start(file_id: int) -> str:
            pf_obj = BVProjectFile.objects.get(pk=file_id)
            pf_obj.processing_status = BVProjectFile.PROCESSING
            pf_obj.save(update_fields=["processing_status"])
            task_id = None
            for func, arg in pf_obj.get_analysis_tasks():
                tid = mock_async(func, arg)
                if task_id is None:
                    task_id = tid
            return task_id or ""

        with patch("core.signals.start_analysis_for_file", side_effect=fake_start):
            resp = self.client.post(
                url,
                {"anlage_nr": 2, "upload": upload, "manual_comment": ""},
                format="multipart",
                HTTP_HX_REQUEST="true",
            )
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.content.decode().count("<tr"), 1)
        pf = self.projekt.anlagen.get(anlage_nr=2)
        self.assertEqual(pf.verification_task_id, "tid1")
        self.assertEqual(pf.processing_status, BVProjectFile.PROCESSING)
        mock_async.assert_any_call(
            "core.llm_tasks.run_conditional_anlage2_check",
            pf.pk,
        )

    def test_second_anlage2_version_skips_ai_check(self):
        func = Anlage2Function.objects.get(name="Anmelden")
        first = BVProjectFile.objects.create(
            project=self.projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("v1.docx", b"x"),
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=first,
            funktion=func,
            quelle="ki",
            technisch_verfuegbar=True,
        )

        doc = Document()
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("Anlage_2.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)

        url = reverse("hx_project_file_upload", args=[self.projekt.pk])
        with patch("core.views.async_task") as mock_async:
            resp = self.client.post(
                url,
                {"anlage_nr": 2, "upload": upload, "manual_comment": ""},
                format="multipart",
                HTTP_HX_REQUEST="true",
            )
        self.assertEqual(resp.status_code, 200)
        self.assertFalse(
            any(
                call.args[0] == "core.llm_tasks.run_conditional_anlage2_check"
                for call in mock_async.call_args_list
            )
        )
        pf_latest = self.projekt.anlagen.filter(anlage_nr=2, is_active=True).first()
        self.assertEqual(pf_latest.version, 2)

    def test_upload_stores_posted_anlage_nr(self):
        doc = Document()
        doc.add_paragraph("x")
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("Anlage_5.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)
        url = reverse("hx_project_file_upload", args=[self.projekt.pk])
        resp = self.client.post(
            url,
            {"anlage_nr": 2, "upload": upload, "manual_comment": ""},
            format="multipart",
            HTTP_HX_REQUEST="true",
        )
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.content.decode().count("<tr"), 1)
        pf = self.projekt.anlagen.get()
        self.assertEqual(pf.anlage_nr, 2)

    def test_save_project_file_respects_form_value(self):
        doc = Document()
        doc.add_paragraph("x")
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("Anlage_5.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)

        form = BVProjectFileForm({}, {"upload": upload}, anlage_nr=1)
        self.assertTrue(form.is_valid())
        pf = _save_project_file(self.projekt, form)
        self.assertEqual(pf.anlage_nr, 1)


    def test_save_multiple_files_unique_numbers(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        for nr in range(1, 7):
            doc = Document()
            doc.add_paragraph(str(nr))
            tmp = NamedTemporaryFile(delete=False, suffix=".docx")
            doc.save(tmp.name)
            tmp.close()
            with open(tmp.name, "rb") as fh:
                upload = SimpleUploadedFile(f"Anlage_{nr}.docx", fh.read())
            Path(tmp.name).unlink(missing_ok=True)
            _save_project_file(projekt, upload=upload, anlage_nr=nr)

        qs = BVProjectFile.objects.filter(project=projekt)
        self.assertEqual(qs.count(), 6)
        self.assertListEqual(
            sorted(qs.values_list("anlage_nr", flat=True)),
            [1, 2, 3, 4, 5, 6],
        )

    def test_upload_uses_filename_when_no_anlage_nr(self):
        doc = Document()
        doc.add_paragraph("x")
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("Anlage_4.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)

        url = reverse("hx_project_file_upload", args=[self.projekt.pk])
        resp = self.client.post(
            url,
            {"upload": upload, "manual_comment": ""},
            format="multipart",
            HTTP_HX_REQUEST="true",
        )
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.content.decode().count("<tr"), 1)
        pf = self.projekt.anlagen.get()
        self.assertEqual(pf.anlage_nr, 4)

    def test_upload_uses_filename_when_anlage_nr_empty(self):
        doc = Document()
        doc.add_paragraph("x")
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("Anlage_3.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)

        url = reverse("hx_project_file_upload", args=[self.projekt.pk])
        resp = self.client.post(
            url,
            {"anlage_nr": "", "upload": upload},
            format="multipart",
            HTTP_HX_REQUEST="true",
        )
        self.assertEqual(resp.status_code, 200)
        pf = self.projekt.anlagen.get()
        self.assertEqual(pf.anlage_nr, 3)


class DropzoneUploadTests(NoesisTestCase):
    """Tests für den neuen Datei-Upload-Workflow."""

    def setUp(self):
        self.user = User.objects.create_user("dz", password="pass")
        self.client.login(username="dz", password="pass")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")

    def test_number_from_filename(self):
        doc = Document()
        doc.add_paragraph("x")
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("Anlage_2.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)

        url = reverse("projekt_file_upload", args=[self.projekt.pk])
        resp = self.client.post(url, {"upload": upload}, format="multipart", HTTP_HX_REQUEST="true")
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.headers.get("X-Upload-Status"), "assigned")
        self.assertTrue(self.projekt.anlagen.filter(anlage_nr=2).exists())

    def test_manual_assignment_flow(self):
        doc = Document()
        doc.add_paragraph("x")
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("foo.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)

        url = reverse("projekt_file_upload", args=[self.projekt.pk])
        resp = self.client.post(url, {"upload": upload}, format="multipart", HTTP_HX_REQUEST="true")
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.headers.get("X-Upload-Status"), "manual")
        self.assertIn("form", resp.content.decode())
        session = self.client.session
        temp_id = next(iter(session.get("pending_uploads", {})))

        resp2 = self.client.post(
            url,
            {"temp_id": temp_id, "anlage_nr": 3},
            format="multipart",
            HTTP_HX_REQUEST="true",
        )
        self.assertEqual(resp2.status_code, 200)
        self.assertEqual(resp2.headers.get("X-Upload-Status"), "assigned")
        self.assertTrue(self.projekt.anlagen.filter(anlage_nr=3).exists())



class AutoApprovalTests(NoesisTestCase):
    """Tests für die automatische Genehmigung von Dokumenten."""

    def setUp(self) -> None:
        self.user = User.objects.create_user("auto", password="pass")
        self.client.login(username="auto", password="pass")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")

    def _upload_doc(self, document: Document) -> BVProjectFile:
        """Hilfsfunktion zum Hochladen eines DOCX-Dokuments."""
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        document.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("Anlage_1.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)
        url = reverse("hx_project_file_upload", args=[self.projekt.pk])
        resp = self.client.post(
            url,
            {"anlage_nr": 1, "upload": upload, "manual_comment": ""},
            format="multipart",
            HTTP_HX_REQUEST="true",
        )
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.content.decode().count("<tr"), 1)
        return self.projekt.anlagen.get(anlage_nr=1)

    def test_single_page_auto_approved(self):
        doc = Document()
        doc.add_paragraph("Seite 1")
        pf = self._upload_doc(doc)
        self.assertFalse(pf.manual_reviewed)
        self.assertFalse(pf.verhandlungsfaehig)

    def test_multi_page_requires_manual_review(self):
        img = Image.new("RGB", (10, 10), color="red")
        img_tmp = NamedTemporaryFile(delete=False, suffix=".png")
        img.save(img_tmp.name)
        img_tmp.close()

        doc = Document()
        doc.add_paragraph("Seite 1")
        doc.add_page_break()
        doc.add_paragraph("Seite 2")
        doc.add_picture(img_tmp.name)
        Path(img_tmp.name).unlink(missing_ok=True)

        pf = self._upload_doc(doc)
        self.assertFalse(pf.manual_reviewed)
        self.assertFalse(pf.verhandlungsfaehig)

    def test_toggle_manual_review_sets_flag(self):
        doc = Document()
        doc.add_paragraph("Seite 1")
        doc.add_page_break()
        doc.add_paragraph("Seite 2")
        pf = self._upload_doc(doc)

        url = reverse("project_file_toggle_flag", args=[pf.pk, "manual_reviewed"])
        resp = self.client.post(url, {"value": "1"})
        self.assertEqual(resp.status_code, 302)
        pf.refresh_from_db()
        self.assertTrue(pf.manual_reviewed)
        self.assertFalse(pf.verhandlungsfaehig)

    def test_project_status_updated_after_all_anlage3_reviewed(self):
        pf1 = BVProjectFile.objects.create(
            project=self.projekt,
            anlage_nr=3,
            upload=SimpleUploadedFile("a.txt", b"x"),
            text_content="x",
        )
        pf2 = BVProjectFile.objects.create(
            project=self.projekt,
            anlage_nr=3,
            upload=SimpleUploadedFile("b.txt", b"x"),
            text_content="y",
        )

        for pf in (pf1, pf2):
            url = reverse("project_file_toggle_flag", args=[pf.pk, "manual_reviewed"])
            self.client.post(url, {"value": "1"})

        self.projekt.refresh_from_db()
        self.assertEqual(self.projekt.status.key, "ENDGEPRUEFT")


class Anlage3AutomationTests(NoesisTestCase):
    def setUp(self) -> None:
        self.user = User.objects.create_user("auto3", password="pass")
        self.client.login(username="auto3", password="pass")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")

    def _upload_docx(self, document: Document) -> BVProjectFile:
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        document.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("Anlage_1.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)
        url = reverse("hx_project_file_upload", args=[self.projekt.pk])
        resp = self.client.post(
            url,
            {"anlage_nr": 3, "upload": upload, "manual_comment": ""},
            format="multipart",
            HTTP_HX_REQUEST="true",
        )
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.content.decode().count("<tr"), 1)
        return BVProjectFile.objects.get(project=self.projekt, anlage_nr=3)

    def test_single_page_sets_negotiable(self):
        doc = Document()
        doc.add_paragraph("Seite 1")
        pf = self._upload_docx(doc)
        self.assertTrue(pf.verhandlungsfaehig)

    def test_review_save_marks_checked(self):
        pf = BVProjectFile.objects.create(
            project=self.projekt,
            anlage_nr=3,
            upload=SimpleUploadedFile("a.docx", b""),
            text_content="",
        )
        url = reverse("projekt_file_edit_json", args=[pf.pk])
        resp = self.client.post(url, {"analysis_json": "{}"})
        self.assertRedirects(resp, reverse("projekt_detail", args=[self.projekt.pk]))
        pf.refresh_from_db()
        self.assertTrue(pf.manual_reviewed)


class BVProjectModelTests(NoesisTestCase):
    def test_title_auto_set_from_software(self):
        projekt = BVProject.objects.create(software_typen="A, B", beschreibung="x")
        self.assertEqual(projekt.title, "A, B")

    def test_title_preserved_when_set(self):
        projekt = BVProject.objects.create(
            title="X", software_typen="A", beschreibung="x"
        )
        self.assertEqual(projekt.title, "X")

    def test_save_accepts_list_for_software_typen(self):
        projekt = BVProject.objects.create(
            software_typen=["A", "", "B "],
            beschreibung="x",
        )
        self.assertEqual(projekt.software_typen, "A, B")


class AnlagenFunktionsMetadatenModelTests(NoesisTestCase):
    def test_manual_result_field(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        func = Anlage2Function.objects.get(name="Anmelden")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"x"),
        )
        res = AnlagenFunktionsMetadaten.objects.create(
            anlage_datei=pf,
            funktion=func,
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=pf,
            funktion=func,
            quelle="manuell",
            technisch_verfuegbar=True,
            ki_beteiligung=False,
        )
        latest = FunktionsErgebnis.objects.filter(
            anlage_datei__project=projekt,
            funktion=func,
            quelle="manuell",
        ).first()
        self.assertTrue(latest.technisch_verfuegbar)
        self.assertFalse(latest.ki_beteiligung)


class WorkflowTests(NoesisTestCase):
    def test_default_status(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        self.assertEqual(projekt.status.key, "NEW")

    def test_set_project_status(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        set_project_status(projekt, "CLASSIFIED")
        projekt.refresh_from_db()
        self.assertEqual(projekt.status.key, "CLASSIFIED")

    def test_invalid_status(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        with self.assertRaises(ValueError):
            set_project_status(projekt, "XXX")

    def test_set_project_status_new_states(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        for status in [
            "IN_PRUEFUNG_ANLAGE_X",
            "FB_IN_PRUEFUNG",
            "ENDGEPRUEFT",
        ]:
            set_project_status(projekt, status)
            projekt.refresh_from_db()
            self.assertEqual(projekt.status.key, status)

    def test_status_history_created(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        self.assertEqual(projekt.status_history.count(), 1)
        set_project_status(projekt, "CLASSIFIED")
        self.assertEqual(projekt.status_history.count(), 2)


class BuildRowDataTests(NoesisTestCase):
    def setUp(self):
        self.func = Anlage2Function.objects.create(name="Testfunktion")
        self.form = Anlage2ReviewForm()

    def test_flag_set_on_difference(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"x"),
        )
        res = AnlagenFunktionsMetadaten.objects.create(
            anlage_datei=pf,
            funktion=self.func,
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=pf,
            funktion=self.func,
            quelle="parser",
            technisch_verfuegbar=True,
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=pf,
            funktion=self.func,
            quelle="ki",
            technisch_verfuegbar=False,
        )
        result_map = {res.get_lookup_key(): res}

        row = _build_row_data(
            "Testfunktion",
            "Testfunktion",
            self.func.id,
            f"func{self.func.id}_",
            self.form,
            {},
            {},
            {},
            {},
            result_map,
        )
        self.assertTrue(row["requires_manual_review"])

    def test_flag_not_set_when_manual(self):
        manual = {"Testfunktion": {"technisch_vorhanden": True}}
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"x"),
        )
        res = AnlagenFunktionsMetadaten.objects.create(
            anlage_datei=pf,
            funktion=self.func,
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=pf,
            funktion=self.func,
            quelle="parser",
            technisch_verfuegbar=True,
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=pf,
            funktion=self.func,
            quelle="ki",
            technisch_verfuegbar=False,
        )
        result_map = {res.get_lookup_key(): res}

        row = _build_row_data(
            "Testfunktion",
            "Testfunktion",
            self.func.id,
            f"func{self.func.id}_",
            self.form,
            {},
            {},
            {},
            manual,
            result_map,
        )
        self.assertFalse(row["requires_manual_review"])

    def test_manual_flags_propagated(self):
        manual = {"Testfunktion": {"technisch_vorhanden": True}}
        row = _build_row_data(
            "Testfunktion",
            "Testfunktion",
            self.func.id,
            f"func{self.func.id}_",
            self.form,
            {},
            {},
            {},
            manual,
            {},
        )
        self.assertTrue(row["manual_flags"]["technisch_vorhanden"])

    def test_manual_flags_false_when_absent(self):
        row = _build_row_data(
            "Testfunktion",
            "Testfunktion",
            self.func.id,
            f"func{self.func.id}_",
            self.form,
            {},
            {},
            {},
            {},
            {},
        )
        self.assertFalse(row["manual_flags"]["technisch_vorhanden"])

    def test_doc_ai_from_result_map_main(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"x"),
        )
        res = AnlagenFunktionsMetadaten.objects.create(
            anlage_datei=pf,
            funktion=self.func,
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=pf,
            funktion=self.func,
            quelle="parser",
            technisch_verfuegbar=True,
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=pf,
            funktion=self.func,
            quelle="ki",
            technisch_verfuegbar=False,
        )
        result_map = {res.get_lookup_key(): res}

        row = _build_row_data(
            self.func.name,
            self.func.name,
            self.func.id,
            f"func{self.func.id}_",
            self.form,
            {},
            {},
            {},
            {},
            result_map,
        )

        self.assertTrue(row["doc_result"]["technisch_vorhanden"])
        self.assertFalse(row["ai_result"]["technisch_vorhanden"])

    def test_doc_ai_from_result_map_subquestion(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        sub = Anlage2SubQuestion.objects.create(
            funktion=self.func, frage_text="Unterfrage?"
        )
        form = Anlage2ReviewForm()  # Formular nach dem Erstellen der Unterfrage
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"x"),
        )
        res = AnlagenFunktionsMetadaten.objects.create(
            anlage_datei=pf,
            funktion=self.func,
            subquestion=sub,
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=pf,
            funktion=self.func,
            subquestion=sub,
            quelle="parser",
            technisch_verfuegbar=False,
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=pf,
            funktion=self.func,
            subquestion=sub,
            quelle="ki",
            technisch_verfuegbar=True,
        )

        lookup = res.get_lookup_key()
        result_map = {lookup: res}

        row = _build_row_data(
            sub.frage_text,
            lookup,
            self.func.id,
            f"sub{sub.id}_",
            form,
            {},
            {},
            {},
            {},
            result_map,
            sub_id=sub.id,
        )

        self.assertFalse(row["doc_result"]["technisch_vorhanden"])
        self.assertTrue(row["ai_result"]["technisch_vorhanden"])


class LLMTasksTests(NoesisTestCase):
    maxDiff = None

    def test_classify_system(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content="Testtext",
        )
        with patch(
            "core.llm_tasks.query_llm",
            return_value='{"kategorie":"X","begruendung":"ok"}',
        ):
            data = classify_system(projekt.pk)
        projekt.refresh_from_db()
        self.assertEqual(projekt.classification_json["kategorie"]["value"], "X")
        self.assertEqual(projekt.status.key, "CLASSIFIED")
        self.assertEqual(data["kategorie"]["value"], "X")

    def test_check_anlage2(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content="Anlagetext",
        )
        func = Anlage2Function.objects.get(name="Anmelden")
        llm_reply = json.dumps({"technisch_verfuegbar": True})
        with patch("core.llm_tasks.query_llm", return_value=llm_reply) as mock_q:
            data = check_anlage2(projekt.pk)
        mock_q.assert_called()
        file_obj = projekt.anlagen.get(anlage_nr=2)
        self.assertTrue(data["functions"][0]["technisch_verfuegbar"])
        self.assertEqual(data["functions"][0]["source"], "llm")
        res = AnlagenFunktionsMetadaten.objects.get(anlage_datei=pf, funktion=func)
        fe = FunktionsErgebnis.objects.filter(
            anlage_datei=pf, funktion=func, quelle="llm"
        ).first()
        self.assertIsNotNone(fe)
        self.assertTrue(fe.technisch_verfuegbar)

    def test_check_anlage2_functions_stores_result(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"data"),
        )
        func = Anlage2Function.objects.get(name="Anmelden")
        llm_reply = json.dumps({"technisch_verfuegbar": True})
        with (
            patch("core.llm_tasks.query_llm", return_value=llm_reply),
            patch("core.llm_tasks.async_task") as mock_async,
            patch("core.llm_tasks.result") as mock_result,
        ):
            mock_async.side_effect = lambda name, *a, **k: (
                worker_verify_feature(*a, **k) or "tid"
            )
            mock_result.side_effect = lambda *a, **k: None
            run_conditional_anlage2_check(pf.pk)

        res = AnlagenFunktionsMetadaten.objects.get(anlage_datei=pf, funktion=func)
        fe = FunktionsErgebnis.objects.filter(
            anlage_datei=pf, funktion=func, quelle="ki"
        ).first()
        self.assertIsNotNone(fe)
        self.assertTrue(fe.technisch_verfuegbar)

    def test_check_anlage2_llm_receives_text(self):
        """Der LLM-Prompt enthält den bekannten Text."""
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content="Testinhalt Anlage2",
        )
        func = Anlage2Function.objects.get(name="Anmelden")
        llm_reply = json.dumps({"technisch_verfuegbar": False})
        with patch("core.llm_tasks.query_llm", return_value=llm_reply) as mock_q:
            data = check_anlage2(projekt.pk)
        self.assertIn("Testinhalt Anlage2", mock_q.call_args_list[0].args[0].text)
        file_obj = projekt.anlagen.get(anlage_nr=2)
        self.assertEqual(data["functions"][0]["funktion"], "Anmelden")

    def test_check_anlage2_prompt_contains_text(self):
        """Der Prompt enth\u00e4lt den gesamten Anlagentext."""
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content="Testinhalt Anlage2",
        )
        func = Anlage2Function.objects.get(name="Anmelden")
        llm_reply = json.dumps({"technisch_verfuegbar": False})
        with patch("core.llm_tasks.query_llm", return_value=llm_reply) as mock_q:
            data = check_anlage2(projekt.pk)
        prompt = mock_q.call_args_list[0].args[0].text
        self.assertIn("Testinhalt Anlage2", prompt)
        file_obj = projekt.anlagen.get(anlage_nr=2)
        self.assertEqual(data["functions"][0]["funktion"], "Anmelden")

    def test_check_anlage2_parser(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        doc = Document()
        table = doc.add_table(rows=2, cols=5)
        table.cell(0, 0).text = "Funktion"
        table.cell(0, 1).text = "Technisch vorhanden"
        table.cell(0, 2).text = "Einsatz bei Telefónica"
        table.cell(0, 3).text = "Zur LV-Kontrolle"
        table.cell(0, 4).text = "KI-Beteiligung"
        table.cell(1, 0).text = "Anmelden"
        table.cell(1, 1).text = "Ja"
        table.cell(1, 2).text = "Nein"
        table.cell(1, 3).text = "Nein"
        table.cell(1, 4).text = "Ja"
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("b.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)
        BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=upload,
            text_content="ignored",
        )
        func = Anlage2Function.objects.get(name="Anmelden")

        with patch("core.llm_tasks.query_llm") as mock_q:
            data = check_anlage2(projekt.pk)
        mock_q.assert_called()
        expected = {
            "task": "check_anlage2",
            "functions": [
                {
                    "funktion": "Anmelden",
                    "technisch_verfuegbar": {"value": True, "note": None},
                    "ki_beteiligung": {"value": True, "note": None},
                    "source": "parser",
                }
            ],
        }
        file_obj = projekt.anlagen.get(anlage_nr=2)
        self.assertEqual(data, expected)
        self.assertEqual(file_obj.analysis_json, expected)

    def test_run_anlage2_analysis_table(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")

        doc = Document()
        table = doc.add_table(rows=2, cols=5)
        table.cell(0, 0).text = "Funktion"
        table.cell(0, 1).text = "Technisch vorhanden"
        table.cell(0, 2).text = "Einsatz bei Telefónica"
        table.cell(0, 3).text = "Zur LV-Kontrolle"
        table.cell(0, 4).text = "KI-Beteiligung"
        table.cell(1, 0).text = "Anmelden"
        table.cell(1, 1).text = "Ja"
        table.cell(1, 2).text = "Nein"
        table.cell(1, 3).text = "Nein"
        table.cell(1, 4).text = "Ja"

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        upload = SimpleUploadedFile("b.docx", buffer.read())

        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=upload,
            text_content="Anmelden: tv: ja; tel: nein; lv: nein; ki: ja",
        )
        func = Anlage2Function.objects.get(name="Anmelden")
        cfg = Anlage2Config.get_instance()
        cfg.parser_mode = "table_only"
        cfg.parser_order = ["table"]
        cfg.text_technisch_verfuegbar_true = ["ja"]
        cfg.text_technisch_verfuegbar_false = []
        cfg.text_einsatz_telefonica_true = []
        cfg.text_einsatz_telefonica_false = ["nein"]
        cfg.text_zur_lv_kontrolle_true = []
        cfg.text_zur_lv_kontrolle_false = ["nein"]
        cfg.text_ki_beteiligung_true = ["ja"]
        cfg.text_ki_beteiligung_false = []
        cfg.save()
        AntwortErkennungsRegel.objects.all().delete()

        result = run_anlage2_analysis(pf)
        expected = [
            {
                "funktion": "Anmelden",
                "technisch_verfuegbar": {"value": True, "note": None},
                "einsatz_telefonica": {"value": False, "note": None},
                "zur_lv_kontrolle": {"value": False, "note": None},
                "ki_beteiligung": {"value": True, "note": None},
            }
        ]

        pf.refresh_from_db()
        fe = FunktionsErgebnis.objects.filter(
            anlage_datei=pf, funktion=func, quelle="parser"
        ).first()
        self.assertIsNotNone(fe)
        self.assertTrue(fe.technisch_verfuegbar)

        login_entry = next(
            f for f in pf.analysis_json["functions"] if f["funktion"] == "Anmelden"
        )
        self.assertTrue(login_entry["technisch_verfuegbar"]["value"])
        self.assertFalse(login_entry["einsatz_telefonica"]["value"])
        self.assertFalse(login_entry["zur_lv_kontrolle"]["value"])
        self.assertTrue(login_entry["ki_beteiligung"]["value"])

        self.assertIsInstance(result, list)
        self.assertTrue(any(r["funktion"] == "Anmelden" for r in result))

    def test_run_anlage2_analysis_sets_negotiable_on_match(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        content = "Anmelden: tv: ja; tel: nein; lv: nein; ki: ja"
        upload = SimpleUploadedFile("b.txt", b"x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=upload,
            text_content=content,
        )
        func = Anlage2Function.objects.get(name="Anmelden")
        cfg = Anlage2Config.get_instance()
        cfg.text_technisch_verfuegbar_true = ["ja"]
        cfg.save()
        AnlagenFunktionsMetadaten.objects.create(
            anlage_datei=pf,
            funktion=func,
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=pf,
            funktion=func,
            quelle="ki",
            technisch_verfuegbar=True,
        )

        run_anlage2_analysis(pf)

        parser_fe = FunktionsErgebnis.objects.filter(
            anlage_datei=pf, funktion=func, quelle="parser"
        ).first()
        ai_fe = FunktionsErgebnis.objects.filter(
            anlage_datei=pf, funktion=func, quelle="ki"
        ).first()

        self.assertTrue(parser_fe.technisch_verfuegbar)
        self.assertTrue(ai_fe.technisch_verfuegbar)

    def test_parser_manager_fallback(self):
        class FailParser(AbstractParser):
            name = "fail"

            def parse(self, project_file):
                raise ValueError("boom")

        class DummyParser(AbstractParser):
            name = "dummy"

            def parse(self, project_file):
                return [{"funktion": "Dummy"}]

        parser_manager.register(FailParser)
        parser_manager.register(DummyParser)
        cfg = Anlage2Config.get_instance()
        cfg.parser_order = ["fail", "dummy"]
        cfg.parser_mode = "manager"
        cfg.save()

        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("c.docx", fh.read())
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=upload,
        )

        try:
            result = parser_manager.parse_anlage2(pf)
        finally:
            Path(tmp.name).unlink(missing_ok=True)
            parser_manager._parsers.pop("fail")
            parser_manager._parsers.pop("dummy")
            cfg.parser_order = ["table"]
            cfg.parser_mode = "auto"
            cfg.save()

        self.assertEqual(result, [{"funktion": "Dummy"}])

    def test_parser_manager_order(self):
        class P1(AbstractParser):
            name = "one"

            def parse(self, project_file):
                return [{"val": 1}]

        class P2(AbstractParser):
            name = "two"

            def parse(self, project_file):
                return [{"val": 2}]

        parser_manager.register(P1)
        parser_manager.register(P2)
        cfg = Anlage2Config.get_instance()
        cfg.parser_order = ["two", "one"]
        cfg.parser_mode = "manager"
        cfg.save()

        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("d.docx", fh.read())
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=upload,
        )
        try:
            result = parser_manager.parse_anlage2(pf)
        finally:
            Path(tmp.name).unlink(missing_ok=True)
            parser_manager._parsers.pop("one")
            parser_manager._parsers.pop("two")
            cfg.parser_order = ["table"]
            cfg.parser_mode = "auto"
            cfg.save()

        self.assertEqual(result, [{"val": 2}])

    def test_parser_manager_uses_first_result(self):
        class P1(AbstractParser):
            name = "p1"

            def parse(self, project_file):
                return [{"funktion": "A", "technisch_verfuegbar": {"value": False}}]

        class P2(AbstractParser):
            name = "p2"

            def parse(self, project_file):
                return [{"funktion": "A", "technisch_verfuegbar": {"value": True}}]

        parser_manager.register(P1)
        parser_manager.register(P2)
        cfg = Anlage2Config.get_instance()
        cfg.parser_order = ["p1", "p2"]
        cfg.parser_mode = "manager"
        cfg.save()

        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        doc = Document()
        doc.add_table(rows=1, cols=1)
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("e.docx", fh.read())
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=upload,
        )

        try:
            with (
                patch("core.parsers.parse_anlage2_table", return_value=[]) as m_tab,
                patch("core.text_parser.parse_anlage2_text", return_value=[]) as m_text,
            ):
                result = parser_manager.parse_anlage2(pf)
        finally:
            Path(tmp.name).unlink(missing_ok=True)
            parser_manager._parsers.pop("p1", None)
            parser_manager._parsers.pop("p2", None)
            cfg.parser_order = ["table"]
            cfg.parser_mode = "auto"
            cfg.save()

        m_tab.assert_not_called()
        m_text.assert_not_called()
        self.assertFalse(result[0]["technisch_verfuegbar"]["value"])

    def test_run_anlage2_analysis_auto_prefers_table(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"x"),
            text_content="t",
        )
        cfg = Anlage2Config.get_instance()

        cfg.parser_mode = "auto"
        cfg.parser_order = ["table", "exact"]
        cfg.save()
        table_result = [{"funktion": "Anmelden"}]
        with (
            patch("core.parsers.parse_anlage2_table", return_value=table_result),
            patch(
                "core.parsers.ExactParser.parse", return_value=[{"funktion": "Alt"}]
            ) as m_exact,
        ):
            result = parser_manager.parse_anlage2(pf)
        self.assertEqual(result, table_result)

    def test_run_anlage2_analysis_auto_fallback_empty_table(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"x"),
            text_content="t",
        )
        cfg = Anlage2Config.get_instance()
        cfg.parser_mode = "auto"
        cfg.parser_order = ["exact", "table"]
        cfg.save()
        with (
            patch("core.parsers.parse_anlage2_table", return_value=[{"funktion": "Anmelden"}]) as m_table,
            patch("core.parsers.ExactParser.parse", return_value=[]) as m_exact,
        ):
            result = parser_manager.parse_anlage2(pf)
        m_exact.assert_called_once()
        m_table.assert_called_once()
        self.assertEqual(result, [{"funktion": "Anmelden"}])

    def test_parser_manager_exact_parser_segments(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"x"),
            text_content="Alpha: aktiv\nBeta: kein einsatz",
        )
        Anlage2Function.objects.create(name="Alpha")
        Anlage2Function.objects.create(name="Beta")
        AntwortErkennungsRegel.objects.create(
            regel_name="aktiv",
            erkennungs_phrase="aktiv",
            actions_json=[{"field": "technisch_verfuegbar", "value": True}],
        )
        AntwortErkennungsRegel.objects.create(
            regel_name="einsatz",
            erkennungs_phrase="kein einsatz",
            actions_json=[{"field": "einsatz_telefonica", "value": False}],
        )
        cfg = Anlage2Config.get_instance()
        cfg.parser_mode = "exact_only"
        cfg.save()
        result = parser_manager.parse_anlage2(pf)
        self.assertEqual(
            result,
            [
                {
                    "funktion": "Alpha",
                    "technisch_verfuegbar": {"value": True, "note": None},
                },
                {
                    "funktion": "Beta",
                    "einsatz_telefonica": {"value": False, "note": None},
                },
            ],
        )

    def test_run_anlage2_analysis_includes_missing_functions(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"x"),
            text_content="",
        )
        func = Anlage2Function.objects.get(name="Anmelden")

        result = run_anlage2_analysis(pf)

        self.assertEqual(len(result), 1)
        self.assertEqual(result[0]["funktion"], "Anmelden")
        self.assertTrue(
            result[0].get("not_found") or result[0].get("technisch_verfuegbar") is None
        )
        pf.refresh_from_db()
        fe = FunktionsErgebnis.objects.filter(
            anlage_datei=pf, funktion=func, quelle="parser"
        ).first()
        self.assertIsNotNone(fe)
        self.assertIsNone(fe.technisch_verfuegbar)

    def test_run_anlage2_analysis_includes_missing_subquestions(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"x"),
            text_content="",
        )
        func = Anlage2Function.objects.get(name="Anmelden")
        Anlage2SubQuestion.objects.filter(funktion=func).delete()
        Anlage2SubQuestion.objects.create(funktion=func, frage_text="Warum?")

        result = run_anlage2_analysis(pf)

        self.assertEqual(len(result), 2)
        names = [row["funktion"] for row in result]
        self.assertIn("Anmelden", names)
        self.assertTrue(any("Warum?" in n for n in names))
        pf.refresh_from_db()
        parser_res = FunktionsErgebnis.objects.filter(
            anlage_datei=pf, funktion=func, quelle="parser"
        )
        self.assertEqual(parser_res.count(), 2)

    def test_run_anlage2_analysis_fuzzy_match(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"x"),
            text_content="Logn: ja",
        )
        func = Anlage2Function.objects.get(name="Anmelden")
        cfg = Anlage2Config.get_instance()
        cfg.text_technisch_verfuegbar_true = ["ja"]
        cfg.save()

        result = run_anlage2_analysis(pf)

        expected = [{
            "funktion": "Anmelden",
            "not_found": True,
            "technisch_verfuegbar": None,
            "einsatz_telefonica": None,
            "zur_lv_kontrolle": None,
            "ki_beteiligung": None,
        }]
        self.assertEqual(result, expected)

    def test_run_anlage2_analysis_sets_complete_status_without_followup(self):
        """Status wird nur ohne anschließende KI-Prüfung auf COMPLETE gesetzt."""

        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"x"),
            text_content="",
            processing_status=BVProjectFile.PROCESSING,
        )
        func = Anlage2Function.objects.get(name="Anmelden")
        FunktionsErgebnis.objects.create(
            anlage_datei=pf,
            funktion=func,
            quelle="ki",
            technisch_verfuegbar=True,
        )

        run_anlage2_analysis(pf)

        pf.refresh_from_db()
        self.assertEqual(pf.processing_status, BVProjectFile.COMPLETE)

    def test_run_anlage2_analysis_keeps_processing_with_followup(self):
        """Bei ausstehender KI-Prüfung bleibt der Status PROCESSING."""

        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"x"),
            text_content="",
            processing_status=BVProjectFile.PROCESSING,
        )
        Anlage2Function.objects.get(name="Anmelden")

        run_anlage2_analysis(pf)

        pf.refresh_from_db()
        self.assertEqual(pf.processing_status, BVProjectFile.PROCESSING)

    def test_check_anlage2_table_error_fallback(self):
        class P1(AbstractParser):
            name = "p1"

            def parse(self, project_file):
                return [{"funktion": "A", "technisch_verfuegbar": {"value": False}}]

        class P2(AbstractParser):
            name = "p2"

            def parse(self, project_file):
                return [{"funktion": "A", "technisch_verfuegbar": {"value": True}}]

        parser_manager.register(P1)
        parser_manager.register(P2)
        cfg = Anlage2Config.get_instance()
        cfg.parser_order = ["p1", "p2"]
        cfg.save()

        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("b.txt", b"x"),
        )

        try:
            result = parser_manager.parse_anlage2(pf)
        finally:
            parser_manager._parsers.pop("p1", None)
            parser_manager._parsers.pop("p2", None)
            cfg.parser_order = ["table"]
            cfg.save()

        self.assertFalse(result[0]["technisch_verfuegbar"]["value"])

    def test_analyse_anlage3_auto_ok(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        doc = Document()
        doc.add_paragraph("Seite 1")
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("c.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)
        BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=3,
            upload=upload,
            text_content="ignored",
        )

        pf = projekt.anlagen.get(anlage_nr=3)
        data = analyse_anlage3(pf.pk)
        pf.refresh_from_db()
        file_obj = pf
        self.assertEqual(data["pages"]["value"], 1)
        self.assertTrue(data["verhandlungsfaehig"]["value"])
        self.assertTrue(file_obj.analysis_json["verhandlungsfaehig"]["value"])
        if hasattr(file_obj, "verhandlungsfaehig"):
            self.assertTrue(file_obj.verhandlungsfaehig)

    def test_analyse_anlage3_manual_required(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        doc = Document()
        doc.add_paragraph("Seite 1")
        doc.add_page_break()
        doc.add_paragraph("Seite 2")
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("d.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)
        BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=3,
            upload=upload,
            text_content="ignored",
        )

        pf = projekt.anlagen.get(anlage_nr=3)
        data = analyse_anlage3(pf.pk)
        pf.refresh_from_db()
        file_obj = pf
        self.assertEqual(data["pages"]["value"], 2)
        self.assertFalse(data["verhandlungsfaehig"]["value"])
        self.assertFalse(file_obj.analysis_json["verhandlungsfaehig"]["value"])
        if hasattr(file_obj, "verhandlungsfaehig"):
            self.assertFalse(file_obj.verhandlungsfaehig)

    def test_analyse_anlage3_pdf_auto_ok(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pdf = fitz.open()
        pdf.new_page()
        tmp = NamedTemporaryFile(delete=False, suffix=".pdf")
        tmp.close()
        pdf.save(tmp.name)
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("c.pdf", fh.read())
        Path(tmp.name).unlink(missing_ok=True)
        BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=3,
            upload=upload,
            text_content="ignored",
        )

        pf = projekt.anlagen.get(anlage_nr=3)
        data = analyse_anlage3(pf.pk)
        pf.refresh_from_db()
        file_obj = pf
        self.assertEqual(data["pages"]["value"], 1)
        self.assertTrue(data["verhandlungsfaehig"]["value"])
        self.assertTrue(file_obj.analysis_json["verhandlungsfaehig"]["value"])

    def test_analyse_anlage3_pdf_manual_required(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pdf = fitz.open()
        pdf.new_page()
        pdf.new_page()
        tmp = NamedTemporaryFile(delete=False, suffix=".pdf")
        tmp.close()
        pdf.save(tmp.name)
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("d.pdf", fh.read())
        Path(tmp.name).unlink(missing_ok=True)
        BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=3,
            upload=upload,
            text_content="ignored",
        )

        pf = projekt.anlagen.get(anlage_nr=3)
        data = analyse_anlage3(pf.pk)
        pf.refresh_from_db()
        file_obj = pf
        self.assertEqual(data["pages"]["value"], 2)
        self.assertFalse(data["verhandlungsfaehig"]["value"])
        self.assertFalse(file_obj.analysis_json["verhandlungsfaehig"]["value"])

    def test_analyse_anlage3_multiple_files(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")

        doc1 = Document()
        doc1.add_paragraph("Seite 1")
        tmp1 = NamedTemporaryFile(delete=False, suffix=".docx")
        doc1.save(tmp1.name)
        tmp1.close()
        with open(tmp1.name, "rb") as fh:
            upload1 = SimpleUploadedFile("e.docx", fh.read())
        Path(tmp1.name).unlink(missing_ok=True)
        pf1 = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=3,
            upload=upload1,
            text_content="x",
        )

        doc2 = Document()
        doc2.add_paragraph("Seite 1")
        doc2.add_page_break()
        doc2.add_paragraph("Seite 2")
        tmp2 = NamedTemporaryFile(delete=False, suffix=".docx")
        doc2.save(tmp2.name)
        tmp2.close()
        with open(tmp2.name, "rb") as fh:
            upload2 = SimpleUploadedFile("f.docx", fh.read())
        Path(tmp2.name).unlink(missing_ok=True)
        pf2 = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=3,
            upload=upload2,
            text_content="y",
        )

        analyse_anlage3(pf1.pk)
        pf1.refresh_from_db()
        pf2.refresh_from_db()
        self.assertIsNotNone(pf1.analysis_json)
        self.assertIsNotNone(pf2.analysis_json)

    def test_check_anlage3_vision_stores_json(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        doc = Document()
        doc.add_paragraph("A")
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("g.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)
        BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=3,
            upload=upload,
            text_content="ignored",
        )

        llm_reply = json.dumps({"ok": True, "hinweis": "x"})
        with patch("core.llm_tasks.query_llm_with_images", return_value=llm_reply):
            data = check_anlage3_vision(projekt.pk)
        file_obj = projekt.anlagen.get(anlage_nr=3)
        self.assertTrue(data["ok"]["value"])
        self.assertTrue(file_obj.analysis_json["ok"]["value"])

    def test_check_anlage1_parser(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        text = (
            "Frage 1: Extrahiere alle Unternehmen als Liste.\u00b6A1\u00b6"
            "Frage 2: Extrahiere alle Fachbereiche als Liste.\u00b6A2"
        )
        BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content=text,
        )
        file_obj = projekt.anlagen.get(anlage_nr=1)
        data = check_anlage1(file_obj.pk)
        expected = {"questions": parse_anlage1_questions(text)}
        self.assertEqual(data, expected)
        self.assertEqual(file_obj.analysis_json, expected)

    def test_parse_anlage1_questions_extra(self):
        Anlage1Question.objects.create(
            num=10,
            text="Frage 10: Test?",
            enabled=True,
            parser_enabled=True,
            llm_enabled=True,
        )
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        text = (
            "Frage 1: Extrahiere alle Unternehmen als Liste.\u00b6A1\u00b6"
            "Frage 10: Test?\u00b6A10"
        )
        BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content=text,
        )
        file_obj = projekt.anlagen.get(anlage_nr=1)
        data = check_anlage1(file_obj.pk)
        q_data = data["questions"]
        self.assertEqual(q_data["10"]["answer"], "A10")

    def test_parse_anlage1_questions_without_numbers(self):
        """Prüft die Extraktion ohne nummerierte Fragen."""
        # Frage-Texte ohne Präfix "Frage X:" speichern
        q1 = Anlage1Question.objects.get(num=1)
        q2 = Anlage1Question.objects.get(num=2)
        q1.text = q1.text.split(": ", 1)[1]
        q2.text = q2.text.split(": ", 1)[1]
        q1.save(update_fields=["text"])
        q2.save(update_fields=["text"])
        v1 = q1.variants.first()
        v2 = q2.variants.first()
        v1.text = q1.text
        v2.text = q2.text
        v1.save()
        v2.save()

        text = f"{q1.text}\u00b6A1\u00b6{q2.text}\u00b6A2"
        parsed = parse_anlage1_questions(text)
        self.assertEqual(
            parsed,
            {
                "1": {"answer": "A1", "found_num": None},
                "2": {"answer": "A2", "found_num": None},
            },
        )

    def test_parse_anlage1_questions_with_variant(self):
        q1 = Anlage1Question.objects.get(num=1)
        q1.variants.create(text="Alternative Frage 1?")
        text = "Alternative Frage 1?\u00b6A1"
        parsed = parse_anlage1_questions(text)
        self.assertEqual(parsed, {"1": {"answer": "A1", "found_num": "1"}})

    def test_parse_anlage1_questions_with_newlines(self):
        """Extraktion funktioniert trotz Zeilenumbr\u00fcche."""
        text = (
            "Frage 1:\nExtrahiere alle Unternehmen als Liste.\nA1\n"
            "Frage 2:\nExtrahiere alle Fachbereiche als Liste.\nA2"
        )
        parsed = parse_anlage1_questions(text)
        self.assertEqual(
            parsed,
            {
                "1": {"answer": "A1", "found_num": "1"},
                "2": {"answer": "A2", "found_num": "2"},
            },
        )

    def test_parse_anlage1_questions_split_lines(self):
        """Fragen werden auch mit Zeilenumbruch nach dem Doppelpunkt erkannt."""
        text = (
            "Frage 1:\n"
            "Extrahiere alle Unternehmen als Liste.\r\n"
            "A1\nFrage 2:\n"
            "Extrahiere alle Fachbereiche als Liste.\r\nA2"
        )
        parsed = parse_anlage1_questions(text)
        self.assertEqual(
            parsed,
            {
                "1": {"answer": "A1", "found_num": "1"},
                "2": {"answer": "A2", "found_num": "2"},
            },
        )

    def test_parse_anlage1_questions_respects_parser_enabled(self):
        q2 = Anlage1Question.objects.get(num=2)
        q2.parser_enabled = False
        q2.save(update_fields=["parser_enabled"])
        text = "Frage 1: Extrahiere alle Unternehmen als Liste.\u00b6A1"
        parsed = parse_anlage1_questions(text)
        self.assertEqual(parsed, {"1": {"answer": "A1", "found_num": "1"}})

    def test_parse_anlage1_questions_detects_wrong_number(self):
        """Die erkannte Nummer wird zur\u00fcckgegeben."""
        text = "Frage 1.2: Extrahiere alle Unternehmen als Liste.\u00b6A1"
        parsed = parse_anlage1_questions(text)
        self.assertEqual(parsed, {"1": {"answer": "A1", "found_num": "1.2"}})

    def test_generate_gutachten_twice_replaces_file(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        first = generate_gutachten(projekt.pk, text="Alt")
        second = generate_gutachten(projekt.pk, text="Neu")
        try:
            self.assertTrue(second.exists())
            self.assertNotEqual(first, second)
            self.assertFalse(first.exists())
        finally:
            second.unlink(missing_ok=True)

    def test_parse_anlage2_question_list(self):
        text = "Welche Funktionen bietet das System?\u00b6- Anmelden\u00b6- Suche"
        parsed = _parse_anlage2(text)
        self.assertEqual(parsed, ["Anmelden", "Suche"])

    def test_parse_anlage2_table_llm(self):
        text = "Funktion | Beschreibung\u00b6Anmelden | a\u00b6Suche | b"
        with patch(
            "core.llm_tasks.query_llm", return_value='["Anmelden", "Suche"]'
        ) as mock_q:
            parsed = _parse_anlage2(text)
        mock_q.assert_called_once()
        self.assertEqual(parsed, ["Anmelden", "Suche"])


class PromptTests(NoesisTestCase):
    def test_get_prompt_returns_default(self):
        self.assertEqual(get_prompt("unknown", "foo"), "foo")

    def test_get_prompt_returns_db_value(self):
        p, _ = Prompt.objects.get_or_create(
            name="classify_system", defaults={"text": "orig"}
        )
        p.text = "DB"
        p.save()
        self.assertEqual(get_prompt("classify_system", "x"), "DB")

    def test_check_anlage3_vision_prompt_text(self):
        p = Prompt.objects.get(name="check_anlage3_vision")
        expected = (
            "Pr\u00fcfe die folgenden Bilder der Anlage. "
            "Gib ein JSON mit 'ok' und 'hinweis' zur\u00fcck:\n\n"
        )
        self.assertEqual(p.text, expected)


class CheckAnlage5Tests(NoesisTestCase):
    def test_check_anlage5_sets_flag(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=5,
            upload=SimpleUploadedFile("a.docx", b""),
            text_content="",
        )
        ZweckKategorieA.objects.all().delete()
        cat1 = ZweckKategorieA.objects.create(beschreibung="A")
        cat2 = ZweckKategorieA.objects.create(beschreibung="B")
        text = f"{cat1.beschreibung} {cat2.beschreibung}"
        with patch("core.llm_tasks.extract_text", return_value=text):
            data = check_anlage5(pf.pk)
        pf.refresh_from_db()
        review = pf.anlage5review
        self.assertEqual(set(data["purposes"]), {cat1.pk, cat2.pk})
        self.assertTrue(pf.verhandlungsfaehig)
        self.assertEqual(
            set(review.found_purposes.values_list("pk", flat=True)), {cat1.pk, cat2.pk}
        )

    def test_check_anlage5_detects_other_text(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=5,
            upload=SimpleUploadedFile("a.docx", b""),
            text_content="",
        )
        cat = ZweckKategorieA.objects.create(beschreibung="A")
        text = f"{cat.beschreibung} Sonstige Zwecke zur Leistungs- oder und Verhaltenskontrolle Test"
        with patch("core.llm_tasks.extract_text", return_value=text):
            data = check_anlage5(pf.pk)
        pf.refresh_from_db()
        self.assertFalse(pf.verhandlungsfaehig)
        self.assertEqual(data["sonstige"], "Test")


class PromptImportTests(NoesisTestCase):
    def setUp(self):
        admin_group = Group.objects.create(name="admin")
        self.user = User.objects.create_user("pimport", password="pass")
        self.user.groups.add(admin_group)
        self.client.login(username="pimport", password="pass")

    def test_import_with_clear_first_replaces_prompts(self):
        Prompt.objects.create(name="old", text="x")
        payload = json.dumps([{"name": "neu", "text": "t"}])
        file = SimpleUploadedFile("p.json", payload.encode("utf-8"))
        url = reverse("admin_prompt_import")
        resp = self.client.post(
            url,
            {"json_file": file, "clear_first": "on"},
            format="multipart",
        )
        self.assertRedirects(resp, reverse("admin_prompts"))
        self.assertEqual(Prompt.objects.count(), 1)
        self.assertTrue(Prompt.objects.filter(name="neu").exists())


class ReportingTests(NoesisTestCase):
    def test_gap_analysis_file_created(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content="Testtext",
            analysis_json={"ok": {"value": True, "editable": True}},
        )
        path = generate_gap_analysis(projekt)
        try:
            self.assertTrue(path.exists())
        finally:
            path.unlink(missing_ok=True)


class ProjektFileCheckViewTests(NoesisTestCase):
    def setUp(self):
        self.user = User.objects.create_user("user2", password="pass")
        self.client.login(username="user2", password="pass")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        BVProjectFile.objects.create(
            project=self.projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content="Text",
            manual_analysis_json={"functions": {}},
            analysis_json={},
            verification_json={"functions": {}},
        )

    def test_file_check_endpoint_saves_json(self):
        url = reverse("projekt_file_check", args=[self.projekt.pk, 1])
        resp = self.client.post(url)
        self.assertEqual(resp.status_code, 200)
        resp_json = resp.json()
        file_obj = self.projekt.anlagen.get(anlage_nr=1)
        expected = {"questions": {}}
        self.assertEqual(file_obj.analysis_json, expected)
        self.assertEqual(resp_json["analysis"], expected)

    def test_file_check_pk_endpoint_saves_json(self):
        file_obj = self.projekt.anlagen.get(anlage_nr=1)
        url = reverse("projekt_file_check_pk", args=[file_obj.pk])
        resp = self.client.post(url)
        self.assertEqual(resp.status_code, 200)
        resp_json = resp.json()
        file_obj.refresh_from_db()
        expected = {"questions": {}}
        self.assertEqual(file_obj.analysis_json, expected)
        self.assertEqual(resp_json["analysis"], expected)


class Anlage2ReviewTests(NoesisTestCase):
    def setUp(self):
        self.user = User.objects.create_user("rev", password="pass")
        self.client.login(username="rev", password="pass")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        self.file = BVProjectFile.objects.create(
            project=self.projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("c.txt", b"d"),
            text_content="Text",
            manual_analysis_json={"functions": {}},
            analysis_json={
                "functions": [
                    {
                        "funktion": "Anmelden",
                        "technisch_vorhanden": {"value": True, "note": None},
                        "einsatz_bei_telefonica": {"value": False, "note": None},
                        "zur_lv_kontrolle": {"value": False, "note": None},
                        "ki_beteiligung": {"value": True, "note": None},
                    }
                ]
            },
            verification_json={"functions": {}},
        )
        self.func = Anlage2Function.objects.get(name="Anmelden")
        Anlage2SubQuestion.objects.filter(funktion=self.func).delete()
        self.sub = Anlage2SubQuestion.objects.create(
            funktion=self.func, frage_text="Warum?"
        )

    def test_get_shows_table(self):
        url = reverse("projekt_file_edit_json", args=[self.file.pk])
        resp = self.client.get(url)
        self.assertContains(resp, "Anmelden")
        self.assertContains(resp, "Warum?")
        self.assertContains(resp, f'name="func{self.func.id}_technisch_vorhanden"')

    def test_post_saves_data(self):
        url = reverse("projekt_file_edit_json", args=[self.file.pk])
        resp = self.client.post(
            url,
            {
                f"func{self.func.id}_technisch_vorhanden": "on",
                f"sub{self.sub.id}_ki_beteiligung": "on",
            },
        )
        self.assertRedirects(resp, reverse("projekt_detail", args=[self.projekt.pk]))
        self.file.refresh_from_db()
        data = self.file.manual_analysis_json["functions"][str(self.func.id)]
        self.assertTrue(data["technisch_vorhanden"])
        self.assertTrue(data["subquestions"][str(self.sub.id)]["ki_beteiligung"])

    def test_prefill_from_analysis(self):
        """Die Formulardaten verwenden Analysewerte als Vorgabe."""
        self.file.manual_analysis_json = None
        self.file.analysis_json = {
            "functions": [
                {
                    "funktion": "Anmelden",
                    "technisch_vorhanden": {"value": True, "note": None},
                    "einsatz_bei_telefonica": {"value": True, "note": None},
                    "zur_lv_kontrolle": {"value": True, "note": None},
                }
            ]
        }
        self.file.save()

        url = reverse("projekt_file_edit_json", args=[self.file.pk])
        resp = self.client.get(url)
        field = f"func{self.func.id}_technisch_vorhanden"
        self.assertTrue(resp.context["form"].initial[field])

    def test_prefill_with_metadaten_no_ergebnis(self):
        """Analysewerte werden auch ohne FunktionsErgebnisse angezeigt."""
        self.file.manual_analysis_json = None
        self.file.save()
        AnlagenFunktionsMetadaten.objects.create(
            anlage_datei=self.file,
            funktion=self.func,
        )
        url = reverse("projekt_file_edit_json", args=[self.file.pk])
        resp = self.client.get(url)
        field = f"func{self.func.id}_technisch_vorhanden"
        self.assertTrue(resp.context["form"].initial[field])

    def test_rows_include_lookup_key(self):
        url = reverse("projekt_file_edit_json", args=[self.file.pk])
        resp = self.client.get(url)
        rows = resp.context["rows"]
        self.assertEqual(rows[0]["verif_key"], self.func.name)
        self.assertEqual(
            rows[1]["verif_key"], f"{self.func.name}: {self.sub.frage_text}"
        )

    def test_subquestion_justification_link(self):
        FunktionsErgebnis.objects.create(
            anlage_datei=self.file,
            funktion=self.func,
            subquestion=self.sub,
            quelle="ki",
            begruendung="Text",
        )

        url = reverse("projekt_file_edit_json", args=[self.file.pk])
        resp = self.client.get(url)
        link = reverse(
            "justification_detail_edit",
            args=[self.file.pk, f"{self.func.name}: {self.sub.frage_text}"],
        )
        self.assertContains(resp, link)

    def test_no_auto_analysis_on_get(self):
        pf = BVProjectFile.objects.create(
            project=self.projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("n.txt", b"x"),
            text_content="t",
        )
        url = reverse("projekt_file_edit_json", args=[pf.pk])

        def _fake(obj):
            obj.analysis_json = {"functions": []}
            obj.save(update_fields=["analysis_json"])
            return []

        with patch("core.views.run_anlage2_analysis", side_effect=_fake) as mock_func:
            self.client.get(url)
            self.client.get(url)
        self.assertEqual(mock_func.call_count, 0)


class WorkerGenerateGutachtenTests(NoesisTestCase):
    def setUp(self):
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        BVProjectFile.objects.create(
            project=self.projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content="Text",
            manual_analysis_json={"functions": {}},
            analysis_json={},
            verification_json={"functions": {}},
        )
        self.knowledge = SoftwareKnowledge.objects.create(
            project=self.projekt,
            software_name="A",
            is_known_by_llm=True,
            description="",
        )

    def test_worker_creates_file(self):
        with patch("core.llm_tasks.query_llm", return_value="Text"):
            path = worker_generate_gutachten(self.projekt.pk, self.knowledge.pk)
        self.projekt.refresh_from_db()
        self.assertTrue(self.projekt.gutachten_file.name)
        self.assertEqual(self.projekt.status.key, "GUTACHTEN_OK")
        self.assertEqual(
            Gutachten.objects.filter(software_knowledge=self.knowledge).count(), 1
        )
        Path(path).unlink(missing_ok=True)

    def test_worker_updates_existing_gutachten(self):
        Gutachten.objects.create(software_knowledge=self.knowledge, text="Alt")
        with patch("core.llm_tasks.query_llm", return_value="Neu"):
            path = worker_generate_gutachten(self.projekt.pk, self.knowledge.pk)
        gutachten = Gutachten.objects.get(software_knowledge=self.knowledge)
        self.assertEqual(gutachten.text, "Neu")
        self.assertEqual(
            Gutachten.objects.filter(software_knowledge=self.knowledge).count(), 1
        )
        Path(path).unlink(missing_ok=True)


class WorkerAnlage3VisionTests(NoesisTestCase):
    def setUp(self):
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        doc = Document()
        doc.add_paragraph("A")
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("v.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)
        BVProjectFile.objects.create(
            project=self.projekt,
            anlage_nr=3,
            upload=upload,
            text_content="ignored",
            manual_analysis_json={"functions": {}},
            analysis_json={},
            verification_json={"functions": {}},
        )

    def test_worker_runs_vision_check(self):
        reply = json.dumps({"ok": True})
        with patch("core.llm_tasks.query_llm_with_images", return_value=reply):
            data = worker_run_anlage3_vision(self.projekt.pk)
        self.assertTrue(data["ok"]["value"])
        file_obj = self.projekt.anlagen.get(anlage_nr=3)
        self.assertTrue(file_obj.analysis_json["ok"]["value"])

    def test_model_name_is_forwarded(self):
        with patch(
            "core.llm_tasks.query_llm_with_images",
            return_value=json.dumps({"ok": False}),
        ) as mock_q:
            worker_run_anlage3_vision(self.projekt.pk, model_name="vision")
        self.assertEqual(mock_q.call_args[0][2], "vision")


class ProjektFileDeleteResultTests(NoesisTestCase):
    def setUp(self):
        self.user = User.objects.create_user("deluser", password="pass")
        self.client.login(username="deluser", password="pass")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        self.file = BVProjectFile.objects.create(
            project=self.projekt,
            anlage_nr=3,
            upload=SimpleUploadedFile("d.txt", b"data"),
            text_content="Text",
            analysis_json={
                "verhandlungsfaehig": {"value": True},
                "pages": {"value": 1},
            },
            manual_analysis_json={"functions": {}},
            verification_json={"functions": {}},
            manual_reviewed=True,
            verhandlungsfaehig=True,
        )

    def test_delete_resets_fields(self):
        url = reverse("projekt_file_delete_result", args=[self.file.pk])
        resp = self.client.post(url)
        self.assertRedirects(resp, reverse("anlage3_review", args=[self.projekt.pk]))
        self.file.refresh_from_db()
        self.assertIsNone(self.file.analysis_json)
        self.assertFalse(self.file.manual_reviewed)
        self.assertFalse(self.file.verhandlungsfaehig)


class ProjektFileVersionDeletionTests(NoesisTestCase):
    def setUp(self):
        self.user.is_staff = True
        self.user.save()
        self.client.login(username=self.user.username, password="pass")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")

    def test_delete_active_restores_parent(self):
        v1 = BVProjectFile.objects.create(
            project=self.projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"d"),
            text_content="t",
        )
        v1.is_active = False
        v1.save(update_fields=["is_active"])
        v2 = BVProjectFile.objects.create(
            project=self.projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("b.txt", b"d"),
            text_content="t",
            version=2,
            parent=v1,
        )
        url = reverse("delete_project_file_version", args=[v2.pk])
        resp = self.client.post(url, follow=True)
        self.assertRedirects(resp, reverse("projekt_detail", args=[self.projekt.pk]))
        self.assertContains(resp, "Die Version wurde erfolgreich gel\u00f6scht.")
        self.assertFalse(BVProjectFile.objects.filter(pk=v2.pk).exists())
        v1.refresh_from_db()
        self.assertTrue(v1.is_active)

    def test_delete_inactive_repairs_chain(self):
        v1 = BVProjectFile.objects.create(
            project=self.projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"d"),
            text_content="t",
            is_active=False,
        )
        v2 = BVProjectFile.objects.create(
            project=self.projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("b.txt", b"d"),
            text_content="t",
            version=2,
            parent=v1,
            is_active=False,
        )
        v3 = BVProjectFile.objects.create(
            project=self.projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("c.txt", b"d"),
            text_content="t",
            version=3,
            parent=v2,
        )
        url = reverse("delete_project_file_version", args=[v2.pk])
        resp = self.client.post(url, follow=True)
        self.assertContains(resp, "Die Version wurde erfolgreich gel\u00f6scht.")
        self.assertFalse(BVProjectFile.objects.filter(pk=v2.pk).exists())
        v3.refresh_from_db()
        self.assertEqual(v3.parent, v1)


class ProjektFileCheckResultTests(NoesisTestCase):
    def setUp(self):
        self.user = User.objects.create_user("vuser", password="pass")
        self.client.login(username="vuser", password="pass")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        self.file = BVProjectFile.objects.create(
            project=self.projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content="Text",
            manual_analysis_json={"functions": {}},
            analysis_json={},
            verification_json={"functions": {}},
        )
        self.file2 = BVProjectFile.objects.create(
            project=self.projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("b.txt", b"data"),
            text_content="Text2",
            manual_analysis_json={"functions": {}},
            analysis_json={},
            verification_json={"functions": {}},
        )

    def test_get_runs_check_and_redirects_to_edit(self):
        url = reverse("projekt_file_check_view", args=[self.file.pk])
        resp = self.client.get(url)
        self.assertRedirects(
            resp, reverse("projekt_file_edit_json", args=[self.file.pk])
        )
        self.file.refresh_from_db()
        expected = {"questions": {}}
        self.assertEqual(self.file.analysis_json, expected)

    def test_post_triggers_check_and_redirects(self):
        url = reverse("projekt_file_check_view", args=[self.file.pk])
        with patch("core.views.check_anlage1") as mock_func:
            mock_func.return_value = {"questions": {}}
            resp = self.client.post(url)
        self.assertRedirects(
            resp, reverse("projekt_file_edit_json", args=[self.file.pk])
        )
        mock_func.assert_called_with(self.file.pk, model_name=None)

    def test_anlage2_uses_parser_by_default(self):
        url = reverse("projekt_file_check_view", args=[self.file2.pk])
        with patch("core.views.run_anlage2_analysis") as mock_func:
            mock_func.return_value = []
            resp = self.client.get(url)
        self.assertRedirects(
            resp, reverse("projekt_file_edit_json", args=[self.file2.pk])
        )
        mock_func.assert_called_with(self.file2)

    def test_llm_param_triggers_full_check(self):
        url = reverse("projekt_file_check_view", args=[self.file2.pk]) + "?llm=1"
        with patch("core.views.check_anlage2") as mock_func:
            mock_func.return_value = {"task": "check_anlage2"}
            resp = self.client.get(url)
        self.assertRedirects(
            resp, reverse("projekt_file_edit_json", args=[self.file2.pk])
        )
        mock_func.assert_called_with(self.file2.pk, model_name=None)

    def test_anlage3_uses_analysis(self):
        pf = BVProjectFile.objects.create(
            project=self.projekt,
            anlage_nr=3,
            upload=SimpleUploadedFile("c.txt", b"x"),
            text_content="T",
        )
        url = reverse("projekt_file_check_view", args=[pf.pk])
        with patch("core.views.analyse_anlage3") as mock_func:
            mock_func.return_value = {"task": "analyse_anlage3"}
            resp = self.client.get(url)
        self.assertRedirects(resp, reverse("anlage3_review", args=[self.projekt.pk]))
        mock_func.assert_called_with(pf.pk, model_name=None)

    def test_anlage3_llm_param_triggers_vision_check(self):
        pf = BVProjectFile.objects.create(
            project=self.projekt,
            anlage_nr=3,
            upload=SimpleUploadedFile("d.txt", b"x"),
            text_content="T",
        )
        url = reverse("projekt_file_check_view", args=[pf.pk]) + "?llm=1"
        with patch("core.views.check_anlage3_vision") as mock_func:
            mock_func.return_value = {"task": "check_anlage3_vision"}
            resp = self.client.get(url)
        self.assertRedirects(resp, reverse("anlage3_review", args=[self.projekt.pk]))
        mock_func.assert_called_with(pf.pk, model_name=None)

    def test_parse_view_runs_parser(self):
        url = reverse("projekt_file_check_view", args=[self.file2.pk])
        with patch("core.views.run_anlage2_analysis") as mock_func:
            mock_func.return_value = []
            resp = self.client.get(url)
        self.assertRedirects(
            resp, reverse("projekt_file_edit_json", args=[self.file2.pk])
        )
        mock_func.assert_called_with(self.file2)

    def test_parse_view_rejects_other_files(self):
        url = reverse("projekt_file_parse_anlage2", args=[self.file.pk])
        resp = self.client.get(url)
        self.assertEqual(resp.status_code, 404)


class LLMConfigTests(NoesisTestCase):
    @override_settings(GOOGLE_API_KEY="x")
    @patch("google.generativeai.list_models")
    @patch("google.generativeai.configure")
    def test_ready_populates_models(self, mock_conf, mock_list):
        mock_list.return_value = [
            type("M", (), {"name": "m1"})(),
            type("M", (), {"name": "m2"})(),
        ]
        from core.signals import init_llm_config

        init_llm_config(None)
        cfg = LLMConfig.objects.first()
        self.assertIsNotNone(cfg)
        self.assertEqual(cfg.available_models, ["m1", "m2"])
        self.assertTrue(cfg.models_changed)

    @override_settings(GOOGLE_API_KEY="x")
    @patch("google.generativeai.list_models")
    @patch("google.generativeai.configure")
    def test_ready_updates_models(self, mock_conf, mock_list):
        LLMConfig.objects.create(available_models=["old"])
        mock_list.return_value = [type("M", (), {"name": "new"})()]
        from core.signals import init_llm_config

        init_llm_config(None)
        cfg = LLMConfig.objects.first()
        self.assertEqual(cfg.available_models, ["new"])
        self.assertTrue(cfg.models_changed)


class Anlage2ConfigSingletonTests(NoesisTestCase):
    def test_single_instance_enforced(self):
        first = Anlage2Config.get_instance()
        from django.db import transaction

        with self.assertRaises(IntegrityError):
            with transaction.atomic():
                Anlage2Config.objects.create()
        self.assertEqual(Anlage2Config.objects.count(), 1)


class TileVisibilityTests(NoesisTestCase):
    def setUp(self):
        admin_group = Group.objects.create(name="admin")
        self.user = User.objects.create_user("tileuser", password="pass")
        self.user.groups.add(admin_group)
        self.group = Group.objects.create(name="role")
        work = Area.objects.get_or_create(slug="work", defaults={"name": "Work"})[0]
        self.personal = Area.objects.get_or_create(
            slug="personal", defaults={"name": "Personal"}
        )[0]
        self.talkdiary = Tile.objects.get_or_create(
            slug="talkdiary",
            defaults={
                "name": "TalkDiary",
                "url_name": "talkdiary_personal",
            },
        )[0]
        self.talkdiary.areas.add(self.personal)
        self.group.areas.add(self.personal)
        self.group.tiles.add(self.talkdiary)
        self.projekt = Tile.objects.get_or_create(
            slug="projektverwaltung",
            defaults={
                "name": "Projektverwaltung",
                "url_name": "projekt_list",
            },
        )[0]
        self.projekt.areas.add(work)
        self.group.areas.add(work)
        self.cfg = LLMConfig.objects.first() or LLMConfig.objects.create(
            models_changed=False
        )
        self.client.login(username="tileuser", password="pass")

    def test_personal_without_access(self):
        resp = self.client.get(reverse("personal"))
        self.assertNotContains(resp, "TalkDiary")

    def test_personal_with_access(self):
        UserTileAccess.objects.create(user=self.user, tile=self.talkdiary)
        resp = self.client.get(reverse("personal"))
        self.assertContains(resp, "TalkDiary")

    def test_personal_with_group_access(self):
        self.user.groups.add(self.group)
        resp = self.client.get(reverse("personal"))
        self.assertContains(resp, "TalkDiary")

    def test_personal_with_image(self):
        UserTileAccess.objects.create(user=self.user, tile=self.talkdiary)
        self.talkdiary.image.save(
            "img.png",
            SimpleUploadedFile("img.png", b"data"),
            save=True,
        )
        resp = self.client.get(reverse("personal"))
        self.assertContains(resp, "<img", html=False)

    def _login(self, name: str) -> User:
        """Erzeugt einen Benutzer und loggt ihn ein."""
        user = User.objects.create_user(name, password="pass")
        self.client.login(username=name, password="pass")
        return user

    def test_talkdiary_access_denied_without_tile(self):
        self._login("nodiac")
        resp = self.client.get(reverse("personal"))
        self.assertNotContains(resp, "TalkDiary")
        resp = self.client.get(reverse("talkdiary_personal"))
        self.assertEqual(resp.status_code, 403)

    def test_talkdiary_access_allowed_with_tile(self):
        user = self._login("withdia")
        UserTileAccess.objects.create(user=user, tile=self.talkdiary)
        resp = self.client.get(reverse("personal"))
        self.assertContains(resp, "TalkDiary")
        resp = self.client.get(reverse("talkdiary_personal"))
        self.assertEqual(resp.status_code, 200)

    def test_projekt_tile_hidden_without_group(self):
        self._login("noproj")
        resp = self.client.get(reverse("work"))
        self.assertNotContains(resp, "Projektverwaltung")
        resp = self.client.get(reverse("projekt_list"))
        self.assertEqual(resp.status_code, 200)

    def test_projekt_access_allowed_with_tile(self):
        user = self._login("withproj")
        UserTileAccess.objects.create(user=user, tile=self.projekt)
        resp = self.client.get(reverse("work"))
        self.assertContains(resp, "Projektverwaltung")
        resp = self.client.get(reverse("projekt_list"))
        self.assertEqual(resp.status_code, 200)

    def test_projekt_tile_visible_with_group(self):
        user = self._login("groupproj")
        user.groups.add(self.group)
        GroupTileAccess.objects.create(group=self.group, tile=self.projekt)
        resp = self.client.get(reverse("work"))
        self.assertContains(resp, "Projektverwaltung")
        resp = self.client.get(reverse("projekt_list"))
        self.assertEqual(resp.status_code, 200)


class LLMConfigNoticeMiddlewareTests(NoesisTestCase):
    def setUp(self):
        admin_group = Group.objects.create(name="admin")
        self.user = User.objects.create_user("llmadmin", password="pass")
        self.user.groups.add(admin_group)
        self.client.login(username="llmadmin", password="pass")
        LLMConfig.objects.create(models_changed=True)

    def test_message_shown(self):
        resp = self.client.get(reverse("home"))
        msgs = [m.message for m in resp.context["messages"]]
        self.assertTrue(any("LLM-Einstellungen" in m for m in msgs))


class HomeRedirectTests(NoesisTestCase):
    def setUp(self):
        self.user = User.objects.create_user("redir", password="pass")
        personal = Area.objects.get_or_create(
            slug="personal", defaults={"name": "Personal"}
        )[0]
        tile = Tile.objects.get_or_create(
            slug="talkdiary",
            defaults={
                "name": "TalkDiary",
                "url_name": "talkdiary_personal",
            },
        )[0]
        tile.areas.add(personal)
        UserTileAccess.objects.create(user=self.user, tile=tile)
        self.client.login(username="redir", password="pass")

    def test_redirect_personal(self):
        resp = self.client.get(reverse("home"))
        self.assertRedirects(resp, reverse("personal"))


class AreaImageTests(NoesisTestCase):
    def setUp(self):
        self.user = User.objects.create_user("areauser", password="pass")
        self.client.login(username="areauser", password="pass")

    def test_home_without_images(self):
        Area.objects.get_or_create(slug="work", defaults={"name": "Work"})
        Area.objects.get_or_create(slug="personal", defaults={"name": "Personal"})
        resp = self.client.get(reverse("home"))
        self.assertNotContains(resp, 'alt="Work"', html=False)
        self.assertNotContains(resp, 'alt="Personal"', html=False)

    def test_home_with_images(self):
        work, _ = Area.objects.get_or_create(slug="work", defaults={"name": "Work"})
        personal, _ = Area.objects.get_or_create(
            slug="personal", defaults={"name": "Personal"}
        )
        work.image.save("w.png", SimpleUploadedFile("w.png", b"d"), save=True)
        personal.image.save("p.png", SimpleUploadedFile("p.png", b"d"), save=True)
        resp = self.client.get(reverse("home"))
        self.assertContains(resp, reverse("work"))
        self.assertContains(resp, reverse("personal"))
        self.assertContains(resp, "Arbeitsassistent")
        self.assertContains(resp, "Persönlicher Bereich")


class RecordingDeleteTests(NoesisTestCase):
    def setUp(self):
        self.user = User.objects.create_user("recuser", password="pass")
        self.client.login(username="recuser", password="pass")
        self.personal = Area.objects.get_or_create(
            slug="personal", defaults={"name": "Personal"}
        )[0]
        self.tile = Tile.objects.get_or_create(
            slug="talkdiary",
            defaults={
                "name": "TalkDiary",
                "url_name": "talkdiary_personal",
            },
        )[0]
        self.tile.areas.add(self.personal)
        UserTileAccess.objects.create(user=self.user, tile=self.tile)
        audio = SimpleUploadedFile("a.wav", b"data")
        transcript = SimpleUploadedFile("a.md", b"text")
        self.rec = Recording.objects.create(
            user=self.user,
            bereich=self.personal,
            audio_file=audio,
            transcript_file=transcript,
        )
        self.audio_path = Path(self.rec.audio_file.path)
        self.trans_path = Path(self.rec.transcript_file.path)

    def test_delete_own_recording(self):
        url = reverse("recording_delete", args=[self.rec.pk])
        resp = self.client.post(url)
        self.assertRedirects(resp, reverse("talkdiary_personal"))
        self.assertFalse(Recording.objects.filter(pk=self.rec.pk).exists())
        self.assertFalse(self.audio_path.exists())
        self.assertFalse(self.trans_path.exists())

    def test_delete_requires_post(self):
        url = reverse("recording_delete", args=[self.rec.pk])
        resp = self.client.get(url)
        self.assertEqual(resp.status_code, 405)
        self.assertTrue(Recording.objects.filter(pk=self.rec.pk).exists())

    def test_delete_other_user_recording(self):
        other = User.objects.create_user("other", password="pass")
        rec = Recording.objects.create(
            user=other,
            bereich=self.personal,
            audio_file=SimpleUploadedFile("b.wav", b"d"),
            transcript_file=SimpleUploadedFile("b.md", b"t"),
        )
        url = reverse("recording_delete", args=[rec.pk])
        resp = self.client.post(url)
        self.assertEqual(resp.status_code, 404)
        self.assertTrue(Recording.objects.filter(pk=rec.pk).exists())


class ModelSelectionTests(NoesisTestCase):
    def setUp(self):
        self.user = User.objects.create_user("modeluser", password="pass")
        self.client.login(username="modeluser", password="pass")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        BVProjectFile.objects.create(
            project=self.projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content="Text",
            manual_analysis_json={"functions": {}},
            analysis_json={},
            verification_json={"functions": {}},
        )
        self.pf2 = BVProjectFile.objects.create(
            project=self.projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("b.txt", b"data"),
        )
        LLMConfig.objects.create(
            default_model="d",
            gutachten_model="g",
            anlagen_model="a",
        )

    def test_file_check_uses_category(self):
        url = reverse("projekt_file_check", args=[self.projekt.pk, 1])
        with patch("core.views.check_anlage1") as mock_func:
            mock_func.return_value = {"questions": {}}
            resp = self.client.post(url, {"model_category": "anlagen"})
        self.assertEqual(resp.status_code, 200)
        mock_func.assert_called_with(
            self.projekt.pk,
            model_name=LLMConfig.get_instance().anlagen_model,
        )

    def test_forms_show_categories(self):
        edit_url = reverse("projekt_edit", args=[self.projekt.pk])
        resp = self.client.get(edit_url)
        self.assertContains(resp, "Standard")
        self.assertContains(resp, "Gutachten")
        self.assertContains(resp, "Anlagen")

        view_url = reverse(
            "projekt_file_check_view", args=[self.projekt.anlagen.first().pk]
        )
        with patch("core.views.check_anlage1") as mock_func:
            mock_func.return_value = {"questions": {}}
            resp = self.client.get(view_url)
        self.assertRedirects(
            resp,
            reverse("projekt_file_edit_json", args=[self.projekt.anlagen.first().pk]),
        )

    def test_functions_check_uses_model(self):
        url = reverse("projekt_functions_check", args=[self.projekt.pk])
        with (
            patch("core.views.async_task", return_value="tid") as mock_task,
            patch("core.views.transaction.on_commit", side_effect=lambda f: f()),
        ):
            resp = self.client.post(url, {"model": "mf"})
        self.assertEqual(resp.status_code, 200)
        mock_task.assert_called_with(
            "core.llm_tasks.run_conditional_anlage2_check",
            self.pf2.pk,
            "mf",
        )
        pf = BVProjectFile.objects.get(pk=self.pf2.pk)
        self.assertEqual(pf.verification_task_id, "tid")
        self.assertEqual(pf.processing_status, BVProjectFile.PROCESSING)


class FunctionImportExportTests(NoesisTestCase):
    def setUp(self):
        Anlage2Function.objects.all().delete()
        admin_group = Group.objects.create(name="admin")
        self.user = User.objects.create_user("adminie", password="pass")
        self.user.groups.add(admin_group)
        self.client.login(username="adminie", password="pass")

    def test_export_returns_json(self):
        func = Anlage2Function.objects.get(name="Anmelden")
        Anlage2SubQuestion.objects.create(funktion=func, frage_text="Warum?")
        url = reverse("anlage2_function_export")
        resp = self.client.get(url)
        self.assertEqual(resp.status_code, 200)
        data = json.loads(resp.content)
        self.assertEqual(data[0]["name"], "Anmelden")
        self.assertEqual(data[0]["subquestions"][0]["frage_text"], "Warum?")

    def test_import_creates_functions(self):
        payload = json.dumps([{"name": "Anmelden", "subquestions": ["Frage"]}])
        file = SimpleUploadedFile("func.json", payload.encode("utf-8"))
        url = reverse("anlage2_function_import")
        resp = self.client.post(
            url,
            {"json_file": file, "clear_first": "on"},
            format="multipart",
        )
        self.assertRedirects(resp, reverse("anlage2_function_list"))
        self.assertTrue(Anlage2Function.objects.filter(name="Anmelden").exists())

    def test_import_accepts_german_keys(self):
        payload = json.dumps(
            [
                {
                    "funktion": "Anwesenheit",
                    "unterfragen": [{"frage": "Testfrage"}],
                }
            ]
        )
        file = SimpleUploadedFile("func_de.json", payload.encode("utf-8"))
        url = reverse("anlage2_function_import")
        resp = self.client.post(
            url,
            {"json_file": file, "clear_first": "on"},
            format="multipart",
        )
        self.assertRedirects(resp, reverse("anlage2_function_list"))
        self.assertTrue(Anlage2Function.objects.filter(name="Anwesenheit").exists())
        self.assertEqual(
            Anlage2SubQuestion.objects.filter(funktion__name="Anwesenheit").count(),
            1,
        )

    def test_roundtrip_preserves_aliases(self):
        func = Anlage2Function.objects.get(name="Anmelden")
        func.detection_phrases = {"name_aliases": ["Sign in"]}
        func.save()
        Anlage2SubQuestion.objects.filter(funktion=func).delete()
        Anlage2SubQuestion.objects.create(
            funktion=func,
            frage_text="Warum?",
            detection_phrases={"name_aliases": ["Why"]},
        )
        export_url = reverse("anlage2_function_export")
        resp = self.client.get(export_url)
        payload = resp.content
        file = SimpleUploadedFile("func.json", payload)
        import_url = reverse("anlage2_function_import")
        resp = self.client.post(
            import_url,
            {"json_file": file, "clear_first": "on"},
            format="multipart",
        )
        self.assertRedirects(resp, reverse("anlage2_function_list"))
        func = Anlage2Function.objects.get(name="Anmelden")
        self.assertEqual(func.detection_phrases.get("name_aliases"), ["Sign in"])
        sub = func.anlage2subquestion_set.get(frage_text="Warum?")
        self.assertEqual(sub.detection_phrases.get("name_aliases"), ["Why"])


class GutachtenLLMCheckTests(NoesisTestCase):
    def setUp(self):
        self.user = User.objects.create_user("gcheck", password="pass")
        self.client.login(username="gcheck", password="pass")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        self.knowledge = SoftwareKnowledge.objects.create(
            project=self.projekt,
            software_name="A",
            is_known_by_llm=True,
        )
        self.gutachten = Gutachten.objects.create(
            software_knowledge=self.knowledge, text="Test"
        )

    def test_endpoint_updates_note(self):
        url = reverse("gutachten_llm_check", args=[self.gutachten.pk])
        with patch("core.views.check_gutachten_functions") as mock_func:
            mock_func.return_value = "Hinweis"
            resp = self.client.post(url)
        self.assertRedirects(resp, reverse("projekt_initial_pruefung", args=[self.projekt.pk]))
        self.projekt.refresh_from_db()
        self.assertEqual(self.projekt.gutachten_function_note, "Hinweis")


class FeatureVerificationTests(NoesisTestCase):
    def setUp(self):
        self.projekt = BVProject.objects.create(
            software_typen="Word, Excel",
            beschreibung="x",
        )
        self.pf = BVProjectFile.objects.create(
            project=self.projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"data"),
            manual_analysis_json={"functions": {}},
            analysis_json={},
            verification_json={"functions": {}},
        )
        self.func = Anlage2Function.objects.get(name="Export")
        Anlage2SubQuestion.objects.filter(funktion=self.func).delete()
        self.sub = Anlage2SubQuestion.objects.create(
            funktion=self.func,
            frage_text="Warum?",
        )

    def test_any_yes_returns_true(self):
        with patch(
            "core.llm_tasks.query_llm",
            side_effect=["Ja", "Nein", "Begruendung", "Nein"],
        ) as mock_q:
            result = worker_verify_feature(self.pf.pk, "function", self.func.pk)
        self.assertEqual(
            result,
            {
                "technisch_verfuegbar": True,
                "ki_begruendung": "Begruendung",
                "ki_beteiligt": False,
                "ki_beteiligt_begruendung": "",
            },
        )
        self.assertEqual(mock_q.call_count, 4)
        pf = BVProjectFile.objects.get(project=self.projekt, anlage_nr=2)
        res = AnlagenFunktionsMetadaten.objects.get(
            anlage_datei=pf,
            funktion=self.func,
        )
        fe = FunktionsErgebnis.objects.filter(
            anlage_datei=pf, funktion=self.func, quelle="ki"
        ).first()
        self.assertIsNotNone(fe)
        self.assertTrue(fe.technisch_verfuegbar)
        self.assertFalse(fe.ki_beteiligung)
        self.assertEqual(fe.begruendung, "Begruendung")

    def test_all_no_returns_false(self):
        with patch(
            "core.llm_tasks.query_llm",
            side_effect=["Nein", "Nein"],
        ):
            result = worker_verify_feature(self.pf.pk, "subquestion", self.sub.pk)
        self.assertEqual(
            result,
            {
                "technisch_verfuegbar": False,
                "ki_begruendung": "",
                "ki_beteiligt": None,
                "ki_beteiligt_begruendung": "",
            },
        )

        pf = BVProjectFile.objects.get(project=self.projekt, anlage_nr=2)
        res = AnlagenFunktionsMetadaten.objects.get(
            anlage_datei=pf,
            funktion=self.func,
        )
        fe = FunktionsErgebnis.objects.filter(
            anlage_datei=pf, funktion=self.func, subquestion=self.sub, quelle="ki"
        ).first()
        self.assertIsNotNone(fe)
        self.assertFalse(fe.technisch_verfuegbar)

    def test_subquestion_context_contains_question(self):
        """Die Subquestion wird korrekt im Kontext übergeben."""
        with patch(
            "core.llm_tasks.query_llm",
            side_effect=["Nein", "Nein"],
        ) as mock_q:
            worker_verify_feature(self.pf.pk, "subquestion", self.sub.pk)
        first_call_ctx = mock_q.call_args_list[0].args[1]
        self.assertEqual(first_call_ctx["subquestion_text"], self.sub.frage_text)

    def test_gutachten_text_is_added_to_context(self):
        doc = Document()
        doc.add_paragraph("Info")
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        dest_dir = Path(settings.MEDIA_ROOT) / "gutachten"
        dest_dir.mkdir(parents=True, exist_ok=True)
        dest = dest_dir / Path(tmp.name).name
        shutil.copy(tmp.name, dest)
        Path(tmp.name).unlink(missing_ok=True)
        self.projekt.gutachten_file.name = f"gutachten/{dest.name}"
        self.projekt.save(update_fields=["gutachten_file"])
        with patch(
            "core.llm_tasks.query_llm",
            side_effect=["Ja", "Nein", "B", "Nein"],
        ) as mock_q:
            worker_verify_feature(self.pf.pk, "function", self.func.pk)
        ctx = mock_q.call_args_list[0].args[1]
        self.assertIn("Info", ctx["gutachten"])
        dest.unlink(missing_ok=True)

    def test_mixed_returns_none(self):
        with patch(
            "core.llm_tasks.query_llm",
            side_effect=["Unsicher", "Nein", "Begruendung"],
        ):
            result = worker_verify_feature(self.pf.pk, "function", self.func.pk)
        self.assertIsNone(result["technisch_verfuegbar"])
        self.assertEqual(result["ki_begruendung"], "Begruendung")
        self.assertIsNone(result["ki_beteiligt"])
        self.assertEqual(result["ki_beteiligt_begruendung"], "")
        pf = BVProjectFile.objects.get(project=self.projekt, anlage_nr=2)
        fe = FunktionsErgebnis.objects.filter(
            anlage_datei=pf, funktion=self.func, quelle="ki"
        ).first()
        self.assertIsNotNone(fe)
        self.assertEqual(fe.begruendung, "Begruendung")

    def test_negotiable_set_on_match(self):
        pf = BVProjectFile.objects.get(project=self.projekt, anlage_nr=2)
        AnlagenFunktionsMetadaten.objects.create(
            anlage_datei=pf,
            funktion=self.func,
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=pf,
            funktion=self.func,
            quelle="parser",
            technisch_verfuegbar=True,
        )
        with patch(
            "core.llm_tasks.query_llm",
            side_effect=["Ja", "Nein", "", "Nein"],
        ):
            worker_verify_feature(self.pf.pk, "function", self.func.pk)
        parser_fe = FunktionsErgebnis.objects.filter(
            anlage_datei=pf,
            funktion=self.func,
            quelle="parser",
        ).first()
        ai_fe = FunktionsErgebnis.objects.filter(
            anlage_datei=pf,
            funktion=self.func,
            quelle="ki",
        ).first()
        self.assertTrue(parser_fe.technisch_verfuegbar)
        self.assertTrue(ai_fe.technisch_verfuegbar)

    def test_negotiable_not_set_on_mismatch(self):
        pf = BVProjectFile.objects.get(project=self.projekt, anlage_nr=2)
        AnlagenFunktionsMetadaten.objects.create(
            anlage_datei=pf,
            funktion=self.func,
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=pf,
            funktion=self.func,
            quelle="parser",
            technisch_verfuegbar=False,
        )
        with patch(
            "core.llm_tasks.query_llm",
            side_effect=["Ja", "Nein", "", "Nein"],
        ):
            worker_verify_feature(self.pf.pk, "function", self.func.pk)
        parser_fe = FunktionsErgebnis.objects.filter(
            anlage_datei=pf,
            funktion=self.func,
            quelle="parser",
        ).first()
        ai_fe = FunktionsErgebnis.objects.filter(
            anlage_datei=pf,
            funktion=self.func,
            quelle="ki",
        ).first()
        self.assertFalse(parser_fe.technisch_verfuegbar)
        self.assertTrue(ai_fe.technisch_verfuegbar)

    def test_warnung_bei_geloeschter_datei(self):
        """L\u00f6schen der Datei f\u00fchrt zu Warnung statt Ausnahme."""

        pf = BVProjectFile.objects.get(project=self.projekt, anlage_nr=2)

        original_update = AnlagenFunktionsMetadaten.objects.update_or_create

        def _wrapped(*args, **kwargs):
            obj, created = original_update(*args, **kwargs)
            pf.delete()
            return obj, created

        class _DummyQS:
            def exists(self) -> bool:  # noqa: D401 - einfache Hilfsklasse
                return False


        with patch("core.llm_tasks.query_llm", side_effect=["Nein", "Nein"]):
            with patch(
                "core.llm_tasks.BVProjectFile.objects.filter",
                return_value=_DummyQS(),
            ):
                with self.assertLogs("core.llm_tasks", level="WARNING") as cm:
                    result = worker_verify_feature(
                        self.pf.pk, "function", self.func.pk
                    )

        self.assertEqual(result, {})
        self.assertTrue(
            any("Ergebnis wird verworfen" in msg for msg in cm.output)
        )

    def test_integrity_error_logs_and_returns_empty(self):
        """Allgemeiner IntegrityError führt zu Fehler-Log und leerem Ergebnis."""
        def _raise(*args, **kwargs):
            raise IntegrityError("boom")

        with patch("core.llm_tasks.query_llm", side_effect=["Nein", "Nein"]):
            with patch(
                "core.llm_tasks.AnlagenFunktionsMetadaten.objects.update_or_create",
                side_effect=_raise,
            ):
                with self.assertLogs("core.llm_tasks", level="ERROR") as cm:
                    result = worker_verify_feature(
                        self.pf.pk, "function", self.func.pk
                    )

        self.assertEqual(result, {})
        self.assertTrue(any("Integrit" in msg for msg in cm.output))


class InitialCheckTests(NoesisTestCase):
    def setUp(self):
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")

    def test_known_software_stores_description(self):
        with patch(
            "core.llm_tasks.query_llm",
            side_effect=["Ja", "Beschreibung"],
        ) as mock_q:
            sk = SoftwareKnowledge.objects.create(
                project=self.projekt,
                software_name="A",
            )
            result = worker_run_initial_check(sk.pk)
        self.assertTrue(result["is_known_by_llm"])
        self.assertEqual(result["description"], "Beschreibung")
        self.assertEqual(mock_q.call_count, 2)
        sk.refresh_from_db()
        self.assertTrue(sk.is_known_by_llm)
        self.assertEqual(sk.description, "Beschreibung")

    def test_unknown_sets_flags(self):
        with patch("core.llm_tasks.query_llm", return_value="Nein") as mock_q:
            sk = SoftwareKnowledge.objects.create(
                project=self.projekt,
                software_name="A",
            )
            result = worker_run_initial_check(sk.pk)
        self.assertFalse(result["is_known_by_llm"])
        self.assertEqual(result["description"], "")
        self.assertEqual(mock_q.call_count, 1)
        sk.refresh_from_db()
        self.assertFalse(sk.is_known_by_llm)
        self.assertEqual(sk.description, "")

    def test_context_is_passed_to_prompt(self):
        with patch("core.llm_tasks.query_llm", return_value="Nein") as mock_q:
            sk = SoftwareKnowledge.objects.create(
                project=self.projekt,
                software_name="A",
            )
            worker_run_initial_check(sk.pk, user_context="Hint")
        prompt_text = mock_q.call_args[0][0].text
        self.assertIn("Hint", prompt_text)


class EditKIJustificationTests(NoesisTestCase):
    def setUp(self):
        self.user = User.objects.create_user("justi", password="pass")
        self.client.login(username="justi", password="pass")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        self.file = BVProjectFile.objects.create(
            project=self.projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"data"),
            manual_analysis_json={"functions": {}},
            analysis_json={},
            verification_json={
                "Export": {"technisch_verfuegbar": True, "ki_begruendung": "Alt"}
            },
        )
        self.func = Anlage2Function.objects.get(name="Export")

    def test_get_returns_form(self):
        url = (
            reverse(
                "edit_ki_justification",
                args=[self.file.pk],
            )
            + f"?function={self.func.pk}"
        )
        resp = self.client.get(url)
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, "Alt")

    def test_post_updates_value(self):
        url = reverse("edit_ki_justification", args=[self.file.pk])
        resp = self.client.post(
            url,
            {"function": self.func.pk, "ki_begruendung": "Neu"},
        )
        self.assertRedirects(
            resp, reverse("projekt_file_edit_json", args=[self.file.pk])
        )
        self.file.refresh_from_db()
        self.assertEqual(
            self.file.verification_json["Export"]["ki_begruendung"],
            "Neu",
        )


class JustificationDetailEditTests(NoesisTestCase):
    def setUp(self):
        self.user = User.objects.create_user("jdet", password="pw")
        self.client.login(username="jdet", password="pw")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        self.file = BVProjectFile.objects.create(
            project=self.projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("jd.txt", b"data"),
            manual_analysis_json={"functions": {}},
            analysis_json={},
            verification_json={"functions": {}},
        )
        self.func = Anlage2Function.objects.get(name="Export")
        FunktionsErgebnis.objects.create(
            anlage_datei=self.file,
            funktion=self.func,
            quelle="ki",
            begruendung="Begruendung",
        )

    def test_get_loads_text(self):
        url = reverse("justification_detail_edit", args=[self.file.pk, self.func.name])
        resp = self.client.get(url)
        self.assertContains(resp, "Begruendung")


class KIInvolvementDetailEditTests(NoesisTestCase):
    def setUp(self):
        self.user = User.objects.create_user("kid", password="pw")
        self.client.login(username="kid", password="pw")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        self.file = BVProjectFile.objects.create(
            project=self.projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("kid.txt", b"data"),
            manual_analysis_json={"functions": {}},
            analysis_json={},
            verification_json={"functions": {}},
        )
        self.func = Anlage2Function.objects.get(name="Export")
        FunktionsErgebnis.objects.create(
            anlage_datei=self.file,
            funktion=self.func,
            quelle="ki",
            ki_beteiligt_begruendung="Initial",
        )

    def test_get_loads_text(self):
        url = reverse("ki_involvement_detail_edit", args=[self.file.pk, self.func.name])
        resp = self.client.get(url)
        self.assertContains(resp, "Initial")

    def test_post_updates_value(self):
        url = reverse("ki_involvement_detail_edit", args=[self.file.pk, self.func.name])
        resp = self.client.post(url, {"justification": "Neu"})
        self.assertRedirects(resp, reverse("projekt_file_edit_json", args=[self.file.pk]))
        fe = FunktionsErgebnis.objects.get(
            anlage_datei=self.file, funktion=self.func, quelle="ki"
        )
        self.assertEqual(fe.ki_beteiligt_begruendung, "Neu")


class VerificationToInitialTests(NoesisTestCase):
    def setUp(self):
        self.project = BVProject.objects.create(software_typen="A", beschreibung="x")
        BVProjectFile.objects.create(
            project=self.project,
            anlage_nr=2,
            upload=SimpleUploadedFile("v.txt", b"data"),
            manual_analysis_json={"functions": {}},
            analysis_json={},
            verification_json={"functions": {}},
        )
        self.func = Anlage2Function.objects.get(name="Export")
        Anlage2SubQuestion.objects.filter(funktion=self.func).delete()
        self.sub = Anlage2SubQuestion.objects.create(
            funktion=self.func,
            frage_text="Warum?",
        )

    def test_conversion_reads_ai_fields(self):
        data = {
            "Export": {
                "technisch_verfuegbar": True,
                "ki_beteiligt": True,
                "ki_beteiligt_begruendung": "Grund",
            },
            "Export: Warum?": {
                "technisch_verfuegbar": False,
                "ki_beteiligt": False,
                "ki_beteiligt_begruendung": "Nein",
            },
        }

        pf = self.project.anlagen.get(anlage_nr=2)
        AnlagenFunktionsMetadaten.objects.create(anlage_datei=pf, funktion=self.func)
        AnlagenFunktionsMetadaten.objects.create(
            anlage_datei=pf, funktion=self.func, subquestion=self.sub
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=pf,
            funktion=self.func,
            quelle="ki",
            technisch_verfuegbar=True,
            ki_beteiligung=True,
            begruendung="Grund",
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=pf,
            funktion=self.func,
            subquestion=self.sub,
            quelle="ki",
            technisch_verfuegbar=False,
            ki_beteiligung=False,
            begruendung="Nein",
        )
        result = _verification_to_initial(pf)
        func_data = result["functions"][str(self.func.id)]
        self.assertTrue(func_data["technisch_vorhanden"])
        self.assertTrue(func_data["ki_beteiligt"])
        self.assertEqual(func_data["begruendung"], "Grund")

        sub_data = func_data["subquestions"][str(self.sub.id)]
        self.assertFalse(sub_data["technisch_vorhanden"])
        self.assertFalse(sub_data["ki_beteiligt"])
        self.assertEqual(sub_data["begruendung"], "Nein")


class UserImportExportTests(NoesisTestCase):
    def setUp(self):
        admin_group = Group.objects.create(name="admin")
        self.user = User.objects.create_user("uadmin", password="pass")
        self.user.groups.add(admin_group)
        self.client.login(username="uadmin", password="pass")

        self.group = Group.objects.create(name="testgroup")
        self.area = Area.objects.get_or_create(slug="work", defaults={"name": "Work"})[
            0
        ]
        self.tile = Tile.objects.create(slug="t1", name="T", url_name="tile")
        self.tile.areas.add(self.area)
        self.group.areas.add(self.area)
        self.group.tiles.add(self.tile)

    def test_export_json(self):
        self.user.groups.add(self.group)


        url = reverse("admin_export_users_permissions")
        resp = self.client.get(url)
        self.assertEqual(resp.status_code, 200)
        data = json.loads(resp.content)
        entry = next(u for u in data if u["username"] == "uadmin")
        self.assertIn("testgroup", entry["groups"])
        self.assertIn("tile", entry["tiles"])

    def test_import_creates_user(self):
        payload = json.dumps(
            [
                {
                    "username": "neu",
                    "email": "a@b.c",
                    "groups": ["testgroup"],

                }
            ]
        )
        file = SimpleUploadedFile("u.json", payload.encode("utf-8"))
        url = reverse("admin_import_users_permissions")
        resp = self.client.post(url, {"json_file": file}, format="multipart")
        self.assertRedirects(resp, reverse("admin_user_list"))
        user = User.objects.get(username="neu")
        self.assertTrue(user.groups.filter(name="testgroup").exists())



class Anlage1ImportTests(NoesisTestCase):
    def setUp(self):
        admin_group = Group.objects.create(name="admin")
        self.user = User.objects.create_user("a1import", password="pass")
        self.user.groups.add(admin_group)
        self.client.login(username="a1import", password="pass")

    def test_clear_first_resets_questions(self):
        Anlage1Question.objects.create(num=99, text="Alt?", enabled=True)
        payload = json.dumps([{"text": "Neu?"}])
        file = SimpleUploadedFile("a1.json", payload.encode("utf-8"))
        url = reverse("admin_anlage1_import")
        resp = self.client.post(
            url,
            {"json_file": file, "clear_first": "on"},
            format="multipart",
        )
        self.assertRedirects(resp, reverse("admin_anlage1"))
        self.assertEqual(Anlage1Question.objects.count(), 1)
        q = Anlage1Question.objects.first()
        self.assertEqual(q.text, "Neu?")
        self.assertEqual(q.num, 1)


class Anlage2ConfigImportExportTests(NoesisTestCase):
    def setUp(self):
        admin_group = Group.objects.create(name="admin")
        self.user = User.objects.create_user("cfgadmin", password="pass")
        self.user.groups.add(admin_group)
        self.client.login(username="cfgadmin", password="pass")
        # Anlage2Config ist ein Singleton. Für die Tests müssen wir die
        # globale Instanz verwenden und die benötigten Felder direkt darauf
        # setzen.
        self.cfg = Anlage2Config.get_instance()
        self.cfg.parser_mode = "auto"
        self.cfg.parser_order = ["exact"]
        self.cfg.text_technisch_verfuegbar_true = ["ja"]
        self.cfg.save()

    def test_export_contains_headings_and_phrases(self):
        Anlage2ColumnHeading.objects.create(
            config=self.cfg,
            field_name="technisch_vorhanden",
            text="Verfügbar?",
        )
        AntwortErkennungsRegel.objects.create(
            regel_name="R1",
            erkennungs_phrase="ja",
            actions_json=[{"field": "technisch_verfuegbar", "value": True}],
            prioritaet=0,
        )
        a4 = Anlage4ParserConfig.objects.first()
        a4.delimiter_phrase = "X"
        a4.save()
        url = reverse("admin_anlage2_config_export")
        resp = self.client.get(url)
        self.assertEqual(resp.status_code, 200)
        data = json.loads(resp.content)
        self.assertEqual(data["config"]["parser_mode"], self.cfg.parser_mode)
        self.assertEqual(data["config"]["parser_order"], self.cfg.parser_order)
        self.assertIn(
            {"field_name": "technisch_vorhanden", "text": "Verfügbar?"},
            data["alias_headings"],
        )
        self.assertIn("answer_rules", data)
        self.assertTrue(any(r["regel_name"] == "R1" for r in data["answer_rules"]))
        self.assertEqual(data["a4_parser"]["delimiter_phrase"], "X")

    def test_import_creates_headings(self):
        payload = json.dumps(
            {
                "config": {
                "parser_mode": "text_only",
                "parser_order": ["exact"],
                    "enforce_subquestion_override": True,
                    "text_technisch_verfuegbar_true": ["ja"],
                },
                "alias_headings": [{"field_name": "ki_beteiligung", "text": "KI?"}],
                "answer_rules": [
                    {
                        "regel_name": "R2",
                        "erkennungs_phrase": "nein",
                        "actions": [{"field": "technisch_verfuegbar", "value": False}],
                        "prioritaet": 1,
                    }
                ],
                "a4_parser": {"delimiter_phrase": "Y"},
            }
        )
        file = SimpleUploadedFile("cfg.json", payload.encode("utf-8"))
        url = reverse("admin_anlage2_config_import")
        resp = self.client.post(url, {"json_file": file}, format="multipart")
        self.assertRedirects(resp, reverse("anlage2_config"))
        self.assertTrue(
            Anlage2ColumnHeading.objects.filter(
                field_name="ki_beteiligung", text="KI?"
            ).exists()
        )
        self.assertTrue(AntwortErkennungsRegel.objects.filter(regel_name="R2").exists())
        a4_cfg = Anlage4ParserConfig.objects.first()
        self.assertEqual(a4_cfg.delimiter_phrase, "Y")
        self.cfg.refresh_from_db()
        self.assertEqual(self.cfg.parser_mode, "text_only")
        self.assertEqual(self.cfg.parser_order, ["exact"])
        self.assertTrue(self.cfg.enforce_subquestion_override)
        self.assertEqual(self.cfg.text_technisch_verfuegbar_true, ["ja"])


class Anlage2ConfigViewTests(NoesisTestCase):
    def setUp(self):
        admin = Group.objects.create(name="admin")
        self.user = User.objects.create_user("cfguser", password="pass")
        self.user.groups.add(admin)
        self.client.login(username="cfguser", password="pass")
        self.cfg = Anlage2Config.get_instance()

    def _build_general_data(self, **extra) -> dict:
        """Erstellt Grunddaten für das Anlage2Config-Formular."""
        data = {name: "" for name in Anlage2ConfigForm.OPTIONAL_JSON_FIELDS}
        data.update(extra)
        return data


    def test_update_parser_mode(self):
        url = reverse("anlage2_config")
        resp = self.client.post(
            url,
            self._build_general_data(
                parser_mode="text_only",
                parser_order=["exact"],
                action="save_general",
                active_tab="general",
            ),
        )
        self.assertRedirects(resp, url + "?tab=general")
        self.cfg.refresh_from_db()
        self.assertEqual(self.cfg.parser_order, ["exact"])

    def test_update_parser_order(self):
        url = reverse("anlage2_config")
        resp = self.client.post(
            url,
            self._build_general_data(
                parser_mode=self.cfg.parser_mode,
                parser_order=["exact"],
                action="save_general",
                active_tab="general",
            ),
        )
        self.assertRedirects(resp, url + "?tab=general")
        self.cfg.refresh_from_db()
        self.assertEqual(self.cfg.parser_order, ["exact"])

    def test_save_table_tab(self):
        url = reverse("anlage2_config")
        resp = self.client.post(
            url,
            {
                "new_field": "technisch_vorhanden",
                "new_text": "Verfügbar?",
                "action": "save_table",
                "active_tab": "table",
            },
        )
        self.assertRedirects(resp, url + "?tab=table")
        self.assertTrue(Anlage2ColumnHeading.objects.filter(text="Verfügbar?").exists())

    def test_multiline_phrases_saved(self):
        url = reverse("anlage2_config")
        resp = self.client.post(
            url,
            self._build_general_data(
                text_technisch_verfuegbar_true="ja\nokay\n",
                parser_mode=self.cfg.parser_mode,
                parser_order=self.cfg.parser_order,
                action="save_general",
                active_tab="general",
            ),
        )
        self.assertRedirects(resp, url + "?tab=general")
        self.cfg.refresh_from_db()
        self.assertEqual(self.cfg.text_technisch_verfuegbar_true, ["ja", "okay"])


class ParserRuleImportExportTests(NoesisTestCase):
    def setUp(self):
        admin_group = Group.objects.create(name="admin")
        self.user = User.objects.create_user(
            "ruleadmin", password="pass", is_staff=True
        )
        self.user.groups.add(admin_group)
        self.client.login(username="ruleadmin", password="pass")

    def test_export_returns_json(self):
        AntwortErkennungsRegel.objects.create(
            regel_name="R1",
            erkennungs_phrase="ja",
            actions_json=[{"field": "tech", "value": True}],
        )
        url = reverse("anlage2_parser_rule_export")
        resp = self.client.get(url)
        self.assertEqual(resp.status_code, 200)
        data = json.loads(resp.content)
        self.assertTrue(any(r["regel_name"] == "R1" for r in data))

    def test_import_creates_rule(self):
        payload = json.dumps(
            [
                {
                    "regel_name": "R2",
                    "erkennungs_phrase": "nein",
                    "actions": [{"field": "tech", "value": False}],
                }
            ]
        )
        file = SimpleUploadedFile("rules.json", payload.encode("utf-8"))
        url = reverse("anlage2_parser_rule_import")
        resp = self.client.post(url, {"json_file": file}, format="multipart")
        self.assertRedirects(resp, reverse("parser_rule_list"))
        self.assertTrue(AntwortErkennungsRegel.objects.filter(regel_name="R2").exists())


class AjaxAnlage2ReviewTests(NoesisTestCase):
    def setUp(self):
        self.user = User.objects.create_user("reviewer", password="pw", is_staff=True)
        self.client.login(username="reviewer", password="pw")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        self.pf = BVProjectFile.objects.create(
            project=self.projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"x"),
            manual_analysis_json={"functions": {}},
            analysis_json={},
            verification_json={"functions": {}},
        )
        self.func = Anlage2Function.objects.get(name="Anmelden")
        Anlage2SubQuestion.objects.filter(funktion=self.func).delete()

    def test_manual_result_saved(self):
        url = reverse("ajax_save_anlage2_review")
        resp = self.client.post(
            url,
            data=json.dumps(
                {
                    "project_file_id": self.pf.pk,
                    "function_id": self.func.pk,
                    "status": True,
                    "field_name": "technisch_vorhanden",
                }
            ),
            content_type="application/json",
        )
        self.assertEqual(resp.status_code, 200)

        self.pf.refresh_from_db()
        func_data = self.pf.manual_analysis_json["functions"][str(self.func.pk)]
        self.assertTrue(func_data["technisch_vorhanden"])

        fe = FunktionsErgebnis.objects.filter(
            anlage_datei__project=self.projekt,
            funktion=self.func,
            quelle="manuell",
        ).first()
        self.assertTrue(fe.technisch_verfuegbar)

    def test_gap_generated_on_difference(self):
        AnlagenFunktionsMetadaten.objects.create(
            anlage_datei=self.pf,
            funktion=self.func,
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=self.pf,
            funktion=self.func,
            quelle="ki",
            technisch_verfuegbar=True,
        )
        url = reverse("ajax_save_anlage2_review")
        resp = self.client.post(
            url,
            data=json.dumps(
                {
                    "project_file_id": self.pf.pk,
                    "function_id": self.func.pk,
                    "status": False,
                }
            ),
            content_type="application/json",
        )
        self.assertEqual(resp.status_code, 200)
        data = resp.json()
        self.assertEqual(data.get("gap_summary"), "")
        res = AnlagenFunktionsMetadaten.objects.get(
            anlage_datei=self.pf,
            funktion=self.func,
        )
        self.assertEqual(res.gap_summary, "")

        gap_url = reverse("ajax_generate_gap_summary", args=[res.pk])

        def immediate(name, *args):
            self.assertEqual(name, "core.llm_tasks.worker_generate_gap_summary")
            worker_generate_gap_summary(*args)

        with patch("core.views.async_task", side_effect=immediate):
            resp = self.client.post(gap_url)
        self.assertEqual(resp.status_code, 200)
        res.refresh_from_db()
        self.assertEqual(res.gap_notiz, "")
        self.assertEqual(res.gap_summary, "")
        gap_entry = FunktionsErgebnis.objects.filter(
            anlage_datei=self.pf,
            funktion=self.func,
            quelle="gap",
        ).latest("created_at")
        self.assertIsNotNone(gap_entry)

    def test_manual_sets_negotiable(self):
        AnlagenFunktionsMetadaten.objects.create(
            anlage_datei=self.pf,
            funktion=self.func,
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=self.pf,
            funktion=self.func,
            quelle="ki",
            technisch_verfuegbar=True,
        )

        url = reverse("ajax_save_anlage2_review")
        resp = self.client.post(
            url,
            data=json.dumps(
                {
                    "project_file_id": self.pf.pk,
                    "function_id": self.func.pk,
                    "status": True,
                    "field_name": "technisch_vorhanden",
                }
            ),
            content_type="application/json",
        )
        self.assertEqual(resp.status_code, 200)
        fe = FunktionsErgebnis.objects.filter(
            anlage_datei__project=self.projekt,
            funktion=self.func,
            quelle="manuell",
        ).first()
        self.assertTrue(fe.technisch_verfuegbar)

    def test_save_einsatz_telefonica(self):
        url = reverse("ajax_save_anlage2_review")
        resp = self.client.post(
            url,
            data=json.dumps(
                {
                    "project_file_id": self.pf.pk,
                    "function_id": self.func.pk,
                    "status": True,
                    "field_name": "einsatz_bei_telefonica",
                }
            ),
            content_type="application/json",
        )
        self.assertEqual(resp.status_code, 200)
        fe = FunktionsErgebnis.objects.filter(
            anlage_datei__project=self.projekt,
            funktion=self.func,
            quelle="manuell",
        ).first()
        self.assertTrue(fe.einsatz_bei_telefonica)
        self.pf.refresh_from_db()
        func_data = self.pf.manual_analysis_json["functions"][str(self.func.pk)]
        self.assertTrue(func_data["einsatz_bei_telefonica"])

    def test_save_lv_kontrolle(self):
        url = reverse("ajax_save_anlage2_review")
        resp = self.client.post(
            url,
            data=json.dumps(
                {
                    "project_file_id": self.pf.pk,
                    "function_id": self.func.pk,
                    "status": False,
                    "field_name": "zur_lv_kontrolle",
                }
            ),
            content_type="application/json",
        )
        self.assertEqual(resp.status_code, 200)
        fe = FunktionsErgebnis.objects.filter(
            anlage_datei__project=self.projekt,
            funktion=self.func,
            quelle="manuell",
        ).first()
        self.assertFalse(fe.zur_lv_kontrolle)
        self.pf.refresh_from_db()
        func_data = self.pf.manual_analysis_json["functions"][str(self.func.pk)]
        self.assertFalse(func_data["zur_lv_kontrolle"])

    def test_manual_result_merge(self):
        url = reverse("ajax_save_anlage2_review")
        self.client.post(
            url,
            data=json.dumps(
                {
                    "project_file_id": self.pf.pk,
                    "function_id": self.func.pk,
                    "status": True,
                    "field_name": "technisch_vorhanden",
                }
            ),
            content_type="application/json",
        )
        self.client.post(
            url,
            data=json.dumps(
                {
                    "project_file_id": self.pf.pk,
                    "function_id": self.func.pk,
                    "status": False,
                    "field_name": "ki_beteiligung",
                }
            ),
            content_type="application/json",
        )

        self.pf.refresh_from_db()
        func_data = self.pf.manual_analysis_json["functions"][str(self.func.pk)]
        self.assertTrue(func_data["technisch_vorhanden"])
        self.assertFalse(func_data["ki_beteiligung"])
        fes = FunktionsErgebnis.objects.filter(
            anlage_datei__project=self.projekt,
            funktion=self.func,
            quelle="manuell",
        )
        self.assertEqual(fes.count(), 2)

    def test_set_negotiable_override(self):
        AnlagenFunktionsMetadaten.objects.create(
            anlage_datei=self.pf,
            funktion=self.func,
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=self.pf,
            funktion=self.func,
            quelle="ki",
            technisch_verfuegbar=True,
        )
        url = reverse("ajax_save_anlage2_review")
        resp = self.client.post(
            url,
            data=json.dumps(
                {
                    "project_file_id": self.pf.pk,
                    "function_id": self.func.pk,
                    "set_negotiable": True,
                }
            ),
            content_type="application/json",
        )
        self.assertEqual(resp.status_code, 200)
        res = AnlagenFunktionsMetadaten.objects.get(
            anlage_datei=self.pf,
            funktion=self.func,
        )
        self.assertTrue(res.is_negotiable_manual_override)

        self.client.post(
            url,
            data=json.dumps(
                {
                    "project_file_id": self.pf.pk,
                    "function_id": self.func.pk,
                    "set_negotiable": None,
                }
            ),
            content_type="application/json",
        )
        res.refresh_from_db()
        self.assertIsNone(res.is_negotiable_manual_override)

    def test_negotiable_does_not_set_manual_value(self):
        AnlagenFunktionsMetadaten.objects.create(
            anlage_datei=self.pf,
            funktion=self.func,
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=self.pf,
            funktion=self.func,
            quelle="ki",
            technisch_verfuegbar=True,
        )
        url = reverse("ajax_save_anlage2_review")
        self.client.post(
            url,
            data=json.dumps(
                {
                    "project_file_id": self.pf.pk,
                    "function_id": self.func.pk,
                    "set_negotiable": True,
                }
            ),
            content_type="application/json",
        )
        res = AnlagenFunktionsMetadaten.objects.get(
            anlage_datei=self.pf,
            funktion=self.func,
        )
        self.assertFalse(
            FunktionsErgebnis.objects.filter(
                anlage_datei__project=self.projekt,
                funktion=self.func,
                quelle="manuell",
            ).exists()
        )

    def test_manual_yes_triggers_subquestion_checks(self):
        sub = Anlage2SubQuestion.objects.create(
            funktion=self.func,
            frage_text="S?",
        )
        url = reverse("ajax_save_anlage2_review")
        with patch("core.views.async_task") as mock_task:
            resp = self.client.post(
                url,
                data=json.dumps(
                    {
                        "project_file_id": self.pf.pk,
                        "function_id": self.func.pk,
                        "status": True,
                        "field_name": "technisch_vorhanden",
                    }
                ),
                content_type="application/json",
            )
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(mock_task.call_count, 2)
        mock_task.assert_any_call(
            "core.llm_tasks.worker_verify_feature",
            self.pf.pk,
            "function",
            self.func.pk,
        )
        mock_task.assert_any_call(
            "core.llm_tasks.worker_verify_feature",
            self.pf.pk,
            "subquestion",
            sub.pk,
        )


class SupervisionGapTests(NoesisTestCase):
    def test_manually_negotiable_function_excluded_from_supervision(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("f.txt", b"x"),
        )
        func = Anlage2Function.objects.get(name="Anmelden")
        Anlage2SubQuestion.objects.filter(funktion=func).delete()
        AnlagenFunktionsMetadaten.objects.create(
            anlage_datei=pf,
            funktion=func,
            is_negotiable_manual_override=True,
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=pf,
            funktion=func,
            quelle="parser",
            technisch_verfuegbar=True,
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=pf,
            funktion=func,
            quelle="ki",
            technisch_verfuegbar=False,
        )
        sub = Anlage2SubQuestion.objects.create(funktion=func, frage_text="S?")
        AnlagenFunktionsMetadaten.objects.create(
            anlage_datei=pf,
            funktion=func,
            subquestion=sub,
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=pf,
            funktion=func,
            subquestion=sub,
            quelle="parser",
            technisch_verfuegbar=True,
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=pf,
            funktion=func,
            subquestion=sub,
            quelle="ki",
            technisch_verfuegbar=False,
        )
        groups = _build_supervision_groups(pf)
        self.assertEqual(groups, [])

    def test_subquestion_before_function_excluded_from_supervision(self):
        """Unterfrage vor Hauptfunktion: Funktion wird dennoch übersprungen."""

        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("f.txt", b"x"),
        )
        func = Anlage2Function.objects.get(name="Anmelden")
        Anlage2SubQuestion.objects.filter(funktion=func).delete()
        sub = Anlage2SubQuestion.objects.create(funktion=func, frage_text="S?")

        # Unterfrage zuerst anlegen, damit sie im Default-Ordering vor der Funktion steht
        AnlagenFunktionsMetadaten.objects.create(
            anlage_datei=pf,
            funktion=func,
            subquestion=sub,
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=pf,
            funktion=func,
            subquestion=sub,
            quelle="parser",
            technisch_verfuegbar=True,
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=pf,
            funktion=func,
            subquestion=sub,
            quelle="ki",
            technisch_verfuegbar=False,
        )

        # Hauptfunktion nachträglich als verhandlungsfähig markieren
        AnlagenFunktionsMetadaten.objects.create(
            anlage_datei=pf,
            funktion=func,
            is_negotiable_manual_override=True,
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=pf,
            funktion=func,
            quelle="parser",
            technisch_verfuegbar=True,
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=pf,
            funktion=func,
            quelle="ki",
            technisch_verfuegbar=False,
        )

        groups = _build_supervision_groups(pf)
        self.assertEqual(groups, [])

    def test_ai_reason_uses_function_begruendung(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("f.txt", b"x"),
        )
        func = Anlage2Function.objects.get(name="Anmelden")
        res = AnlagenFunktionsMetadaten.objects.create(anlage_datei=pf, funktion=func)
        FunktionsErgebnis.objects.create(
            anlage_datei=pf,
            funktion=func,
            quelle="parser",
            technisch_verfuegbar=True,
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=pf,
            funktion=func,
            quelle="ki",
            technisch_verfuegbar=False,
            begruendung="Func reason",
            ki_beteiligt_begruendung="KI involvement",
        )
        row = _build_supervision_row(res, pf)
        self.assertEqual(row["ai_reason"], "Func reason")


class Anlage2ResetTests(NoesisTestCase):
    """Tests zum Zurücksetzen von Anlage-2-Ergebnissen."""

    def setUp(self):
        self.user = User.objects.create_user("reset", password="pw")
        self.client.login(username="reset", password="pw")

    def test_run_anlage2_analysis_resets_results(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf_old = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("old.txt", b"x"),
        )
        Anlage2Function.objects.all().delete()
        func = Anlage2Function.objects.get(name="Anmelden")
        AnlagenFunktionsMetadaten.objects.create(
            anlage_datei=pf_old,
            funktion=func,
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=pf_old,
            funktion=func,
            quelle="parser",
        )
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("new.txt", b"x"),
        )
        with (
            patch(
                "core.llm_tasks.parse_anlage2_table",
                return_value=[{"funktion": "Anmelden"}],
            ),
            patch("core.text_parser.parse_anlage2_text", return_value=[]),
        ):
            run_anlage2_analysis(pf)
        results = AnlagenFunktionsMetadaten.objects.filter(
            anlage_datei__project=projekt
        )
        self.assertEqual(results.count(), Anlage2Function.objects.count())
        fe = FunktionsErgebnis.objects.filter(
            anlage_datei__project=projekt,
            funktion=func,
            quelle="parser",
        ).first()
        self.assertIsNotNone(fe)

    def test_conditional_check_resets_results(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf_old = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("old.txt", b"x"),
        )
        func = Anlage2Function.objects.get(name="Anmelden")
        AnlagenFunktionsMetadaten.objects.create(
            anlage_datei=pf_old,
            funktion=func,
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=pf_old,
            funktion=func,
            quelle="ki",
            technisch_verfuegbar=False,
        )

        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("new.txt", b"x"),
        )

        def fake(_pid, _typ, fid, _model=None):
            pf_latest = BVProjectFile.objects.filter(
                project=projekt, anlage_nr=2
            ).first()
            AnlagenFunktionsMetadaten.objects.update_or_create(
                anlage_datei=pf_latest,
                funktion_id=fid,
                defaults={},
            )
            FunktionsErgebnis.objects.create(
                anlage_datei=pf_latest,
                funktion_id=fid,
                quelle="ki",
                technisch_verfuegbar=True,
            )
            return {}

        with (
            patch("core.llm_tasks.worker_verify_feature", side_effect=fake) as mock_verify,
            patch("core.llm_tasks.async_task") as mock_async,
            patch("core.llm_tasks.result") as mock_result,
        ):
            mock_async.side_effect = lambda name, *a, **k: (
                mock_verify(*a, **k) or "tid"
            )
            mock_result.side_effect = lambda *a, **k: None
            run_conditional_anlage2_check(pf.pk)
        results = AnlagenFunktionsMetadaten.objects.filter(
            anlage_datei__project=projekt
        )
        self.assertEqual(results.count(), Anlage2Function.objects.count())
        fe = FunktionsErgebnis.objects.filter(
            anlage_datei__project=projekt,
            funktion=func,
            quelle="ki",
        ).first()
        self.assertTrue(fe.technisch_verfuegbar)

    def test_conditional_check_deletes_only_current_file_metadata(self):
        """Nur Metadaten der geprüften Anlage werden entfernt."""
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"x"),
        )
        other_pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=3,
            upload=SimpleUploadedFile("b.txt", b"x"),
        )
        func = Anlage2Function.objects.get(name="Anmelden")
        AnlagenFunktionsMetadaten.objects.create(anlage_datei=pf, funktion=func)
        AnlagenFunktionsMetadaten.objects.create(
            anlage_datei=other_pf, funktion=func
        )
        with (
            patch("core.llm_tasks.async_task", return_value="tid"),
            patch("core.llm_tasks.result", return_value=None),
        ):
            run_conditional_anlage2_check(pf.pk)

        self.assertFalse(
            AnlagenFunktionsMetadaten.objects.filter(anlage_datei=pf).exists()
        )
        self.assertTrue(
            AnlagenFunktionsMetadaten.objects.filter(
                anlage_datei=other_pf
            ).exists()
        )

    def test_ajax_reset_all_reviews_resets_manual_fields(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"x"),
        )
        func = Anlage2Function.objects.get(name="Anmelden")
        res = AnlagenFunktionsMetadaten.objects.create(
            anlage_datei=pf,
            funktion=func,
            is_negotiable_manual_override=True,
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=pf,
            funktion=func,
            quelle="parser",
            technisch_verfuegbar=True,
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=pf,
            funktion=func,
            quelle="ki",
            technisch_verfuegbar=True,
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=pf,
            funktion=func,
            quelle="manuell",
            technisch_verfuegbar=True,
        )
        self.client.login(username=self.user.username, password="pw")
        url = reverse("ajax_reset_all_reviews", args=[pf.pk])
        resp = self.client.post(url)
        self.assertEqual(resp.status_code, 200)
        res.refresh_from_db()
        self.assertFalse(
            FunktionsErgebnis.objects.filter(
                anlage_datei__project=projekt,
                anlage_datei=pf,
                funktion=func,
                quelle="manuell",
            ).exists()
        )
        self.assertIsNone(res.is_negotiable_manual_override)
        self.assertTrue(
            FunktionsErgebnis.objects.filter(
                anlage_datei__project=projekt,
                anlage_datei=pf,
                funktion=func,
                quelle="parser",
                technisch_verfuegbar=True,
            ).exists()
        )
        self.assertTrue(
            FunktionsErgebnis.objects.filter(
                anlage_datei__project=projekt,
                anlage_datei=pf,
                funktion=func,
                quelle="ki",
                technisch_verfuegbar=True,
            ).exists()
        )

    def test_hx_update_review_cell_toggles_manual_entry(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"x"),
        )
        func = Anlage2Function.objects.get(name="Anmelden")
        result = AnlagenFunktionsMetadaten.objects.create(
            anlage_datei=pf,
            funktion=func,
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=pf,
            funktion=func,
            quelle="parser",
            technisch_verfuegbar=True,
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=pf,
            funktion=func,
            quelle="ki",
            technisch_verfuegbar=False,
        )

        self.client.login(username=self.superuser.username, password="pass")
        url = reverse("hx_update_review_cell", args=[result.pk, "technisch_vorhanden"])

        resp = self.client.post(url, HTTP_HX_REQUEST="true")
        self.assertEqual(resp.status_code, 200)
        self.assertTrue(
            FunktionsErgebnis.objects.filter(
                anlage_datei=pf,
                funktion=func,
                quelle="manuell",
            ).exists()
        )
        pf.refresh_from_db()
        self.assertTrue(pf.manual_analysis_json["functions"][str(func.id)]["technisch_vorhanden"])

        resp = self.client.post(url, HTTP_HX_REQUEST="true")
        self.assertEqual(resp.status_code, 200)
        self.assertFalse(
            FunktionsErgebnis.objects.filter(
                anlage_datei=pf,
                funktion=func,
                quelle="manuell",
            ).exists()
        )
        pf.refresh_from_db()
        self.assertIsNone(pf.manual_analysis_json)


class GapReportTests(NoesisTestCase):
    def setUp(self):
        self.client.login(username=self.superuser.username, password="pass")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        self.pf1 = BVProjectFile.objects.create(
            project=self.projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"data"),
            question_review={"1": {"hinweis": "Hinweis", "vorschlag": "V"}},
        )
        self.pf2 = BVProjectFile.objects.create(
            project=self.projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("b.txt", b"data"),
        )
        self.func = Anlage2Function.objects.get(name="Anmelden")
        AnlagenFunktionsMetadaten.objects.create(
            anlage_datei=self.pf2,
            funktion=self.func,
            gap_summary="Extern",
            gap_notiz="Intern",
        )

    def test_tasks_return_text(self):
        with patch("core.llm_tasks.query_llm", return_value="T1") as mock_q:
            text = summarize_anlage1_gaps(self.projekt)
        self.assertEqual(text, "T1")
        self.assertTrue(mock_q.called)

        with patch("core.llm_tasks.query_llm", return_value="T2") as mock_q:
            text = summarize_anlage2_gaps(self.projekt)
        self.assertEqual(text, "T2")
        self.assertTrue(mock_q.called)

    def test_view_saves_text(self):
        url = reverse("gap_report_view", args=[self.projekt.pk])
        with patch("core.views.summarize_anlage1_gaps", return_value="A1"), patch(
            "core.views.summarize_anlage2_gaps", return_value="A2"
        ):
            resp = self.client.get(url)
            self.assertContains(resp, "A1")
            self.assertContains(resp, "A2")
            resp = self.client.post(url, {"text1": "E1", "text2": "E2"})
        self.assertRedirects(resp, reverse("projekt_detail", args=[self.projekt.pk]))
        self.pf1.refresh_from_db()
        self.pf2.refresh_from_db()
        self.assertEqual(self.pf1.gap_summary, "E1")
        self.assertEqual(self.pf2.gap_summary, "E2")

    def test_delete_gap_report(self):
        self.pf1.gap_summary = "E1"
        self.pf2.gap_summary = "E2"
        self.pf1.save(update_fields=["gap_summary"])
        self.pf2.save(update_fields=["gap_summary"])
        url = reverse("delete_gap_report", args=[self.projekt.pk])
        resp = self.client.post(url)
        self.assertRedirects(resp, reverse("projekt_detail", args=[self.projekt.pk]))
        self.pf1.refresh_from_db()
        self.pf2.refresh_from_db()
        self.assertEqual(self.pf1.gap_summary, "")
        self.assertEqual(self.pf2.gap_summary, "")


class ProjektDetailGapTests(NoesisTestCase):
    def setUp(self):
        self.client.login(username=self.superuser.username, password="pass")

    def test_anlage1_gap_sets_flag(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"data"),
            question_review={"1": {"vorschlag": "V"}},
        )
        resp = self.client.get(reverse("projekt_detail", args=[projekt.pk]))
        self.assertTrue(resp.context["can_gap_report"])

    def test_anlage2_gap_sets_flag(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf2 = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("b.txt", b"data"),
        )
        func = Anlage2Function.objects.get(name="Anmelden")
        AnlagenFunktionsMetadaten.objects.create(
            anlage_datei=pf2,
            funktion=func,
            supervisor_notes="Hinweis",
        )
        resp = self.client.get(reverse("projekt_detail", args=[projekt.pk]))
        self.assertTrue(resp.context["can_gap_report"])

    def test_anlage4_gap_sets_flag(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=4,
            upload=SimpleUploadedFile("c.txt", b"data"),
            manual_comment="Hinweis",
        )
        resp = self.client.get(reverse("projekt_detail", args=[projekt.pk]))
        self.assertTrue(resp.context["can_gap_report"])

    def test_anlage5_gap_sets_flag(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf5 = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=5,
            upload=SimpleUploadedFile("d.txt", b"data"),
        )
        ZweckKategorieA.objects.all().delete()
        ZweckKategorieA.objects.create(beschreibung="A")
        ZweckKategorieA.objects.create(beschreibung="B")
        Anlage5Review.objects.create(project_file=pf5)
        resp = self.client.get(reverse("projekt_detail", args=[projekt.pk]))
        self.assertTrue(resp.context["can_gap_report"])

class ManualGapDetectionTests(TestCase):
    """Tests für die Funktion ``_has_manual_gap``."""

    def test_detects_difference(self) -> None:
        """Erkennt eine Abweichung zwischen Dokument und manuellem Wert."""
        doc_data = {"technisch_vorhanden": True}
        manual_data = {"technisch_vorhanden": False}
        self.assertTrue(_has_manual_gap(doc_data, manual_data))

    def test_no_gap_when_equal(self) -> None:
        """Kein GAP, wenn Werte übereinstimmen."""
        doc_data = {"technisch_vorhanden": True}
        manual_data = {"technisch_vorhanden": True}
        self.assertFalse(_has_manual_gap(doc_data, manual_data))

    def test_gap_when_doc_missing(self) -> None:
        """Erkennt eine Lücke, wenn Dokumentdaten fehlen."""
        doc_data: dict = {}
        manual_data = {"technisch_vorhanden": True}
        self.assertTrue(_has_manual_gap(doc_data, manual_data))

    def test_no_gap_when_manual_missing(self) -> None:
        """Kein GAP, wenn manuelle Daten fehlen."""
        doc_data = {"technisch_vorhanden": True}
        manual_data = {"technisch_vorhanden": None}
        self.assertFalse(_has_manual_gap(doc_data, manual_data))

    def test_gap_with_additional_manual_field(self) -> None:
        """Erkennt eine Lücke bei zusätzlichen manuellen Feldern."""
        doc_data = {"technisch_vorhanden": True}
        manual_data = {
            "technisch_vorhanden": True,
            "neues_feld": False,
        }
        self.assertTrue(_has_manual_gap(doc_data, manual_data))

    def test_no_gap_when_parser_missing_special(self) -> None:
        """Kein GAP bei Spezialfeldern ohne Parser-Wert."""
        doc_data: dict = {}
        manual_data = {"einsatz_bei_telefonica": True}
        self.assertFalse(_has_manual_gap(doc_data, manual_data))

    def test_gap_only_on_difference_special(self) -> None:
        """GAP bei Spezialfeldern nur bei Abweichung."""
        doc_data = {"einsatz_bei_telefonica": False}
        manual_data = {"einsatz_bei_telefonica": True}
        self.assertTrue(_has_manual_gap(doc_data, manual_data))
        manual_data2 = {"einsatz_bei_telefonica": False}
        self.assertFalse(_has_manual_gap(doc_data, manual_data2))


class ResolveValueLogicTests(TestCase):
    """Tests für die Feldpriorisierung in ``_resolve_value``."""

    def test_doc_overrides_ai_for_special_fields(self) -> None:
        """Bei Spezialfeldern hat Parser-Vorrang vor KI."""
        val, src = _resolve_value(
            None, False, True, "einsatz_bei_telefonica", False, True
        )
        self.assertTrue(val)
        self.assertEqual(src, "Dokumenten-Analyse")

    def test_ai_still_used_for_regular_fields(self) -> None:
        """Bei normalen Feldern dominiert der KI-Wert."""
        val, _ = _resolve_value(
            None, False, True, "technisch_vorhanden", False, True
        )
        self.assertFalse(val)
