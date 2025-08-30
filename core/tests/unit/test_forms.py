from pathlib import Path
from tempfile import NamedTemporaryFile

import pytest
from django.http import QueryDict
from django.urls import reverse
from django.contrib.auth.models import User
from django.core.files.uploadedfile import SimpleUploadedFile
from django.conf import settings
from docx import Document

from ..base import NoesisTestCase
from ...forms import (
    BVProjectForm,
    BVProjectUploadForm,
    BVProjectFileJSONForm,
    BVProjectFileForm,
    Anlage2ConfigForm,
    Anlage2ReviewForm,
    Anlage5ReviewForm,
)
from ...models import (
    BVProject,
    BVProjectFile,
    ZweckKategorieA,
    SoftwareKnowledge,
    Gutachten,
    Anlage2Config,
)
from ...reporting import generate_gap_analysis

pytestmark = [pytest.mark.unit, pytest.mark.usefixtures("seed_db"), pytest.mark.django_db]

class BVProjectFormTests(NoesisTestCase):
    def test_project_form_docx_validation(self):
        data = QueryDict(mutable=True)
        data.update({"title": "Testprojekt"})
        data.setlist("software", ["A"])
        form = BVProjectForm(data)
        self.assertTrue(form.is_valid())

    def test_upload_form_docx_validation(self):
        valid = BVProjectUploadForm(
            {}, {"docx_file": SimpleUploadedFile("t.docx", b"d")}
        )
        self.assertTrue(valid.is_valid())
        invalid = BVProjectUploadForm(
            {}, {"docx_file": SimpleUploadedFile("t.txt", b"d")}
        )
        self.assertFalse(invalid.is_valid())


class Anlage2ConfigFormTests(NoesisTestCase):
    def test_parser_order_field(self):
        cfg = Anlage2Config.get_instance()
        form = Anlage2ConfigForm({"parser_order": ["table"]}, instance=cfg)
        self.assertTrue(form.is_valid())
        inst = form.save()
        self.assertEqual(inst.parser_order, ["table"])


@pytest.fixture
def projekt_file_setup(client, user_factory, bv_project_factory, bv_project_file_factory):
    """Erzeugt Projekt und zugehörige Dateien."""

    user = user_factory(username="user3")
    client.login(username=user.username, password="pw")
    projekt = bv_project_factory(software_typen="A", beschreibung="x")
    file = bv_project_file_factory(
        project=projekt,
        anlage_nr=4,
        analysis_json={"items": ["Alt"], "manual_review": {"0": {"ok": False, "nego": False, "note": ""}}},
    )
    anlage1 = bv_project_file_factory(
        project=projekt,
        anlage_nr=1,
        analysis_json={
            "questions": {
                "1": {
                    "answer": "foo",
                    "status": None,
                    "hinweis": "",
                    "vorschlag": "",
                }
            }
        },
    )
    return {"user": user, "projekt": projekt, "file": file, "anlage1": anlage1}


def test_edit_json_updates_and_reports(client, projekt_file_setup):
    file = projekt_file_setup["file"]
    projekt = projekt_file_setup["projekt"]

    url = reverse("projekt_file_edit_json", args=[file.pk])
    resp = client.post(
        url,
        {
            "analysis_json": '{"items": ["Neu"], "manual_review": {"0": {"note": "Hinweis"}}}',
        },
    )
    assert resp.status_code == 302
    file.refresh_from_db()
    assert file.analysis_json["items"] == ["Neu"]
    assert file.analysis_json["manual_review"]["0"]["note"] == "Hinweis"
    path = generate_gap_analysis(projekt)
    try:
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert '"note": "Hinweis"' in text
        assert '"Alt"' not in text
    finally:
        path.unlink(missing_ok=True)


def test_invalid_json_shows_error(client, projekt_file_setup):
    file = projekt_file_setup["file"]

    url = reverse("projekt_file_edit_json", args=[file.pk])
    file.analysis_json = {"items": []}
    file.save(update_fields=["analysis_json"])
    resp = client.post(url, {"analysis_json": "{"})
    assert resp.status_code == 200
    file.refresh_from_db()
    assert file.analysis_json == {"items": []}


def test_question_review_saved(client, projekt_file_setup):
    projekt = projekt_file_setup["projekt"]
    anlage1 = projekt_file_setup["anlage1"]

    url = reverse("hx_toggle_anlage1_ok", args=[anlage1.pk, 1])
    resp = client.post(url)
    assert resp.status_code == 302
    assert resp.headers["Location"] == reverse("projekt_detail", args=[projekt.pk])
    anlage1.refresh_from_db()
    assert anlage1.question_review["1"]["ok"]
    assert "note" not in anlage1.question_review["1"]


def test_question_review_saved_htmx(client, projekt_file_setup):
    anlage1 = projekt_file_setup["anlage1"]

    url = reverse("hx_toggle_anlage1_ok", args=[anlage1.pk, 1])
    resp = client.post(url, HTTP_HX_REQUEST="true")
    assert resp.status_code == 200
    anlage1.refresh_from_db()
    assert anlage1.question_review["1"]["ok"]


def test_question_review_extended_fields_saved(client, projekt_file_setup):
    anlage1 = projekt_file_setup["anlage1"]

    url = reverse("hx_anlage1_note", args=[anlage1.pk, 1, "hinweis"])
    client.post(url, {"text": "Fehlt"})
    url = reverse("hx_anlage1_note", args=[anlage1.pk, 1, "vorschlag"])
    resp = client.post(url, {"text": "Mehr Infos"})
    assert resp.status_code == 200
    anlage1.refresh_from_db()
    data = anlage1.question_review["1"]
    assert data["hinweis"] == "Fehlt"
    assert data["vorschlag"] == "Mehr Infos"
    assert "status" not in data


def test_question_review_prefill_from_analysis(client, projekt_file_setup):
    anlage1 = projekt_file_setup["anlage1"]

    anlage1.question_review = None
    anlage1.analysis_json = {
        "questions": {
            "1": {
                "answer": "A",
                "status": "ok",
                "hinweis": "H",
                "vorschlag": "V",
            }
        }
    }
    anlage1.save()

    url = reverse("projekt_file_edit_json", args=[anlage1.pk])
    resp = client.get(url)
    qa = resp.context["qa"]
    assert qa[0]["hinweis"] == "H"
    assert qa[0]["vorschlag"] == "V"


def test_edit_page_has_mde(client, projekt_file_setup):
    projekt = projekt_file_setup["projekt"]
    pf = BVProjectFile.objects.create(
        project=projekt,
        anlage_nr=5,
        upload=SimpleUploadedFile("c.txt", b"data"),
        text_content="Text",
    )
    url = reverse("projekt_file_edit_json", args=[pf.pk])
    resp = client.get(url)
    assert "markdown_editor.js" in resp.content.decode()


class GutachtenEditDeleteTests(NoesisTestCase):
    def setUp(self):
        self.user = User.objects.create_user("euser", password="pass")
        self.client.login(username="euser", password="pass")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        doc = Document()
        doc.add_paragraph("Alt")
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            self.projekt.gutachten_file.save(
                "g.docx", SimpleUploadedFile("g.docx", fh.read())
            )
        Path(tmp.name).unlink(missing_ok=True)
        self.knowledge = SoftwareKnowledge.objects.create(
            project=self.projekt,
            software_name="A",
            is_known_by_llm=True,
        )
        self.gutachten = Gutachten.objects.create(software_knowledge=self.knowledge, text="Alt")

    def test_view_shows_content(self):
        url = reverse("hx_project_software_tab", args=[self.projekt.pk, "gutachten"])
        resp = self.client.get(url, HTTP_HX_REQUEST="true")
        self.assertContains(resp, "Alt")

    def test_edit_updates_text(self):
        url = reverse("gutachten_edit", args=[self.gutachten.pk])
        resp = self.client.post(url, {"text": "Neu"})
        self.assertRedirects(resp, reverse("projekt_initial_pruefung", args=[self.projekt.pk]))
        self.gutachten.refresh_from_db()
        self.assertEqual(self.gutachten.text, "Neu")

    def test_edit_page_has_mde(self):
        url = reverse("gutachten_edit", args=[self.gutachten.pk])
        resp = self.client.get(url)
        self.assertContains(resp, "markdown_editor.js")

    def test_delete_removes_file(self):
        url = reverse("gutachten_delete", args=[self.gutachten.pk])
        resp = self.client.post(url)
        self.assertRedirects(resp, reverse("projekt_detail", args=[self.projekt.pk]))
        self.assertFalse(Gutachten.objects.filter(pk=self.gutachten.pk).exists())


class KnowledgeDescriptionEditTests(NoesisTestCase):
    def setUp(self):
        self.user = User.objects.create_user("kwuser", password="pass")
        self.client.login(username="kwuser", password="pass")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        self.knowledge = SoftwareKnowledge.objects.create(
            project=self.projekt,
            software_name="A",
            is_known_by_llm=True,
            description="Alt",
        )

    def test_edit_page_has_mde(self):
        url = reverse("edit_knowledge_description", args=[self.knowledge.pk])
        resp = self.client.get(url)
        self.assertContains(resp, "markdown_editor.js")

    def test_edit_updates_description(self):
        url = reverse("edit_knowledge_description", args=[self.knowledge.pk])
        resp = self.client.post(url, {"description": "Neu"})
        self.assertRedirects(resp, reverse("projekt_detail", args=[self.projekt.pk]))
        self.knowledge.refresh_from_db()
        self.assertEqual(self.knowledge.description, "Neu")


class Anlage5ReviewFormTests(NoesisTestCase):
    def test_get_json(self):
        cat = ZweckKategorieA.objects.create(beschreibung="A")
        form = Anlage5ReviewForm({"purposes": [cat.pk], "sonstige": "x"})
        self.assertTrue(form.is_valid())
        data = form.get_json()
        self.assertEqual(data["purposes"], [cat.pk])
        self.assertEqual(data["sonstige"], "x")


class BVProjectFileFormTests(NoesisTestCase):
    def test_extension_validation(self):
        form = BVProjectFileForm(
            {}, {"upload": SimpleUploadedFile("Anlage_1.pdf", b"d")}, anlage_nr=1
        )
        self.assertFalse(form.is_valid())

    def test_filename_pattern(self):
        form = BVProjectFileForm(
            {}, {"upload": SimpleUploadedFile("foo.docx", b"d")}, anlage_nr=1
        )
        self.assertTrue(form.is_valid())

    def test_max_size(self):
        form = BVProjectFileForm(
            {},
            {"upload": SimpleUploadedFile("Anlage_1.docx", b"x" * (settings.MAX_UPLOAD_SIZE + 1))},
            anlage_nr=1,
        )
        self.assertFalse(form.is_valid())

    def test_posted_anlage_nr_is_used(self):
        doc = Document()
        doc.add_paragraph("x")
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("Anlage_5.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)

        form = BVProjectFileForm({}, {"upload": upload}, anlage_nr=2)
        self.assertTrue(form.is_valid())
        obj = form.save(commit=False)
        self.assertEqual(obj.anlage_nr, 2)

    def test_manual_comment_is_optional(self):
        form = BVProjectFileForm(
            {"manual_comment": "Hinweis"},
            {"upload": SimpleUploadedFile("Anlage_1.docx", b"d")},
            anlage_nr=1,
        )
        self.assertTrue(form.is_valid())
        obj = form.save(commit=False)
        self.assertEqual(obj.manual_comment, "Hinweis")


