from pathlib import Path
from tempfile import NamedTemporaryFile

import pytest
from django.core.files.uploadedfile import SimpleUploadedFile
from docx import Document

from ..base import NoesisTestCase
from ...models import BVProject, BVProjectFile
from ...anlage3_parser import parse_anlage3

pytestmark = [pytest.mark.unit, pytest.mark.usefixtures("seed_db")]


class Anlage3ParserTests(NoesisTestCase):
    """Tests für den Anlage-3-Parser."""

    def _create_file(self, document: Document) -> BVProjectFile:
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        document.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("anlage3.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        return BVProjectFile.objects.create(project=projekt, anlage_nr=3, upload=upload)

    def test_verhandlungsfaehig_true(self):
        doc = Document()
        doc.add_paragraph("Nur eine Seite")
        pf = self._create_file(doc)
        data = parse_anlage3(pf)
        self.assertTrue(data["verhandlungsfaehig"])

    def test_verhandlungsfaehig_false(self):
        doc = Document()
        doc.add_paragraph("Seite 1")
        doc.add_page_break()
        doc.add_paragraph("Seite 2")
        pf = self._create_file(doc)
        data = parse_anlage3(pf)
        self.assertFalse(data["verhandlungsfaehig"])
