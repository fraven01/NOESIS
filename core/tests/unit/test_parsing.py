import pytest
from pathlib import Path
from tempfile import NamedTemporaryFile
from unittest.mock import patch

from django.core.files.uploadedfile import SimpleUploadedFile
from django.urls import reverse
from django.contrib.auth.models import User
from PIL import Image
from docx import Document

from ..base import NoesisTestCase
from ...anlage4_parser import parse_anlage4, parse_anlage4_dual
from ...docx_utils import (
    extract_text,
    get_docx_page_count,
    get_pdf_page_count,
    parse_anlage2_table,
    _normalize_header_text,
    extract_images,
)
from ...models import (
    BVProject,
    BVProjectFile,
    Anlage2Function,
    Anlage2Config,
    Anlage2ColumnHeading,
    Anlage2SubQuestion,
    Anlage4ParserConfig,
    Anlage4Config,
    AntwortErkennungsRegel,
)
from ...parsers import ExactParser
from ...text_parser import parse_anlage2_text
from ...llm_tasks import (
    analyse_anlage4,
    analyse_anlage4_async,
    worker_anlage4_evaluate,
    worker_a4_plausibility as check_anlage4_item_plausibility,
)

pytestmark = [pytest.mark.unit, pytest.mark.usefixtures("seed_db")]

@pytest.mark.usefixtures("prepared_files")
class DocxExtractTests(NoesisTestCase):
    def setUp(self) -> None:  # pragma: no cover - setup
        super().setUp()
        Anlage2Function.objects.all().delete()
        self.anmelden = Anlage2Function.objects.create(name="Anmelden")
        self.analyse = Anlage2Function.objects.create(name="Analyse")
        self.reporting = Anlage2Function.objects.create(name="Analyse-/Reportingfunktionen")
        self.anwesenheit = Anlage2Function.objects.create(name="Anwesenheit")

    def test_extract_text(self):
        text = extract_text(self.docx_content_path)
        self.assertIn("Docx Inhalt", text)

    def test_get_docx_page_count_single(self):
        count = get_docx_page_count(self.docx_content_path)
        self.assertEqual(count, 1)

    @pytest.mark.slow
    def test_get_docx_page_count_two_pages(self):
        count = get_docx_page_count(self.docx_two_page_path)
        self.assertEqual(count, 2)

    def test_get_pdf_page_count_single(self):
        count = get_pdf_page_count(self.pdf_one_page_path)
        self.assertEqual(count, 1)

    @pytest.mark.slow
    def test_get_pdf_page_count_two_pages(self):
        count = get_pdf_page_count(self.pdf_two_page_path)
        self.assertEqual(count, 2)

    def test_normalize_header_text_variants(self):
        cases = {
            "Technisch vorhanden?": "technisch vorhanden",
            "Technisch vorhanden:" : "technisch vorhanden",
            "Technisch   vorhanden": "technisch vorhanden",
            "Technisch\tvorhanden": "technisch vorhanden",
            " Verf\u00fcgbar?\t": "verf\u00fcgbar",
        }
        for raw, expected in cases.items():
            self.assertEqual(_normalize_header_text(raw), expected)

    @pytest.mark.slow
    def test_parse_anlage2_table(self):
        with patch("core.docx_utils.logging.getLogger") as mock_get_logger:
            mock_logger = mock_get_logger.return_value
            data = parse_anlage2_table(self.anlage2_table_docx_path)
            expected_raw = [
                "Funktion",
                "Technisch vorhanden",
                "Einsatz bei Telefónica",
                "Zur LV-Kontrolle",
                "KI-Beteiligung",
            ]
            expected_norm = [
                "funktion",
                "technisch vorhanden",
                "einsatz bei telefónica",
                "zur lv-kontrolle",
                "ki-beteiligung",
            ]
            mock_logger.debug.assert_any_call(
                f"Tabelle 0: Roh-Header = {expected_raw}, Normiert = {expected_norm}"
            )

        self.assertEqual(
            data,
            [
                {
                    "funktion": "Anmelden",
                    "technisch_verfuegbar": {"value": True, "note": None},
                    "einsatz_telefonica": {"value": False, "note": None},
                    "zur_lv_kontrolle": {"value": False, "note": None},
                    "ki_beteiligung": {"value": True, "note": None},
                }
            ],
        )

    @pytest.mark.slow
    def test_parse_anlage2_table_multiple_headers(self):
        cfg = Anlage2Config.get_instance()
        Anlage2ColumnHeading.objects.create(
            config=cfg, field_name="technisch_vorhanden", text="Verfügbar?"
        )
        Anlage2ColumnHeading.objects.create(
            config=cfg,
            field_name="technisch_vorhanden",
            text="Steht technisch zur Verfügung?",
        )
        Anlage2ColumnHeading.objects.create(
            config=cfg, field_name="einsatz_bei_telefonica", text="Telefonica Einsatz"
        )
        Anlage2ColumnHeading.objects.create(
            config=cfg, field_name="zur_lv_kontrolle", text="LV Kontrolle"
        )
        Anlage2ColumnHeading.objects.create(
            config=cfg, field_name="ki_beteiligung", text="KI?"
        )
        doc = Document()
        table = doc.add_table(rows=2, cols=5)
        table.cell(0, 0).text = "Funktion"
        table.cell(0, 1).text = "Steht technisch zur Verfügung?"
        table.cell(0, 2).text = "Telefonica Einsatz"
        table.cell(0, 3).text = "LV Kontrolle"
        table.cell(0, 4).text = "KI?"
        table.cell(1, 0).text = "Anmelden"
        table.cell(1, 1).text = "Ja"
        table.cell(1, 2).text = "Nein"
        table.cell(1, 3).text = "Nein"
        table.cell(1, 4).text = "Ja"
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        try:
            data = parse_anlage2_table(Path(tmp.name))
        finally:
            Path(tmp.name).unlink(missing_ok=True)

        self.assertTrue(data[0]["technisch_verfuegbar"]["value"])

    @pytest.mark.slow
    def test_parse_anlage2_table_alias_headers(self):
        cfg = Anlage2Config.get_instance()
        Anlage2ColumnHeading.objects.create(
            config=cfg,
            field_name="technisch_vorhanden",
            text="Steht technisch zur Verfügung?",
        )
        Anlage2ColumnHeading.objects.create(
            config=cfg,
            field_name="einsatz_bei_telefonica",
            text="einsatzweise bei telefónica: soll die funktion verwendet werden?",
        )
        Anlage2ColumnHeading.objects.create(
            config=cfg,
            field_name="zur_lv_kontrolle",
            text="einsatzweise bei telefónica: soll zur überwachung von leistung oder verhalten verwendet werden?",
        )

        doc = Document()
        table = doc.add_table(rows=2, cols=5)
        table.cell(0, 0).text = "Funktion"
        table.cell(0, 1).text = "Steht technisch zur Verfügung?"
        table.cell(0, 2).text = "einsatzweise bei Telefónica: soll die Funktion verwendet werden?"
        table.cell(0, 3).text = "einsatzweise bei Telefónica: soll zur Überwachung von Leistung oder Verhalten verwendet werden?"
        table.cell(0, 4).text = "KI-Beteiligung"

        table.cell(1, 0).text = "Anmelden"
        table.cell(1, 1).text = "Ja"
        table.cell(1, 2).text = "Nein"
        table.cell(1, 3).text = "Nein"
        table.cell(1, 4).text = "Ja"

        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        try:
            data = parse_anlage2_table(Path(tmp.name))
        finally:
            Path(tmp.name).unlink(missing_ok=True)

        self.assertEqual(
            data,
            [
                {
                    "funktion": "Anmelden",
                    "technisch_verfuegbar": {"value": True, "note": None},
                    "einsatz_telefonica": {"value": False, "note": None},
                    "zur_lv_kontrolle": {"value": False, "note": None},
                    "ki_beteiligung": {"value": True, "note": None},
                }
            ],
        )


    @pytest.mark.slow
    def test_parse_anlage2_table_extra_text(self):
        cfg = Anlage2Config.get_instance()
        Anlage2ColumnHeading.objects.create(
            config=cfg,
            field_name="technisch_vorhanden",
            text="Steht technisch zur Verf\u00fcgung?",
        )
        Anlage2ColumnHeading.objects.create(
            config=cfg,
            field_name="einsatz_bei_telefonica",
            text="einsatzweise bei telef\u00f3nica: soll die funktion verwendet werden?",
        )
        Anlage2ColumnHeading.objects.create(
            config=cfg,
            field_name="zur_lv_kontrolle",
            text="einsatzweise bei telef\u00f3nica: soll zur \u00fcberwachung von leistung oder verhalten verwendet werden?",
        )

        doc = Document()
        table = doc.add_table(rows=2, cols=5)
        table.cell(0, 0).text = "Funktion"
        table.cell(0, 1).text = "Steht technisch zur Verf\u00fcgung?\n\nJa/nein"
        table.cell(0, 2).text = (
            "einsatzweise bei Telef\u00f3nica: soll die Funktion verwendet werden?\n\nJa/nein"
        )
        table.cell(0, 3).text = (
            "einsatzweise bei Telef\u00f3nica: soll zur \u00dcberwachung von Leistung oder Verhalten verwendet werden?\n\nJa / nein"
        )
        table.cell(0, 4).text = "KI-Beteiligung"

        table.cell(1, 0).text = "Anmelden"
        table.cell(1, 1).text = "Ja"
        table.cell(1, 2).text = "Nein"
        table.cell(1, 3).text = "Nein"
        table.cell(1, 4).text = "Ja"

        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        try:
            data = parse_anlage2_table(Path(tmp.name))
        finally:
            Path(tmp.name).unlink(missing_ok=True)

        self.assertEqual(
            data,
            [
                {
                    "funktion": "Anmelden",
                    "technisch_verfuegbar": {"value": True, "note": None},
                    "einsatz_telefonica": {"value": False, "note": None},
                    "zur_lv_kontrolle": {"value": False, "note": None},
                    "ki_beteiligung": {"value": True, "note": None},
                }
            ],
        )

    @pytest.mark.slow
    def test_parse_anlage2_table_notes(self):
        doc = Document()
        table = doc.add_table(rows=2, cols=4)
        table.cell(0, 0).text = "Funktion"
        table.cell(0, 1).text = "Technisch vorhanden"
        table.cell(0, 2).text = "Einsatz bei Telefónica"
        table.cell(0, 3).text = "Zur LV-Kontrolle"

        table.cell(1, 0).text = "Anmelden"
        table.cell(1, 1).text = "Ja (nur intern)"
        table.cell(1, 2).text = "Nein, später"
        table.cell(1, 3).text = "Nein (k.A.)"

        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        try:
            data = parse_anlage2_table(Path(tmp.name))
        finally:
            Path(tmp.name).unlink(missing_ok=True)

        self.assertEqual(
            data,
            [
                {
                    "funktion": "Anmelden",
                    "technisch_verfuegbar": {"value": True, "note": "nur intern"},
                    "einsatz_telefonica": {"value": False, "note": "später"},
                    "zur_lv_kontrolle": {"value": False, "note": "k.A."},
                }
            ],
        )

    @pytest.mark.slow
    def test_parse_anlage2_table_alias_conflict(self):
        cfg = Anlage2Config.get_instance()
        conflict = "Konflikt"
        Anlage2ColumnHeading.objects.create(
            config=cfg, field_name="technisch_vorhanden", text=conflict
        )
        Anlage2ColumnHeading.objects.create(
            config=cfg, field_name="einsatz_bei_telefonica", text=conflict
        )

        doc = Document()
        table = doc.add_table(rows=2, cols=4)
        table.cell(0, 0).text = "Funktion"
        table.cell(0, 1).text = conflict
        table.cell(0, 2).text = "Einsatz bei Telefónica"
        table.cell(0, 3).text = "Zur LV-Kontrolle"

        table.cell(1, 0).text = "Anmelden"
        table.cell(1, 1).text = "Ja"
        table.cell(1, 2).text = "Nein"
        table.cell(1, 3).text = "Nein"

        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        try:
            with self.assertRaises(ValueError):
                parse_anlage2_table(Path(tmp.name))
        finally:
            Path(tmp.name).unlink(missing_ok=True)

    @pytest.mark.slow
    def test_parse_anlage2_table_invalid_docx(self):
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        tmp.write(b"invalid")
        tmp.close()
        try:
            data = parse_anlage2_table(Path(tmp.name))
        finally:
            Path(tmp.name).unlink(missing_ok=True)

        self.assertEqual(data, [])


    def test_parse_anlage2_text(self):
        func = self.anmelden
        func.detection_phrases = {"name_aliases": ["login"]}
        func.save()
        Anlage2SubQuestion.objects.filter(funktion=func).delete()
        Anlage2SubQuestion.objects.create(
            funktion=func,
            frage_text="Warum?",
            detection_phrases={"name_aliases": ["warum"]},
        )
        cfg = Anlage2Config.get_instance()
        cfg.text_technisch_verfuegbar_true = ["tv ja"]
        cfg.text_technisch_verfuegbar_false = ["tv nein"]
        cfg.text_ki_beteiligung_false = ["ki nein"]
        cfg.save()
        text = "Anmelden tv ja ki nein\nWarum? tv nein"
        data = parse_anlage2_text(text)
        self.assertEqual(
            data,
            [
                {
                    "funktion": "Anmelden",
                    "technisch_verfuegbar": {"value": True, "note": None},
                    "ki_beteiligung": {"value": False, "note": None},
                },
                {"funktion": "Anmelden: Warum?"},
            ],
        )

    def test_parse_anlage2_text_default_aliases(self):
        func = self.anmelden
        Anlage2SubQuestion.objects.filter(funktion=func).delete()
        Anlage2SubQuestion.objects.create(
            funktion=func,
            frage_text="Warum?",
        )
        cfg = Anlage2Config.get_instance()
        cfg.text_technisch_verfuegbar_true = ["tv ja"]
        cfg.text_technisch_verfuegbar_false = ["tv nein"]
        cfg.text_ki_beteiligung_false = ["ki nein"]
        cfg.save()
        text = "Anmelden tv ja ki nein\nWarum? tv nein"
        data = parse_anlage2_text(text)
        self.assertEqual(
            data,
            [
                {
                    "funktion": "Anmelden",
                    "technisch_verfuegbar": {"value": True, "note": None},
                    "ki_beteiligung": {"value": False, "note": None},
                },
                {"funktion": "Anmelden: Warum?"},
            ],
        )

    def test_parse_anlage2_text_name_always_alias(self):
        func = self.anmelden
        func.detection_phrases = {"name_aliases": ["login"]}
        func.save()
        Anlage2SubQuestion.objects.create(
            funktion=func,
            frage_text="Grund?",
            detection_phrases={"name_aliases": ["reason"]},
        )
        cfg = Anlage2Config.get_instance()
        cfg.text_technisch_verfuegbar_true = ["tv ja"]
        cfg.text_technisch_verfuegbar_false = ["tv nein"]
        cfg.save()
        text = "Anmelden tv ja\nGrund? tv nein"
        data = parse_anlage2_text(text)
        self.assertEqual(
            data,
            [
                {
                    "funktion": "Anmelden",
                    "technisch_verfuegbar": {"value": True, "note": None},
                },
                {"funktion": "Anmelden: Grund?"},
            ],
        )

    def test_parse_anlage2_text_normalizes_variants(self):
        func = self.anmelden
        cfg = Anlage2Config.get_instance()
        cfg.text_technisch_verfuegbar_true = ["tv ja"]
        cfg.save()
        text = "Anmelden   tv ja"
        data = parse_anlage2_text(text)
        self.assertEqual(
            data,
            [
                {
                    "funktion": "Anmelden",
                    "technisch_verfuegbar": {"value": True, "note": None},
                }
            ],
        )

    def test_parse_anlage2_text_punctuation_variants(self):
        func = self.reporting
        func.detection_phrases = {
            "name_aliases": ["Analyse-/Reportingfunktionen"]
        }
        func.save()
        cfg = Anlage2Config.get_instance()
        cfg.text_technisch_verfuegbar_true = ["tv ja"]
        cfg.save()
        text = "Analyse- / Reportingfunktionen tv ja"
        data = parse_anlage2_text(text)
        self.assertEqual(
            data,
            [
                {
                    "funktion": "Analyse-/Reportingfunktionen",
                    "technisch_verfuegbar": {"value": True, "note": None},
                }
            ],
        )

    def test_parse_anlage2_text_prefers_specific_subquestion(self):
        func = self.reporting
        Anlage2SubQuestion.objects.filter(funktion=func).delete()
        Anlage2SubQuestion.objects.create(
            funktion=func,
            frage_text="Bitte wähle zutreffendes aus",
        )
        cfg = Anlage2Config.get_instance()
        cfg.text_technisch_verfuegbar_true = ["ja"]
        cfg.save()
        text = "Analyse- / Reportingfunktionen - Bitte wähle zutreffendes aus: ja"
        data = parse_anlage2_text(text)
        self.assertEqual(
            data,
            [
                {
                    "funktion": "Analyse-/Reportingfunktionen: Bitte wähle zutreffendes aus",
                }
            ],
        )

    def test_parse_anlage2_text_merges_duplicate_functions(self):
        func = self.anmelden
        cfg = Anlage2Config.get_instance()
        cfg.text_technisch_verfuegbar_true = ["tv ja"]
        cfg.text_ki_beteiligung_false = ["ki nein"]
        cfg.save()
        text = "Anmelden tv ja\nAnmelden ki nein"
        data = parse_anlage2_text(text)
        self.assertEqual(
            data,
            [
                {
                    "funktion": "Anmelden",
                    "technisch_verfuegbar": {"value": True, "note": None},
                    "ki_beteiligung": {"value": False, "note": None},
                }
            ],
        )

    def test_parse_anlage2_text_updates_values_without_function(self):
        func = self.analyse
        cfg = Anlage2Config.get_instance()
        cfg.text_technisch_verfuegbar_true = ["verfuegbar"]
        cfg.text_zur_lv_kontrolle_false = ["kein lv"]
        cfg.save()
        text = "Analyse verfuegbar\nkein lv"
        data = parse_anlage2_text(text)
        self.assertEqual(
            data,
            [
                {
                    "funktion": "Analyse",
                    "technisch_verfuegbar": {"value": True, "note": None},
                    "zur_lv_kontrolle": {"value": False, "note": None},
                }
            ],
        )

    def test_parse_anlage2_text_fuzzy_match(self):
        func = self.anmelden
        cfg = Anlage2Config.get_instance()
        cfg.text_technisch_verfuegbar_true = ["ja"]
        cfg.save()
        data = parse_anlage2_text("Logn: ja")
        self.assertEqual(data, [])

    def test_parse_anlage2_text_fuzzy_token_phrase(self):
        func = self.anmelden
        cfg = Anlage2Config.get_instance()
        cfg.text_technisch_verfuegbar_true = ["ja bitte"]
        cfg.save()
        data = parse_anlage2_text("Lgin: ja bitte")
        self.assertEqual(data, [])

    def test_parse_anlage2_text_fuzzy_rule_phrase(self):
        func = self.anmelden
        AntwortErkennungsRegel.objects.create(
            regel_name="aktiv",
            erkennungs_phrase="aktivv",
            actions_json=[{"field": "einsatz_telefonica", "value": True}],
        )
        data = parse_anlage2_text("Lgin: aktivv")
        self.assertEqual(data, [])

    def test_parse_anlage2_text_multiple_rules_priority(self):
        func = self.anmelden
        AntwortErkennungsRegel.objects.create(
            regel_name="a",
            erkennungs_phrase="foo",
            actions_json=[{"field": "technisch_verfuegbar", "value": True}],
            prioritaet=2,
        )
        AntwortErkennungsRegel.objects.create(
            regel_name="b",
            erkennungs_phrase="bar",
            actions_json=[{"field": "einsatz_telefonica", "value": False}],
            prioritaet=1,
        )
        data = parse_anlage2_text("Anmelden: foo bar rest")
        self.assertEqual(
            data,
            [
                {
                    "funktion": "Anmelden",
                    "technisch_verfuegbar": {"value": True, "note": None},
                    "einsatz_telefonica": {"value": False, "note": "rest"},
                }
            ],
        )

    def test_parse_anlage2_text_unknown_question(self):
        _ = self.anmelden
        cfg = Anlage2Config.get_instance()
        cfg.text_technisch_verfuegbar_true = ["ja"]
        cfg.save()
        data = parse_anlage2_text("Unbekannt: ja\nAnmelden: ja")
        self.assertEqual(len(data), 1)
        self.assertEqual(
            data[0],
            {
                "funktion": "Anmelden",
                "technisch_verfuegbar": {"value": True, "note": None},
            },
        )

    def test_subquestion_skipped_when_main_absent(self):
        func = self.anwesenheit
        Anlage2SubQuestion.objects.filter(funktion=func).delete()
        Anlage2SubQuestion.objects.create(
            funktion=func,
            frage_text="Grund?",
        )
        cfg = Anlage2Config.get_instance()
        cfg.text_technisch_verfuegbar_false = ["nicht verfuegbar"]
        cfg.save()
        text = "Anwesenheit: technisch nicht verfuegbar\nGrund? nicht verfuegbar"
        data = parse_anlage2_text(text)
        self.assertEqual(
            data,
            [
                {
                    "funktion": "Anwesenheit",
                    "technisch_verfuegbar": {"value": False, "note": None},
                }
            ],
        )

    def test_subquestion_processed_when_main_present(self):
        func = self.anwesenheit
        Anlage2SubQuestion.objects.filter(funktion=func).delete()
        Anlage2SubQuestion.objects.create(
            funktion=func,
            frage_text="Grund?",
        )
        cfg = Anlage2Config.get_instance()
        cfg.text_technisch_verfuegbar_true = ["verfuegbar"]
        cfg.text_technisch_verfuegbar_false = ["nicht verfuegbar"]
        cfg.save()
        text = "Anwesenheit: verfuegbar\nGrund? nicht verfuegbar"
        data = parse_anlage2_text(text)
        self.assertEqual(
            data,
            [
                {
                    "funktion": "Anwesenheit",
                    "technisch_verfuegbar": {"value": True, "note": None},
                },
                {"funktion": "Anwesenheit: Grund?"},
            ],
        )

    def test_exact_parser_handles_segments(self):
        func1 = self.anmelden
        func2 = self.analyse
        AntwortErkennungsRegel.objects.create(
            regel_name="aktiv", erkennungs_phrase="aktiv",
            actions_json=[{"field": "technisch_verfuegbar", "value": True}],
        )
        AntwortErkennungsRegel.objects.create(
            regel_name="einsatz", erkennungs_phrase="kein einsatz",
            actions_json=[{"field": "einsatz_telefonica", "value": False}],
        )
        text = "Anmelden: aktiv\nAnalyse: kein einsatz"
        pf = BVProjectFile.objects.create(
            project=BVProject.objects.create(software_typen="A", beschreibung="x"),
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"x"),
            text_content=text,
        )
        data = ExactParser().parse(pf)
        self.assertEqual(
            data,
            [
                {
                    "funktion": "Anmelden",
                    "technisch_verfuegbar": {"value": True, "note": None},
                },
                {
                    "funktion": "Analyse",
                    "einsatz_telefonica": {"value": False, "note": None},
                },
            ],
        )

    def test_exact_parser_subquestion_requires_main(self):
        func = self.anmelden
        Anlage2SubQuestion.objects.filter(funktion=func).delete()
        Anlage2SubQuestion.objects.create(funktion=func, frage_text="Grund?")
        AntwortErkennungsRegel.objects.create(
            regel_name="tv", erkennungs_phrase="ja",
            actions_json=[{"field": "technisch_verfuegbar", "value": True}],
        )
        text = "Grund? ja"
        pf = BVProjectFile.objects.create(
            project=BVProject.objects.create(software_typen="A", beschreibung="x"),
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"x"),
            text_content=text,
        )
        data = ExactParser().parse(pf)
        self.assertEqual(data, [])

class DocxUtilsTests(NoesisTestCase):
    def test_extract_images(self):
        img = Image.new("RGB", (1, 1), color="blue")
        img_tmp = NamedTemporaryFile(delete=False, suffix=".png")
        img.save(img_tmp.name)
        img_tmp.close()
        doc = Document()
        doc.add_picture(img_tmp.name)
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        try:
            data = extract_images(Path(tmp.name))
        finally:
            Path(tmp.name).unlink(missing_ok=True)
            Path(img_tmp.name).unlink(missing_ok=True)
        self.assertEqual(len(data), 1)
        self.assertTrue(data[0].startswith(b"\x89PNG"))


class Anlage4ParserTests(NoesisTestCase):
    def test_parse_table_and_regex(self):
        cfg = Anlage4Config.objects.create(
            table_columns=["auswertung"],
            regex_patterns=[r"Zweck: (.+)"]
        )
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Zweck"
        table.cell(0, 1).text = "A"
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("t.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)
        pf = BVProjectFile.objects.create(
            project=BVProject.objects.create(software_typen="A"),
            anlage_nr=4,
            upload=upload,
            text_content="Zweck: B",
            anlage4_config=cfg,
        )
        items = parse_anlage4(pf)
        self.assertEqual(items, ["A"])

    def test_logs_table_detection(self):
        cfg = Anlage4Config.objects.create(table_columns=["auswertung"])
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Zweck"
        table.cell(0, 1).text = "A"
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("t.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)
        pf = BVProjectFile.objects.create(
            project=BVProject.objects.create(software_typen="A"),
            anlage_nr=4,
            upload=upload,
            anlage4_config=cfg,
        )
        with self.assertLogs("anlage4_detail", level="DEBUG") as cm:
            parse_anlage4(pf)
        self.assertIn("table detected - 1 items", cm.output[0])

    def test_negative_pattern(self):
        cfg = Anlage4Config.objects.create(negative_patterns=["keine zwecke"])
        pf = BVProjectFile.objects.create(
            project=BVProject.objects.create(software_typen="A"),
            anlage_nr=4,
            upload=SimpleUploadedFile("x.docx", b""),
            text_content="Keine Zwecke vorhanden",
            anlage4_config=cfg,
        )
        self.assertEqual(parse_anlage4(pf), [])

    def test_logs_free_text_detection(self):
        cfg = Anlage4Config.objects.create(regex_patterns=[r"Zweck: (.+)"])
        pf = BVProjectFile.objects.create(
            project=BVProject.objects.create(software_typen="A"),
            anlage_nr=4,
            upload=SimpleUploadedFile("x.txt", b""),
            text_content="Zweck: A",
            anlage4_config=cfg,
        )
        with self.assertLogs("anlage4_detail", level="DEBUG") as cm:
            parse_anlage4(pf)
        self.assertIn("free text found - 1 items", "".join(cm.output))

    def test_dual_parser_handles_invalid_rules(self):
        pcfg = Anlage4ParserConfig.objects.create(delimiter_phrase="")
        pf = BVProjectFile.objects.create(
            project=BVProject.objects.create(software_typen="A"),
            anlage_nr=4,
            upload=SimpleUploadedFile("x.txt", b""),
            text_content="",
            anlage4_parser_config=pcfg,
        )
        items = parse_anlage4_dual(pf)
        self.assertEqual(items, [])

    def test_dual_parser_extracts_fields(self):
        pcfg = Anlage4ParserConfig.objects.create(
            delimiter_phrase="Name",
            gesellschaften_phrase="Gesellschaft",
            fachbereiche_phrase="Bereich",
        )
        text = (
            "Name A\nGesellschaft X\nBereich Y\n"
            "Name B\nGesellschaft Z\nBereich W"
        )
        pf = BVProjectFile.objects.create(
            project=BVProject.objects.create(software_typen="A"),
            anlage_nr=4,
            upload=SimpleUploadedFile("x.txt", b""),
            text_content=text,
            anlage4_parser_config=pcfg,
        )
        items = parse_anlage4_dual(pf)
        self.assertEqual(
            items,
            [
                {"name_der_auswertung": "A", "gesellschaften": "X", "fachbereiche": "Y"},
                {"name_der_auswertung": "B", "gesellschaften": "Z", "fachbereiche": "W"},
            ],
        )

    def test_dual_parser_uses_alias_lists(self):
        pcfg = Anlage4ParserConfig.objects.create(
            name_aliases=["Name", "Auswertung"],
            gesellschaft_aliases=["Gesellschaft"],
            fachbereich_aliases=["Bereich"],
        )
        text = (
            "Name A\nGesellschaft X\nBereich Y\n"
            "Auswertung B\nGesellschaft Z\nBereich W"
        )
        pf = BVProjectFile.objects.create(
            project=BVProject.objects.create(software_typen="A"),
            anlage_nr=4,
            upload=SimpleUploadedFile("x.txt", b""),
            text_content=text,
            anlage4_parser_config=pcfg,
        )
        items = parse_anlage4_dual(pf)
        self.assertEqual(
            items,
            [
                {"name_der_auswertung": "A", "gesellschaften": "X", "fachbereiche": "Y"},
                {"name_der_auswertung": "B", "gesellschaften": "Z", "fachbereiche": "W"},
            ],
        )

    def test_dual_parser_strips_colon(self):
        pcfg = Anlage4ParserConfig.objects.create(
            delimiter_phrase="Zweck",
            gesellschaften_phrase="",
            fachbereiche_phrase="",
        )
        pf = BVProjectFile.objects.create(
            project=BVProject.objects.create(software_typen="A"),
            anlage_nr=4,
            upload=SimpleUploadedFile("x.txt", b""),
            text_content="Zweck: Logdaten",
            anlage4_parser_config=pcfg,
        )
        items = parse_anlage4_dual(pf)
        self.assertEqual(items[0]["name_der_auswertung"], "Logdaten")

    def test_dual_parser_default_config_example(self):
        pcfg = Anlage4ParserConfig.objects.create()
        text = (
            "Name der 1. Auswertung\n"
            "Alpha\n"
            "   Gesellschaften, in denen die Auswertung verwendet wird: Foo\n"
            "   Fachbereiche, in denen die Auswertung eingesetzt wird: Bar\n"
            "Name der 2. Auswertung\n"
            "Beta\n"
            "   Gesellschaften, in denen die Auswertung verwendet wird: Baz\n"
            "   Fachbereiche, in denen die Auswertung eingesetzt wird: Qux"
        )
        pf = BVProjectFile.objects.create(
            project=BVProject.objects.create(software_typen="A"),
            anlage_nr=4,
            upload=SimpleUploadedFile("x.txt", b""),
            text_content=text,
            anlage4_parser_config=pcfg,
        )
        items = parse_anlage4_dual(pf)
        self.assertEqual(
            items,
            [
                {"name_der_auswertung": "Alpha", "gesellschaften": "Foo", "fachbereiche": "Bar"},
                {"name_der_auswertung": "Beta", "gesellschaften": "Baz", "fachbereiche": "Qux"},
            ],
        )

    def test_dual_parser_table_columns(self):
        pcfg = Anlage4ParserConfig.objects.create(
            table_columns=[
                "Name der Auswertung",
                "Gesellschaften, in denen die Auswertung verwendet wird",
                "Fachbereiche, in denen die Auswertung eingesetzt wird",
            ]
        )
        doc = Document()
        table = doc.add_table(rows=3, cols=3)
        table.cell(0, 0).text = "Name der Auswertung"
        table.cell(0, 1).text = "Alpha"
        table.cell(0, 2).text = "Beta"
        table.cell(1, 0).text = "Gesellschaften, in denen die Auswertung verwendet wird"
        table.cell(1, 1).text = "G1"
        table.cell(1, 2).text = "G2"
        table.cell(2, 0).text = "Fachbereiche, in denen die Auswertung eingesetzt wird"
        table.cell(2, 1).text = "F1"
        table.cell(2, 2).text = "F2"
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("t.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)
        pf = BVProjectFile.objects.create(
            project=BVProject.objects.create(software_typen="A"),
            anlage_nr=4,
            upload=upload,
            anlage4_parser_config=pcfg,
        )
        items = parse_anlage4_dual(pf)
        self.assertEqual(
            items,
            [
                {"name_der_auswertung": "Alpha", "gesellschaften": "G1", "fachbereiche": "F1"},
                {"name_der_auswertung": "Beta", "gesellschaften": "G2", "fachbereiche": "F2"},
            ],
        )

    def test_dual_parser_multiple_tables(self):
        pcfg = Anlage4ParserConfig.objects.create(
            table_columns=[
                "Name der Auswertung",
                "Gesellschaften, in denen die Auswertung verwendet wird",
                "Fachbereiche, in denen die Auswertung eingesetzt wird",
            ]
        )
        doc = Document()
        t1 = doc.add_table(rows=3, cols=3)
        t1.cell(0, 0).text = "Name der Auswertung"
        t1.cell(0, 1).text = "Alpha"
        t1.cell(0, 2).text = "Beta"
        t1.cell(1, 0).text = "Gesellschaften, in denen die Auswertung verwendet wird"
        t1.cell(1, 1).text = "G1"
        t1.cell(1, 2).text = "G2"
        t1.cell(2, 0).text = "Fachbereiche, in denen die Auswertung eingesetzt wird"
        t1.cell(2, 1).text = "F1"
        t1.cell(2, 2).text = "F2"

        t2 = doc.add_table(rows=3, cols=3)
        t2.cell(0, 0).text = "Name der Auswertung"
        t2.cell(0, 1).text = "Gamma"
        t2.cell(0, 2).text = "Delta"
        t2.cell(1, 0).text = "Gesellschaften, in denen die Auswertung verwendet wird"
        t2.cell(1, 1).text = "G3"
        t2.cell(1, 2).text = "G4"
        t2.cell(2, 0).text = "Fachbereiche, in denen die Auswertung eingesetzt wird"
        t2.cell(2, 1).text = "F3"
        t2.cell(2, 2).text = "F4"

        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("t.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)
        pf = BVProjectFile.objects.create(
            project=BVProject.objects.create(software_typen="A"),
            anlage_nr=4,
            upload=upload,
            anlage4_parser_config=pcfg,
        )
        items = parse_anlage4_dual(pf)
        self.assertEqual(
            items,
            [
                {"name_der_auswertung": "Alpha", "gesellschaften": "G1", "fachbereiche": "F1"},
                {"name_der_auswertung": "Beta", "gesellschaften": "G2", "fachbereiche": "F2"},
                {"name_der_auswertung": "Gamma", "gesellschaften": "G3", "fachbereiche": "F3"},
                {"name_der_auswertung": "Delta", "gesellschaften": "G4", "fachbereiche": "F4"},
            ],
        )

    def test_dual_parser_column_key_mapping(self):
        pcfg = Anlage4ParserConfig.objects.create(
            table_columns=[
                "Name der Auswertung",
                "Gesellschaften, in denen die Auswertung verwendet wird",
                "Fachbereiche, in denen die Auswertung eingesetzt wird",
            ]
        )
        doc = Document()
        table = doc.add_table(rows=3, cols=2)
        table.cell(0, 0).text = "Name der Auswertung"
        table.cell(0, 1).text = "Alpha"
        table.cell(1, 0).text = "Gesellschaften, in denen die Auswertung verwendet wird"
        table.cell(1, 1).text = "G1"
        table.cell(2, 0).text = "Fachbereiche, in denen die Auswertung eingesetzt wird"
        table.cell(2, 1).text = "F1"
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("t.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)
        pf = BVProjectFile.objects.create(
            project=BVProject.objects.create(software_typen="A"),
            anlage_nr=4,
            upload=upload,
            anlage4_parser_config=pcfg,
        )
        items = parse_anlage4_dual(pf)
        self.assertEqual(list(items[0].keys()), ["name_der_auswertung", "gesellschaften", "fachbereiche"])

    def test_dual_parser_negative_pattern(self):
        pcfg = Anlage4ParserConfig.objects.create(
            name_aliases=["Name"],
            gesellschaft_aliases=["Gesellschaft"],
            fachbereich_aliases=["Bereich"],
            negative_patterns=["Keine Auswertung"],
        )
        text = "Keine Auswertung vorhanden\nName A\nGesellschaft X\nBereich Y"
        pf = BVProjectFile.objects.create(
            project=BVProject.objects.create(software_typen="A"),
            anlage_nr=4,
            upload=SimpleUploadedFile("x.txt", b""),
            text_content=text,
            anlage4_parser_config=pcfg,
        )
        self.assertEqual(parse_anlage4_dual(pf), [])

    def test_dual_parser_no_config(self):
        pf = BVProjectFile.objects.create(
            project=BVProject.objects.create(software_typen="A"),
            anlage_nr=4,
            upload=SimpleUploadedFile("x.txt", b""),
            text_content="",
        )
        result = {"items": parse_anlage4(pf)}
        self.assertEqual(result["items"], [])


class AnalyseAnlage4Tests(NoesisTestCase):
    def test_task_stores_json(self):
        """Prüft, dass der Parser Stammdaten korrekt speichert."""
        pcfg = Anlage4ParserConfig.objects.create(
            delimiter_phrase="Name",
            gesellschaften_phrase="Gesellschaft",
            fachbereiche_phrase="Bereich",
        )
        projekt = BVProject.objects.create(software_typen="A", title="Test")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=4,
            upload=SimpleUploadedFile("a.txt", b""),
            text_content="Name A\nGesellschaft X\nBereich Y",
            anlage4_parser_config=pcfg,
        )
        with patch("core.llm_tasks.worker_a4_plausibility"), patch(
            "core.llm_tasks.worker_anlage4_evaluate"
        ):
            data = analyse_anlage4_async(pf.pk)
        pf.refresh_from_db()
        expected = {
            "name_der_auswertung": "A",
            "gesellschaften": "X",
            "fachbereiche": "Y",
        }
        self.assertEqual(data["items"][0]["structured"], expected)
        self.assertEqual(pf.analysis_json["items"][0]["structured"], expected)

    def test_check_anlage4_item_plausibility(self):
        """LLM-Task ergänzt Plausibilitätsdaten im JSON."""
        pcfg = Anlage4ParserConfig.objects.create(
            delimiter_phrase="Name",
            gesellschaften_phrase="Gesellschaft",
            fachbereiche_phrase="Bereich",
        )
        projekt = BVProject.objects.create(software_typen="A", title="Projekt")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=4,
            upload=SimpleUploadedFile("a.txt", b""),
            text_content="Name A\nGesellschaft X\nBereich Y",
            anlage4_parser_config=pcfg,
        )
        # Parser ausführen ohne LLM-Prüfung
        with patch("core.llm_tasks.worker_a4_plausibility"), patch(
            "core.llm_tasks.worker_anlage4_evaluate"
        ):
            analyse_anlage4_async(pf.pk)
        pf.refresh_from_db()

        structured = {
            **pf.analysis_json["items"][0]["structured"],
            "kontext": projekt.title,
        }
        with patch(
            "core.llm_tasks.query_llm",
            return_value='{"plausibilitaet":"hoch","score":0.8,"begruendung":"ok"}',
        ):
            check_anlage4_item_plausibility(structured, pf.pk, 0)

        pf.refresh_from_db()
        item = pf.analysis_json["items"][0]["plausibility"]
        self.assertEqual(item["plausibilitaet"], "hoch")
        self.assertEqual(item["score"], 0.8)
        self.assertEqual(item["begruendung"], "ok")

    def test_passes_config_to_parser(self):
        """Die Analyse nutzt die übergebene Konfiguration."""
        cfg = Anlage4Config.objects.create(regex_patterns=[r"Zweck: (.+)"])
        cfg.delimiter_phrase = "Y"
        cfg.save()
        Anlage4ParserConfig.objects.all().delete()
        projekt = BVProject.objects.create(software_typen="A")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=4,
            upload=SimpleUploadedFile("a.txt", b""),
            text_content="Zweck: A",
            anlage4_config=cfg,
        )
        from core.llm_tasks import analyse_anlage4_async as _analyse_anlage4_worker

        with patch(
            "core.llm_tasks.parse_anlage4", return_value=[]
        ) as m_parse, patch("core.llm_tasks.worker_a4_plausibility"), patch(
            "core.llm_tasks.worker_anlage4_evaluate"
        ):
            _analyse_anlage4_worker(pf.pk)
        m_parse.assert_called_once_with(pf, cfg)

    def test_dual_parser_used_when_parser_config(self):
        pcfg = Anlage4ParserConfig.objects.create(table_columns=["auswertung"])
        projekt = BVProject.objects.create(software_typen="A")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=4,
            upload=SimpleUploadedFile("a.txt", b""),
            text_content="Zweck: A",
            anlage4_parser_config=pcfg,
        )
        with patch("core.llm_tasks.parse_anlage4_dual", return_value=[] ) as m_dual, patch(
            "core.llm_tasks.parse_anlage4"
        ) as m_std, patch("core.llm_tasks.query_llm", return_value="{}"):
            analyse_anlage4(projekt.pk)
        m_dual.assert_called_once_with(pf)
        m_std.assert_not_called()

    def test_template_allows_json_data_placeholder(self):
        cfg = Anlage4Config.objects.create(prompt_template="Vorlage {json_data}")
        projekt = BVProject.objects.create(software_typen="A")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=4,
            upload=SimpleUploadedFile("a.txt", b""),
            text_content="Zweck: A",
            anlage4_config=cfg,
        )
        with patch("core.llm_tasks.query_llm", return_value="{}"):  # kein Fehler
            analyse_anlage4(projekt.pk)
        pf.refresh_from_db()
        self.assertIsNotNone(pf.analysis_json)


class AnalyseAnlage4AsyncTests(NoesisTestCase):
    def test_async_analysis_stores_results(self):
        cfg = Anlage4Config.objects.create(regex_patterns=[r"Zweck: (.+)"])
        projekt = BVProject.objects.create(software_typen="A")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=4,
            upload=SimpleUploadedFile("a.txt", b""),
            text_content="Zweck: A\nZweck: B",
            anlage4_config=cfg,
        )

        def immediate(name, *args):
            self.assertEqual(name, "core.llm_tasks.worker_anlage4_evaluate")
            worker_anlage4_evaluate(*args)

        with patch("core.llm_tasks.async_task", side_effect=immediate), patch(
            "core.llm_tasks.query_llm",
            return_value='{"plausibilitaet":"hoch","score":0.8,"begruendung":"ok"}',
        ):
            analyse_anlage4_async(pf.pk)

        pf.refresh_from_db()
        results = pf.analysis_json["items"]
        expected = ["A", "B"]
        for idx, item in enumerate(results):
            self.assertEqual(item["text"], expected[idx])
            self.assertEqual(item["structured"]["name_der_auswertung"], expected[idx])
            self.assertEqual(item["plausibility"]["plausibilitaet"], "hoch")
            self.assertEqual(item["plausibility"]["begruendung"], "ok")

    def test_async_dual_parser_used_when_parser_config(self):
        pcfg = Anlage4ParserConfig.objects.create(delimiter_phrase="Zweck")
        projekt = BVProject.objects.create(software_typen="A")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=4,
            upload=SimpleUploadedFile("a.txt", b""),
            text_content="Zweck: A",
            anlage4_parser_config=pcfg,
        )

        with patch(
            "core.llm_tasks.parse_anlage4_dual",
            return_value=[{"name_der_auswertung": "A"}],
        ) as m_dual, patch(
            "core.llm_tasks.parse_anlage4"
        ) as m_std, patch("core.llm_tasks.async_task") as m_task, patch(
            "core.llm_tasks.query_llm", return_value="{}"
        ):
            analyse_anlage4_async(pf.pk)
        m_dual.assert_called_once_with(pf)
        m_std.assert_not_called()


class Anlage4ReviewViewTests(NoesisTestCase):
    def setUp(self):
        self.user = User.objects.create_user("rev4", password="pass")
        self.client.login(username="rev4", password="pass")
        self.projekt = BVProject.objects.create(software_typen="A")
        self.file = BVProjectFile.objects.create(
            project=self.projekt,
            anlage_nr=4,
            upload=SimpleUploadedFile("a.txt", b""),
            analysis_json={
                "items": [{"text": "A"}, {"text": "B"}],
                "manual_review": {
                    "0": {"ok": True, "nego": False, "note": "alt"},
                    "1": {"ok": False, "nego": True, "note": "vorher"},
                },
            },
        )

    def test_post_saves_manual_review(self):
        url = reverse("anlage4_review", args=[self.file.pk])
        resp = self.client.post(url, {"item1_note": "neu"})
        self.assertRedirects(resp, reverse("projekt_detail", args=[self.projekt.pk]))
        self.file.refresh_from_db()
        self.assertEqual(
            self.file.analysis_json.get("manual_review"),
            {
                "0": {"ok": True, "nego": False, "note": "alt"},
                "1": {"ok": False, "nego": True, "note": "neu"},
            },
        )


class ProjektFileAnalyseAnlage4ViewTests(NoesisTestCase):
    def setUp(self):
        self.user = User.objects.create_user("a4user", password="pass")
        self.client.login(username="a4user", password="pass")
        self.projekt = BVProject.objects.create(software_typen="A")
        self.file = BVProjectFile.objects.create(
            project=self.projekt,
            anlage_nr=4,
            upload=SimpleUploadedFile("a.txt", b""),
            text_content="Zwecke",
            analysis_json={},
        )

    def test_get_runs_analysis_and_redirects(self):
        url = reverse("projekt_file_analyse_anlage4", args=[self.file.pk])
        with patch("core.views.connection.vendor", new="postgresql"), patch(
            "core.views.async_task"
        ) as mock_task:
            resp = self.client.get(url)
        self.assertRedirects(resp, reverse("anlage4_review", args=[self.file.pk]))
        mock_task.assert_called_with(
            "core.llm_tasks.analyse_anlage4_async", self.file.pk
        )

    def test_post_runs_analysis_and_redirects(self):
        url = reverse("projekt_file_analyse_anlage4", args=[self.file.pk])
        with patch("core.views.connection.vendor", new="postgresql"), patch(
            "core.views.async_task"
        ) as mock_task:
            resp = self.client.post(url)
        self.assertRedirects(resp, reverse("anlage4_review", args=[self.file.pk]))
        mock_task.assert_called_with(
            "core.llm_tasks.analyse_anlage4_async", self.file.pk
        )


class WorkerAnlage4EvaluateTests(NoesisTestCase):
    def test_worker_adds_structured(self):
        projekt = BVProject.objects.create(software_typen="A")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=4,
            upload=SimpleUploadedFile("a.txt", b""),
            text_content="",
        )

        with patch(
            "core.llm_tasks.query_llm",
            return_value='{"plausibilitaet":"hoch","score":0.9,"begruendung":"ok"}',
        ):
            worker_anlage4_evaluate("A", pf.pk, 0)

        pf.refresh_from_db()
        item = pf.analysis_json["items"][0]
        self.assertEqual(item["structured"]["name_der_auswertung"], "A")
        self.assertEqual(item["plausibility"]["plausibilitaet"], "hoch")

    def test_worker_handles_code_fences(self):
        projekt = BVProject.objects.create(software_typen="A")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=4,
            upload=SimpleUploadedFile("a.txt", b""),
            text_content="",
        )

        with patch(
            "core.llm_tasks.query_llm",
            return_value="```json\n{\"plausibilitaet\":\"hoch\",\"score\":0.9,\"begruendung\":\"ok\"}\n```",
        ):
            worker_anlage4_evaluate("A", pf.pk, 0)

        pf.refresh_from_db()
        item = pf.analysis_json["items"][0]
        self.assertEqual(item["plausibility"]["score"], 0.9)

