from pathlib import Path
from tempfile import NamedTemporaryFile

from django.core.files.uploadedfile import SimpleUploadedFile
from docx import Document

from .base import NoesisTestCase
from ..models import BVProject, BVProjectFile
from ..anlage3_parser import parse_anlage3


class Anlage3ParserTests(NoesisTestCase):
    """Tests für den Anlage-3-Parser."""

    def _create_file(self, document: Document) -> BVProjectFile:
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        document.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("anlage3.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        return BVProjectFile.objects.create(project=projekt, anlage_nr=3, upload=upload)

    def test_verhandlungsfaehig_true(self):
        doc = Document()
        table = doc.add_table(rows=4, cols=2)
        table.cell(0, 0).text = "Name der Auswertung"
        table.cell(0, 1).text = "Test"
        table.cell(1, 0).text = "Beschreibung"
        table.cell(1, 1).text = "Desc"
        table.cell(2, 0).text = "Zeitraum"
        table.cell(2, 1).text = "2024"
        table.cell(3, 0).text = "Art der Auswertung"
        table.cell(3, 1).text = "Art"
        pf = self._create_file(doc)
        data = parse_anlage3(pf)
        self.assertTrue(data["verhandlungsfaehig"])

    def test_verhandlungsfaehig_false(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Name der Auswertung"
        table.cell(0, 1).text = "Test"
        pf = self._create_file(doc)
        data = parse_anlage3(pf)
        self.assertFalse(data["verhandlungsfaehig"])
