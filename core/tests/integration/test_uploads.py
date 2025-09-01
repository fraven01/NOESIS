import pytest
from django.contrib.auth.models import User
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test.utils import override_settings
from django.urls import reverse
from pathlib import Path
from tempfile import NamedTemporaryFile
from docx import Document
from PIL import Image
from unittest.mock import Mock, patch, call

from ...models import (
    BVProject,
    BVProjectFile,
    FunktionsErgebnis,
    Anlage2Function,
    AnlagenFunktionsMetadaten,
)
from ...forms import BVProjectFileForm
from ...views import _save_project_file
from ..base import NoesisTestCase

pytestmark = [pytest.mark.integration, pytest.mark.usefixtures("seed_db")]


@override_settings(ALLOWED_HOSTS=["testserver"])
@pytest.mark.usefixtures("prepared_files")
class ProjektFileUploadTests(NoesisTestCase):
    def setUp(self):
        self.user = User.objects.create_user("user", password="pass")
        self.client.login(username="user", password="pass")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        self.anmelden_func = Anlage2Function.objects.create(name="Anmelden")

    @pytest.mark.slow
    def test_docx_upload_extracts_text(self):
        with open(self.docx_content_path, "rb") as fh:
            upload = SimpleUploadedFile("Anlage_1.docx", fh.read())

        url = reverse("hx_project_file_upload", args=[self.projekt.pk])
        resp = self.client.post(
            url,
            {"anlage_nr": 1, "upload": upload, "manual_comment": ""},
            format="multipart",
            HTTP_HX_REQUEST="true",
        )
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.content.decode().count("<tr"), 1)
        file_obj = self.projekt.anlagen.first()
        self.assertIsNotNone(file_obj)
        self.assertIn("Docx Inhalt", file_obj.text_content)

    def test_ownerless_project_allows_upload(self):
        """Prüft, dass ein neuer Nutzer Dateien in ein besitzerloses Projekt laden darf."""
        with open(self.docx_content_path, "rb") as fh:
            upload = SimpleUploadedFile("Anlage_1.docx", fh.read())

        url = reverse("hx_project_file_upload", args=[self.projekt.pk])
        resp = self.client.post(
            url,
            {"anlage_nr": 1, "upload": upload, "manual_comment": ""},
            format="multipart",
            HTTP_HX_REQUEST="true",
        )
        self.assertEqual(resp.status_code, 200)

    def test_pdf_upload_stores_bytes(self):
        with open(self.pdf_one_page_path, "rb") as fh:
            upload = SimpleUploadedFile("Anlage_3.pdf", fh.read())

        url = reverse("hx_project_file_upload", args=[self.projekt.pk])
        resp = self.client.post(
            url,
            {"anlage_nr": 3, "upload": upload, "manual_comment": ""},
            format="multipart",
            HTTP_HX_REQUEST="true",
        )
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.content.decode().count("<tr"), 1)
        file_obj = self.projekt.anlagen.get(anlage_nr=3)
        self.assertEqual(file_obj.text_content, "")

    def test_upload_without_anlage_nr_uses_filename(self):
        """Nutzt die Anlagen-Nummer aus dem Dateinamen."""
        with open(self.docx_content_path, "rb") as fh:
            upload = SimpleUploadedFile("Anlage 4 - Entwurf.docx", fh.read())

        url = reverse("hx_project_file_upload", args=[self.projekt.pk])
        resp = self.client.post(
            url,
            {"upload": upload, "manual_comment": ""},
            format="multipart",
            HTTP_HX_REQUEST="true",
        )
        self.assertEqual(resp.status_code, 200)
        self.assertTrue(self.projekt.anlagen.filter(anlage_nr=4).exists())

    def test_anlage2_upload_queues_check(self):
        with open(self.docx_content_path, "rb") as fh:
            upload = SimpleUploadedFile("Anlage_2.docx", fh.read())

        _ = self.anmelden_func

        url = reverse("hx_project_file_upload", args=[self.projekt.pk])
        mock_async = Mock(side_effect=["tid1", "tid2"])

        def fake_start(file_id: int) -> str:
            pf_obj = BVProjectFile.objects.get(pk=file_id)
            pf_obj.processing_status = BVProjectFile.PROCESSING
            pf_obj.save(update_fields=["processing_status"])
            task_id = None
            for func, arg in pf_obj.get_analysis_tasks():
                tid = mock_async(func, arg)
                if task_id is None:
                    task_id = tid
            return task_id or ""

        with patch("core.signals.start_analysis_for_file", side_effect=fake_start):
            resp = self.client.post(
                url,
                {"anlage_nr": 2, "upload": upload, "manual_comment": ""},
                format="multipart",
                HTTP_HX_REQUEST="true",
            )
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.content.decode().count("<tr"), 1)
        pf = self.projekt.anlagen.get(anlage_nr=2)
        self.assertEqual(pf.verification_task_id, "tid1")
        self.assertEqual(pf.processing_status, BVProjectFile.PROCESSING)
        mock_async.assert_any_call(
            "core.llm_tasks.run_conditional_anlage2_check",
            pf.pk,
        )

    def test_second_anlage2_version_skips_ai_check(self):
        func = self.anmelden_func
        first = BVProjectFile.objects.create(
            project=self.projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("v1.docx", b"x"),
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=first,
            funktion=func,
            quelle="ki",
            technisch_verfuegbar=True,
        )

        with open(self.docx_content_path, "rb") as fh:
            upload = SimpleUploadedFile("Anlage_2.docx", fh.read())

        url = reverse("hx_project_file_upload", args=[self.projekt.pk])
        with patch("core.views.async_task") as mock_async:
            resp = self.client.post(
                url,
                {"anlage_nr": 2, "upload": upload, "manual_comment": ""},
                format="multipart",
                HTTP_HX_REQUEST="true",
            )
        self.assertEqual(resp.status_code, 200)
        self.assertFalse(
            any(
                call.args[0] == "core.llm_tasks.run_conditional_anlage2_check"
                for call in mock_async.call_args_list
            )
        )
        pf_latest = self.projekt.anlagen.filter(anlage_nr=2, is_active=True).first()
        self.assertEqual(pf_latest.version, 2)

    def test_upload_stores_posted_anlage_nr(self):
        with open(self.docx_content_path, "rb") as fh:
            upload = SimpleUploadedFile("Anlage_5.docx", fh.read())
        url = reverse("hx_project_file_upload", args=[self.projekt.pk])
        resp = self.client.post(
            url,
            {"anlage_nr": 2, "upload": upload, "manual_comment": ""},
            format="multipart",
            HTTP_HX_REQUEST="true",
        )
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.content.decode().count("<tr"), 1)
        pf = self.projekt.anlagen.get()
        self.assertEqual(pf.anlage_nr, 2)

    def test_save_project_file_respects_form_value(self):
        with open(self.docx_content_path, "rb") as fh:
            upload = SimpleUploadedFile("Anlage_5.docx", fh.read())

        form = BVProjectFileForm({}, {"upload": upload}, anlage_nr=1)
        self.assertTrue(form.is_valid())
        pf = _save_project_file(self.projekt, form)
        self.assertEqual(pf.anlage_nr, 1)

    def test_new_version_has_empty_gap_fields(self):
        with open(self.docx_content_path, "rb") as fh:
            upload = SimpleUploadedFile("Anlage_2.docx", fh.read())
        pf1 = _save_project_file(self.projekt, upload=upload, anlage_nr=2)
        pf1.gap_summary = "Alt"
        pf1.gap_notiz = "Notiz"
        pf1.save(update_fields=["gap_summary", "gap_notiz"])
        FunktionsErgebnis.objects.create(
            anlage_datei=pf1, funktion=self.anmelden_func, quelle="ki"
        )
        AnlagenFunktionsMetadaten.objects.create(
            anlage_datei=pf1,
            funktion=self.anmelden_func,
            gap_summary="S",
            gap_notiz="N",
        )

        with open(self.docx_content_path, "rb") as fh:
            upload = SimpleUploadedFile("Anlage_2_neu.docx", fh.read())
        pf2 = _save_project_file(self.projekt, upload=upload, anlage_nr=2)
        meta = AnlagenFunktionsMetadaten.objects.get(
            anlage_datei=pf2, funktion=self.anmelden_func
        )
        self.assertEqual(meta.gap_summary, "")
        self.assertEqual(meta.gap_notiz, "")
        self.assertEqual(pf2.gap_summary, "")
        self.assertEqual(pf2.gap_notiz, "")

    def test_save_multiple_files_unique_numbers(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        for nr in range(1, 7):
            with open(self.docx_content_path, "rb") as fh:
                upload = SimpleUploadedFile(f"Anlage_{nr}.docx", fh.read())
            _save_project_file(projekt, upload=upload, anlage_nr=nr)

        qs = BVProjectFile.objects.filter(project=projekt)
        self.assertEqual(qs.count(), 6)
        self.assertListEqual(
            sorted(qs.values_list("anlage_nr", flat=True)),
            [1, 2, 3, 4, 5, 6],
        )

    def test_upload_uses_filename_when_no_anlage_nr(self):
        with open(self.docx_content_path, "rb") as fh:
            upload = SimpleUploadedFile("Anlage_4.docx", fh.read())

        url = reverse("hx_project_file_upload", args=[self.projekt.pk])
        resp = self.client.post(
            url,
            {"upload": upload, "manual_comment": ""},
            format="multipart",
            HTTP_HX_REQUEST="true",
        )
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.content.decode().count("<tr"), 1)
        pf = self.projekt.anlagen.get()
        self.assertEqual(pf.anlage_nr, 4)

    def test_upload_uses_filename_when_anlage_nr_empty(self):
        with open(self.docx_content_path, "rb") as fh:
            upload = SimpleUploadedFile("Anlage_3.docx", fh.read())

        url = reverse("hx_project_file_upload", args=[self.projekt.pk])
        resp = self.client.post(
            url,
            {"anlage_nr": "", "upload": upload},
            format="multipart",
            HTTP_HX_REQUEST="true",
        )
        self.assertEqual(resp.status_code, 200)
        pf = self.projekt.anlagen.get()
        self.assertEqual(pf.anlage_nr, 3)


@pytest.mark.usefixtures("prepared_files")
class DropzoneUploadTests(NoesisTestCase):
    """Tests für den neuen Datei-Upload-Workflow."""

    def setUp(self):
        self.user = User.objects.create_user("dz", password="pass")
        self.client.login(username="dz", password="pass")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")

    def test_number_from_filename(self):
        with open(self.docx_content_path, "rb") as fh:
            upload = SimpleUploadedFile("Anlage_2.docx", fh.read())

        url = reverse("projekt_file_upload", args=[self.projekt.pk])
        resp = self.client.post(url, {"upload": upload}, format="multipart", HTTP_HX_REQUEST="true")
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.headers.get("X-Upload-Status"), "assigned")
        self.assertTrue(self.projekt.anlagen.filter(anlage_nr=2).exists())

    def test_manual_assignment_flow(self):
        with open(self.docx_content_path, "rb") as fh:
            upload = SimpleUploadedFile("foo.docx", fh.read())

        url = reverse("projekt_file_upload", args=[self.projekt.pk])
        resp = self.client.post(url, {"upload": upload}, format="multipart", HTTP_HX_REQUEST="true")
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.headers.get("X-Upload-Status"), "manual")
        self.assertIn("form", resp.content.decode())
        session = self.client.session
        temp_id = next(iter(session.get("pending_uploads", {})))

        resp2 = self.client.post(
            url,
            {"temp_id": temp_id, "anlage_nr": 3},
            format="multipart",
            HTTP_HX_REQUEST="true",
        )
        self.assertEqual(resp2.status_code, 200)
        self.assertEqual(resp2.headers.get("X-Upload-Status"), "assigned")
        self.assertTrue(self.projekt.anlagen.filter(anlage_nr=3).exists())


class AutoApprovalTests(NoesisTestCase):
    """Tests für die automatische Genehmigung von Dokumenten."""

    def setUp(self) -> None:
        self.user = User.objects.create_user("auto", password="pass")
        self.client.login(username="auto", password="pass")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")

    def _upload_doc(self, document: Document) -> BVProjectFile:
        """Hilfsfunktion zum Hochladen eines DOCX-Dokuments."""
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        document.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("Anlage_1.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)
        url = reverse("hx_project_file_upload", args=[self.projekt.pk])
        resp = self.client.post(
            url,
            {"anlage_nr": 1, "upload": upload, "manual_comment": ""},
            format="multipart",
            HTTP_HX_REQUEST="true",
        )
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.content.decode().count("<tr"), 1)
        return self.projekt.anlagen.get(anlage_nr=1)

    def test_single_page_auto_approved(self):
        doc = Document()
        doc.add_paragraph("Seite 1")
        pf = self._upload_doc(doc)
        self.assertFalse(pf.manual_reviewed)
        self.assertFalse(pf.verhandlungsfaehig)

    def test_multi_page_requires_manual_review(self):
        img = Image.new("RGB", (10, 10), color="red")
        img_tmp = NamedTemporaryFile(delete=False, suffix=".png")
        img.save(img_tmp.name)
        img_tmp.close()

        doc = Document()
        doc.add_paragraph("Seite 1")
        doc.add_page_break()
        doc.add_paragraph("Seite 2")
        doc.add_picture(img_tmp.name)
        Path(img_tmp.name).unlink(missing_ok=True)

        pf = self._upload_doc(doc)
        self.assertFalse(pf.manual_reviewed)
        self.assertFalse(pf.verhandlungsfaehig)

    def test_toggle_manual_review_sets_flag(self):
        doc = Document()
        doc.add_paragraph("Seite 1")
        doc.add_page_break()
        doc.add_paragraph("Seite 2")
        pf = self._upload_doc(doc)

        url = reverse("project_file_toggle_flag", args=[pf.pk, "manual_reviewed"])
        resp = self.client.post(url, {"value": "1"})
        self.assertEqual(resp.status_code, 302)
        pf.refresh_from_db()
        self.assertTrue(pf.manual_reviewed)
        self.assertFalse(pf.verhandlungsfaehig)


class Anlage3AutomationTests(NoesisTestCase):
    def setUp(self) -> None:
        self.user = User.objects.create_user("auto3", password="pass")
        self.client.login(username="auto3", password="pass")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")

    def _upload_docx(self, document: Document) -> BVProjectFile:
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        document.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("Anlage_1.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)
        url = reverse("hx_project_file_upload", args=[self.projekt.pk])
        resp = self.client.post(
            url,
            {"anlage_nr": 3, "upload": upload, "manual_comment": ""},
            format="multipart",
            HTTP_HX_REQUEST="true",
        )
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.content.decode().count("<tr"), 1)
        return BVProjectFile.objects.get(project=self.projekt, anlage_nr=3)

    def test_single_page_sets_negotiable(self):
        doc = Document()
        doc.add_paragraph("Seite 1")
        pf = self._upload_docx(doc)
        self.assertTrue(pf.verhandlungsfaehig)

    def test_review_save_marks_checked(self):
        pf = BVProjectFile.objects.create(
            project=self.projekt,
            anlage_nr=3,
            upload=SimpleUploadedFile("a.docx", b""),
            text_content="",
        )
        url = reverse("projekt_file_edit_json", args=[pf.pk])
        resp = self.client.post(url, {"analysis_json": "{}"})
        self.assertRedirects(resp, reverse("projekt_detail", args=[self.projekt.pk]))
        pf.refresh_from_db()
        self.assertTrue(pf.manual_reviewed)


class QuestionReviewPropagationTests(NoesisTestCase):
    """Tests zur Übernahme von Fragenbewertungen bei neuen Versionen."""

    def setUp(self) -> None:
        self.user = User.objects.create_user("review", password="pass")
        self.client.login(username="review", password="pass")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")

    def test_answers_unchanged_keep_ok_flag(self) -> None:
        BVProjectFile.objects.create(
            project=self.projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("old.docx", b"d"),
            analysis_json={"questions": {"1": {"answer": "a"}, "2": {"answer": "b"}}},
            question_review={
                "1": {"hinweis": "H1", "vorschlag": "V1", "ok": True},
                "2": {"hinweis": "H2", "vorschlag": "V2", "ok": True},
            },
        )
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        Document().save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload_new = SimpleUploadedFile("new.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)
        form = BVProjectFileForm(
            data={"manual_comment": ""},
            files={"upload": upload_new},
            anlage_nr=1,
            instance=BVProjectFile(
                analysis_json={
                    "questions": {"1": {"answer": "a"}, "2": {"answer": "c"}}
                }
            ),
        )
        assert form.is_valid()
        new_pf = _save_project_file(self.projekt, form=form)

        review = new_pf.question_review
        self.assertTrue(review["1"]["ok"])
        self.assertFalse(review["2"]["ok"])
        self.assertEqual(review["1"]["hinweis"], "H1")
        self.assertEqual(review["2"]["vorschlag"], "V2")

