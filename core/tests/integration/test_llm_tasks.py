import json
import re
from io import BytesIO
from pathlib import Path
from tempfile import NamedTemporaryFile

import fitz
import pytest
from docx import Document
from django.core.files.uploadedfile import SimpleUploadedFile
from unittest.mock import patch

from ...models import (
    BVProject,
    BVProjectFile,
    Anlage2Function,
    Anlage2Config,
    AnlagenFunktionsMetadaten,
    FunktionsErgebnis,
    AntwortErkennungsRegel,
    Anlage2SubQuestion,
    Anlage1Question,
)
from ...parser_manager import parser_manager
from ...parsers import AbstractParser
from ...llm_tasks import (
    check_anlage1,
    check_anlage2,
    analyse_anlage3,
    run_anlage2_analysis,
    run_conditional_anlage2_check,
    worker_verify_feature,
    generate_gutachten,
    parse_anlage1_questions,
)
from ..base import NoesisTestCase

pytestmark = [pytest.mark.integration, pytest.mark.usefixtures("seed_db")]

class LLMTasksTests(NoesisTestCase):
    maxDiff = None

    def setUp(self) -> None:  # pragma: no cover - setup
        super().setUp()
        self.func = Anlage2Function.objects.create(name="Anmelden")

    # test_classify_system entfernt: Feature ausgebaut

    def test_check_anlage2(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content="Anlagetext",
        )
        func = self.func
        data = check_anlage2(projekt.pk)
        file_obj = projekt.anlagen.get(anlage_nr=2)
        self.assertEqual(data["functions"][0]["source"], "parser")
        res = AnlagenFunktionsMetadaten.objects.get(anlage_datei=pf, funktion=func)
        fe = FunktionsErgebnis.objects.filter(
            anlage_datei=pf, funktion=func, quelle="parser"
        ).first()
        self.assertIsNotNone(fe)
        self.assertIsNone(fe.technisch_verfuegbar)

    def test_check_anlage2_functions_stores_result(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"data"),
        )
        func = self.func
        llm_reply = json.dumps({"technisch_verfuegbar": True})
        with (
            patch("core.llm_tasks.query_llm", return_value=llm_reply),
            patch("core.llm_tasks.async_task") as mock_async,
            patch("core.llm_tasks.result") as mock_result,
        ):
            mock_async.side_effect = lambda name, *a, **k: (
                worker_verify_feature(*a, **k) or "tid"
            )
            mock_result.side_effect = lambda *a, **k: None
            run_conditional_anlage2_check(pf.pk)

        res = AnlagenFunktionsMetadaten.objects.get(anlage_datei=pf, funktion=func)
        fe = FunktionsErgebnis.objects.filter(
            anlage_datei=pf, funktion=func, quelle="ki"
        ).first()
        self.assertIsNotNone(fe)
        self.assertTrue(fe.technisch_verfuegbar)

    def x_test_check_anlage2_llm_receives_text(self):
        """Der LLM-Prompt enthält den bekannten Text."""
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content="Testinhalt Anlage2",
        )
        func = self.func
        llm_reply = json.dumps({"technisch_verfuegbar": False})
        with patch("core.llm_tasks.query_llm", return_value=llm_reply) as mock_q:
            data = check_anlage2(projekt.pk)
        self.assertIn("Testinhalt Anlage2", mock_q.call_args_list[0].args[0].text)
        file_obj = projekt.anlagen.get(anlage_nr=2)
        self.assertTrue(
            any(f["funktion"] == "Anmelden" for f in data["functions"])
        )

    def x_test_check_anlage2_prompt_contains_text(self):
        """Der Prompt enth\u00e4lt den gesamten Anlagentext."""
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content="Testinhalt Anlage2",
        )
        func = self.func
        llm_reply = json.dumps({"technisch_verfuegbar": False})
        with patch("core.llm_tasks.query_llm", return_value=llm_reply) as mock_q:
            data = check_anlage2(projekt.pk)
        prompt = mock_q.call_args_list[0].args[0].text
        self.assertIn("Testinhalt Anlage2", prompt)
        file_obj = projekt.anlagen.get(anlage_nr=2)
        self.assertTrue(
            any(f["funktion"] == "Anmelden" for f in data["functions"])
        )

    def test_check_anlage2_parser(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("b.txt", b"x"),
            text_content="Anmelden tv ja ki ja",
        )
        cfg = Anlage2Config.get_instance()
        cfg.text_technisch_verfuegbar_true = ["ja"]
        cfg.text_ki_beteiligung_true = ["ja"]
        cfg.save()
        AntwortErkennungsRegel.objects.create(
            regel_name="tv ja",
            erkennungs_phrase="tv ja",
            actions_json=[{"field": "technisch_verfuegbar", "value": True}],
        )
        AntwortErkennungsRegel.objects.create(
            regel_name="ki ja",
            erkennungs_phrase="ki ja",
            actions_json=[{"field": "ki_beteiligung", "value": True}],
        )
        func = self.func

        data = check_anlage2(projekt.pk)
        expected_function = {
            "funktion": "Anmelden",
            "technisch_verfuegbar": {"value": True, "note": None},
            "ki_beteiligung": {"value": True, "note": None},
            "source": "parser",
        }
        pf.refresh_from_db()
        self.assertEqual(data["task"], "check_anlage2")
        self.assertIn(expected_function, data["functions"])
        self.assertIn(expected_function, pf.analysis_json["functions"])

    def test_run_anlage2_analysis_table(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")

        doc = Document()
        table = doc.add_table(rows=2, cols=5)
        table.cell(0, 0).text = "Funktion"
        table.cell(0, 1).text = "Technisch vorhanden"
        table.cell(0, 2).text = "Einsatz bei Telefónica"
        table.cell(0, 3).text = "Zur LV-Kontrolle"
        table.cell(0, 4).text = "KI-Beteiligung"
        table.cell(1, 0).text = "Anmelden"
        table.cell(1, 1).text = "Ja"
        table.cell(1, 2).text = "Nein"
        table.cell(1, 3).text = "Nein"
        table.cell(1, 4).text = "Ja"

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        upload = SimpleUploadedFile("b.docx", buffer.read())

        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=upload,
            text_content="Anmelden: tv: ja; tel: nein; lv: nein; ki: ja",
        )
        func = self.func
        cfg = Anlage2Config.get_instance()
        cfg.parser_mode = "table_only"
        cfg.parser_order = ["table"]
        cfg.text_technisch_verfuegbar_true = ["ja"]
        cfg.text_technisch_verfuegbar_false = []
        cfg.text_einsatz_telefonica_true = []
        cfg.text_einsatz_telefonica_false = ["nein"]
        cfg.text_zur_lv_kontrolle_true = []
        cfg.text_zur_lv_kontrolle_false = ["nein"]
        cfg.text_ki_beteiligung_true = ["ja"]
        cfg.text_ki_beteiligung_false = []
        cfg.save()

        result = run_anlage2_analysis(pf)
        expected = [
            {
                "funktion": "Anmelden",
                "technisch_verfuegbar": {"value": True, "note": None},
                "einsatz_telefonica": {"value": False, "note": None},
                "zur_lv_kontrolle": {"value": False, "note": None},
                "ki_beteiligung": {"value": True, "note": None},
            }
        ]

        pf.refresh_from_db()
        fe = FunktionsErgebnis.objects.filter(
            anlage_datei=pf, funktion=func, quelle="parser"
        ).first()
        self.assertIsNotNone(fe)
        self.assertTrue(fe.technisch_verfuegbar)

        login_entry = next(
            f for f in pf.analysis_json["functions"] if f["funktion"] == "Anmelden"
        )
        self.assertTrue(login_entry["technisch_verfuegbar"]["value"])
        self.assertFalse(login_entry["einsatz_telefonica"]["value"])
        self.assertFalse(login_entry["zur_lv_kontrolle"]["value"])
        self.assertTrue(login_entry["ki_beteiligung"]["value"])

        self.assertIsInstance(result, list)
        self.assertTrue(any(r["funktion"] == "Anmelden" for r in result))

    def test_run_anlage2_analysis_sets_negotiable_on_match(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        content = "Anmelden: tv: ja; tel: nein; lv: nein; ki: ja"
        upload = SimpleUploadedFile("b.txt", b"x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=upload,
            text_content=content,
        )
        func = self.func
        cfg = Anlage2Config.get_instance()
        cfg.text_technisch_verfuegbar_true = ["ja"]
        cfg.save()
        AnlagenFunktionsMetadaten.objects.create(
            anlage_datei=pf,
            funktion=func,
        )
        FunktionsErgebnis.objects.create(
            anlage_datei=pf,
            funktion=func,
            quelle="ki",
            technisch_verfuegbar=True,
        )

        run_anlage2_analysis(pf)

        parser_fe = FunktionsErgebnis.objects.filter(
            anlage_datei=pf, funktion=func, quelle="parser"
        ).first()
        ai_fe = FunktionsErgebnis.objects.filter(
            anlage_datei=pf, funktion=func, quelle="ki"
        ).first()

        self.assertTrue(parser_fe.technisch_verfuegbar)
        self.assertTrue(ai_fe.technisch_verfuegbar)

    def test_parser_manager_no_fallback_on_error(self):
        class FailParser(AbstractParser):
            name = "fail"

            def parse(self, project_file):
                raise ValueError("boom")

        class DummyParser(AbstractParser):
            name = "dummy"

            def parse(self, project_file):
                return [{"funktion": "Dummy"}]

        parser_manager.register(FailParser)
        parser_manager.register(DummyParser)
        cfg = Anlage2Config.get_instance()
        old_order, old_mode = cfg.parser_order, cfg.parser_mode
        cfg.parser_order = ["fail", "dummy"]
        cfg.save()

        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        doc = Document()
        doc.add_table(rows=1, cols=1)
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("c.docx", fh.read())
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=upload,
        )

        try:
            with patch.object(DummyParser, "parse", wraps=DummyParser.parse) as m_dummy:
                result = parser_manager.parse_anlage2(pf)
        finally:
            Path(tmp.name).unlink(missing_ok=True)
            parser_manager._parsers.pop("fail")
            parser_manager._parsers.pop("dummy")
            cfg.parser_order = old_order
            cfg.parser_mode = old_mode
            cfg.save()

        self.assertEqual(result, [])
        m_dummy.assert_not_called()

    def test_parser_manager_order(self):
        class P1(AbstractParser):
            name = "one"

            def parse(self, project_file):
                return [{"val": 1}]

        class P2(AbstractParser):
            name = "two"

            def parse(self, project_file):
                return [{"val": 2}]

        parser_manager.register(P1)
        parser_manager.register(P2)
        cfg = Anlage2Config.get_instance()
        old_order, old_mode = cfg.parser_order, cfg.parser_mode
        cfg.parser_order = ["two", "one"]
        cfg.save()

        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("d.docx", fh.read())
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=upload,
        )
        try:
            result = parser_manager.parse_anlage2(pf)
        finally:
            Path(tmp.name).unlink(missing_ok=True)
            parser_manager._parsers.pop("one")
            parser_manager._parsers.pop("two")
            cfg.parser_order = old_order
            cfg.parser_mode = old_mode
            cfg.save()

        self.assertEqual(result, [{"val": 2}])

    def test_parser_manager_uses_first_result(self):
        class P1(AbstractParser):
            name = "p1"

            def parse(self, project_file):
                return [{"funktion": "A", "technisch_verfuegbar": {"value": False}}]

        class P2(AbstractParser):
            name = "p2"

            def parse(self, project_file):
                return [{"funktion": "A", "technisch_verfuegbar": {"value": True}}]

        parser_manager.register(P1)
        parser_manager.register(P2)
        cfg = Anlage2Config.get_instance()
        old_order, old_mode = cfg.parser_order, cfg.parser_mode
        cfg.parser_order = ["p1", "p2"]
        cfg.save()

        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        doc = Document()
        doc.add_table(rows=1, cols=1)
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("e.docx", fh.read())
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=upload,
        )

        try:
            with (
                patch("core.parsers.parse_anlage2_table", return_value=[]) as m_tab,
                patch("core.text_parser.parse_anlage2_text", return_value=[]) as m_text,
            ):
                result = parser_manager.parse_anlage2(pf)
        finally:
            Path(tmp.name).unlink(missing_ok=True)
            parser_manager._parsers.pop("p1", None)
            parser_manager._parsers.pop("p2", None)
            cfg.parser_order = old_order
            cfg.parser_mode = old_mode
            cfg.save()

        m_tab.assert_not_called()
        m_text.assert_not_called()
        self.assertFalse(result[0]["technisch_verfuegbar"]["value"])

    def test_parser_manager_uses_exact_by_default(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"x"),
            text_content="t",
        )
        with (
            patch("core.parsers.ExactParser.parse", return_value=[{"funktion": "Alt"}]) as m_exact,
            patch("core.parsers.TableParser.parse") as m_table,
        ):
            result = parser_manager.parse_anlage2(pf)
        m_exact.assert_called_once()
        m_table.assert_not_called()
        self.assertEqual(result, [{"funktion": "Alt"}])

    def test_parser_manager_exact_parser_segments(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"x"),
            text_content="Alpha: aktiv\nBeta: kein einsatz",
        )
        Anlage2Function.objects.create(name="Alpha")
        Anlage2Function.objects.create(name="Beta")
        AntwortErkennungsRegel.objects.create(
            regel_name="aktiv",
            erkennungs_phrase="aktiv",
            actions_json=[{"field": "technisch_verfuegbar", "value": True}],
        )
        AntwortErkennungsRegel.objects.create(
            regel_name="einsatz",
            erkennungs_phrase="kein einsatz",
            actions_json=[{"field": "einsatz_telefonica", "value": False}],
        )
        cfg = Anlage2Config.get_instance()
        old_order, old_mode = cfg.parser_order, cfg.parser_mode
        cfg.parser_mode = "exact_only"
        cfg.save()
        result = parser_manager.parse_anlage2(pf)
        cfg.parser_order = old_order
        cfg.parser_mode = old_mode
        cfg.save()
        self.assertEqual(
            result,
            [
                {
                    "funktion": "Alpha",
                    "technisch_verfuegbar": {"value": True, "note": None},
                },
                {
                    "funktion": "Beta",
                    "einsatz_telefonica": {"value": False, "note": None},
                },
            ],
        )

    def test_run_anlage2_analysis_includes_missing_functions(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"x"),
            text_content="",
        )
        func = self.func

        result = run_anlage2_analysis(pf)

        entry = next(r for r in result if r["funktion"] == "Anmelden")
        self.assertTrue(
            entry.get("not_found") or entry.get("technisch_verfuegbar") is None
        )
        pf.refresh_from_db()
        fe = FunktionsErgebnis.objects.filter(
            anlage_datei=pf, funktion=func, quelle="parser"
        ).first()
        self.assertIsNotNone(fe)
        self.assertIsNone(fe.technisch_verfuegbar)

    def test_run_anlage2_analysis_includes_missing_subquestions(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"x"),
            text_content="",
        )
        func = self.func
        Anlage2SubQuestion.objects.filter(funktion=func).delete()
        Anlage2SubQuestion.objects.create(funktion=func, frage_text="Warum?")

        result = run_anlage2_analysis(pf)

        names = [row["funktion"] for row in result]
        self.assertIn("Anmelden", names)
        self.assertTrue(any("Warum?" in n for n in names))
        pf.refresh_from_db()
        parser_res = FunktionsErgebnis.objects.filter(
            anlage_datei=pf, funktion=func, quelle="parser"
        )
        self.assertEqual(parser_res.count(), 2)

    def test_run_anlage2_analysis_sets_complete_status_without_followup(self):
        """Status wird nur ohne anschließende KI-Prüfung auf COMPLETE gesetzt."""

        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"x"),
            text_content="",
            processing_status=BVProjectFile.PROCESSING,
        )
        func = self.func
        FunktionsErgebnis.objects.create(
            anlage_datei=pf,
            funktion=func,
            quelle="ki",
            technisch_verfuegbar=True,
        )

        run_anlage2_analysis(pf)

        pf.refresh_from_db()
        self.assertEqual(pf.processing_status, BVProjectFile.COMPLETE)

    def test_run_anlage2_analysis_keeps_processing_with_followup(self):
        """Bei ausstehender KI-Prüfung bleibt der Status PROCESSING."""

        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"x"),
            text_content="",
            processing_status=BVProjectFile.PROCESSING,
        )
        _ = self.func

        run_anlage2_analysis(pf)

        pf.refresh_from_db()
        self.assertEqual(pf.processing_status, BVProjectFile.PROCESSING)

    def test_check_anlage2_table_error_fallback(self):
        class P1(AbstractParser):
            name = "p1"

            def parse(self, project_file):
                return [{"funktion": "A", "technisch_verfuegbar": {"value": False}}]

        class P2(AbstractParser):
            name = "p2"

            def parse(self, project_file):
                return [{"funktion": "A", "technisch_verfuegbar": {"value": True}}]

        parser_manager.register(P1)
        parser_manager.register(P2)
        cfg = Anlage2Config.get_instance()
        cfg.parser_order = ["p1", "p2"]
        cfg.save()

        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pf = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("b.txt", b"x"),
        )

        try:
            result = parser_manager.parse_anlage2(pf)
        finally:
            parser_manager._parsers.pop("p1", None)
            parser_manager._parsers.pop("p2", None)
            cfg.parser_order = ["table"]
            cfg.save()

        self.assertFalse(result[0]["technisch_verfuegbar"]["value"])

    def test_analyse_anlage3_auto_ok(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        doc = Document()
        doc.add_paragraph("Seite 1")
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("c.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)
        BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=3,
            upload=upload,
            text_content="ignored",
        )

        pf = projekt.anlagen.get(anlage_nr=3)
        data = analyse_anlage3(pf.pk)
        pf.refresh_from_db()
        file_obj = pf
        self.assertEqual(data["pages"], 1)

        self.assertTrue(data["auto_ok"])
        self.assertTrue(file_obj.analysis_json["auto_ok"])

        if hasattr(file_obj, "verhandlungsfaehig"):
            self.assertTrue(file_obj.verhandlungsfaehig)

    def test_analyse_anlage3_manual_required(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        doc = Document()
        doc.add_paragraph("Seite 1")
        doc.add_page_break()
        doc.add_paragraph("Seite 2")
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("d.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)
        BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=3,
            upload=upload,
            text_content="ignored",
        )

        pf = projekt.anlagen.get(anlage_nr=3)
        data = analyse_anlage3(pf.pk)
        pf.refresh_from_db()
        file_obj = pf
        self.assertEqual(data["pages"], 2)

        self.assertTrue(data["manual_required"])
        self.assertTrue(file_obj.analysis_json["manual_required"])

        if hasattr(file_obj, "verhandlungsfaehig"):
            self.assertFalse(file_obj.verhandlungsfaehig)

    def test_analyse_anlage3_pdf_auto_ok(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pdf = fitz.open()
        pdf.new_page()
        tmp = NamedTemporaryFile(delete=False, suffix=".pdf")
        tmp.close()
        pdf.save(tmp.name)
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("c.pdf", fh.read())
        Path(tmp.name).unlink(missing_ok=True)
        BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=3,
            upload=upload,
            text_content="ignored",
        )

        pf = projekt.anlagen.get(anlage_nr=3)
        data = analyse_anlage3(pf.pk)
        pf.refresh_from_db()
        file_obj = pf
        self.assertEqual(data["pages"], 1)

        self.assertTrue(data["auto_ok"])
        self.assertTrue(file_obj.analysis_json["auto_ok"])


    def test_analyse_anlage3_pdf_manual_required(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        pdf = fitz.open()
        pdf.new_page()
        pdf.new_page()
        tmp = NamedTemporaryFile(delete=False, suffix=".pdf")
        tmp.close()
        pdf.save(tmp.name)
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("d.pdf", fh.read())
        Path(tmp.name).unlink(missing_ok=True)
        BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=3,
            upload=upload,
            text_content="ignored",
        )

        pf = projekt.anlagen.get(anlage_nr=3)
        data = analyse_anlage3(pf.pk)
        pf.refresh_from_db()
        file_obj = pf
        self.assertEqual(data["pages"], 2)

        self.assertTrue(data["manual_required"])
        self.assertTrue(file_obj.analysis_json["manual_required"])

    def test_analyse_anlage3_multiple_files(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")

        doc1 = Document()
        doc1.add_paragraph("Seite 1")
        tmp1 = NamedTemporaryFile(delete=False, suffix=".docx")
        doc1.save(tmp1.name)
        tmp1.close()
        with open(tmp1.name, "rb") as fh:
            upload1 = SimpleUploadedFile("e.docx", fh.read())
        Path(tmp1.name).unlink(missing_ok=True)
        pf1 = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=3,
            upload=upload1,
            text_content="x",
        )

        doc2 = Document()
        doc2.add_paragraph("Seite 1")
        doc2.add_page_break()
        doc2.add_paragraph("Seite 2")
        tmp2 = NamedTemporaryFile(delete=False, suffix=".docx")
        doc2.save(tmp2.name)
        tmp2.close()
        with open(tmp2.name, "rb") as fh:
            upload2 = SimpleUploadedFile("f.docx", fh.read())
        Path(tmp2.name).unlink(missing_ok=True)
        pf2 = BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=3,
            upload=upload2,
            text_content="y",
        )

        analyse_anlage3(pf1.pk)
        pf1.refresh_from_db()
        pf2.refresh_from_db()
        self.assertIsNotNone(pf1.analysis_json)
        self.assertIsNotNone(pf2.analysis_json)

    def test_check_anlage1_parser(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        q1_text = Anlage1Question.objects.get(num=1).text
        q2_text = Anlage1Question.objects.get(num=2).text
        text = f"{q1_text}\u00b6A1\u00b6{q2_text}\u00b6A2"
        BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content=text,
        )
        file_obj = projekt.anlagen.get(anlage_nr=1)
        data = check_anlage1(file_obj.pk)
        file_obj.refresh_from_db()
        expected = {"questions": parse_anlage1_questions(text)}
        file_obj.refresh_from_db()
        self.assertEqual(data, expected)
        self.assertEqual(file_obj.analysis_json, expected)

    def test_parse_anlage1_questions_extra(self):
        Anlage1Question.objects.create(
            num=10,
            text="Frage 10: Test?",
            enabled=True,
            parser_enabled=True,
            llm_enabled=True,
        )
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        q1_text = Anlage1Question.objects.get(num=1).text
        text = f"{q1_text}\u00b6A1\u00b6Frage 10: Test?\u00b6A10"
        BVProjectFile.objects.create(
            project=projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content=text,
        )
        file_obj = projekt.anlagen.get(anlage_nr=1)
        data = check_anlage1(file_obj.pk)
        q_data = data["questions"]
        self.assertEqual(q_data["10"]["answer"], "A10")

    def test_parse_anlage1_questions_without_numbers(self):
        """Prüft die Extraktion ohne nummerierte Fragen."""
        # Frage-Texte ohne Präfix "Frage X:" speichern
        q1 = Anlage1Question.objects.get(num=1)
        q2 = Anlage1Question.objects.get(num=2)
        prefix = r"^Frage\s+\d+(?:\.\d+)?[:.]?\s*"
        q1.text = re.sub(prefix, "", q1.text)
        q2.text = re.sub(prefix, "", q2.text)
        q1.save(update_fields=["text"])
        q2.save(update_fields=["text"])
        v1 = q1.variants.first()
        v2 = q2.variants.first()
        v1.text = q1.text
        v2.text = q2.text
        v1.save()
        v2.save()

        text = f"{q1.text}\u00b6A1\u00b6{q2.text}\u00b6A2"
        parsed = parse_anlage1_questions(text)
        self.assertEqual(
            parsed,
            {
                "1": {"answer": "A1", "found_num": None},
                "2": {"answer": "A2", "found_num": None},
            },
        )

    def test_parse_anlage1_questions_with_variant(self):
        q1 = Anlage1Question.objects.get(num=1)
        q1.variants.create(text="Alternative Frage 1?")
        text = "Alternative Frage 1?\u00b6A1"
        parsed = parse_anlage1_questions(text)
        self.assertEqual(parsed, {"1": {"answer": "A1", "found_num": "1"}})

    def test_parse_anlage1_questions_with_newlines(self):
        """Extraktion funktioniert trotz Zeilenumbr\u00fcche."""
        q1_text = Anlage1Question.objects.get(num=1).text
        q2_text = Anlage1Question.objects.get(num=2).text
        text = f"{q1_text}\nA1\n{q2_text}\nA2"
        parsed = parse_anlage1_questions(text)
        self.assertEqual(
            parsed,
            {
                "1": {"answer": "A1", "found_num": None},
                "2": {"answer": "A2", "found_num": None},
            },
        )

    def test_parse_anlage1_questions_split_lines(self):
        """Fragen werden auch mit Zeilenumbrüchen innerhalb des Textes erkannt."""
        q1_text = Anlage1Question.objects.get(num=1).text
        q2_text = Anlage1Question.objects.get(num=2).text
        q1_split = q1_text.replace(" ", "\n", 1)
        q2_split = q2_text.replace(" ", "\n", 1)
        text = f"{q1_split}\r\nA1\n{q2_split}\r\nA2"
        parsed = parse_anlage1_questions(text)
        self.assertEqual(
            parsed,
            {
                "1": {"answer": "A1", "found_num": None},
                "2": {"answer": "A2", "found_num": None},
            },
        )

    def test_parse_anlage1_questions_respects_parser_enabled(self):
        q2 = Anlage1Question.objects.get(num=2)
        q2.parser_enabled = False
        q2.save(update_fields=["parser_enabled"])
        q1_text = Anlage1Question.objects.get(num=1).text
        text = f"{q1_text}\u00b6A1"
        parsed = parse_anlage1_questions(text)
        self.assertEqual(parsed, {"1": {"answer": "A1", "found_num": None}})

    def test_parse_anlage1_questions_returns_empty_dict(self):
        """Bei fehlenden Treffern wird ein leeres Dict zur\u00fcckgegeben."""
        text = "Es gibt hier keine Fragen."
        parsed = parse_anlage1_questions(text)
        self.assertEqual(parsed, {})

    def test_generate_gutachten_twice_replaces_file(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        first = generate_gutachten(projekt.pk, text="Alt")
        second = generate_gutachten(projekt.pk, text="Neu")
        try:
            self.assertTrue(second.exists())
            self.assertNotEqual(first, second)
            self.assertFalse(first.exists())
        finally:
            second.unlink(missing_ok=True)

    # Hinweis: Tests für die ehemalige Hilfsfunktion `_parse_anlage2` wurden entfernt,
    # da die Funktion bewusst aus `core.llm_tasks` entfernt wurde. Die neue Parser-
    # Pipeline wird an anderen Stellen in diesem Modul bereits ausführlich getestet
    # (siehe Verwendungen von `parser_manager.parse_anlage2`).


