from __future__ import annotations

from pathlib import Path
import json
import uuid

from django.conf import settings
from docx import Document

from django.db.models import Q
from .models import BVProject, AnlagenFunktionsMetadaten


def _get_value(obj):
    """Extrahiert den Wert aus Strukturen mit ``{"value": x}``."""
    if isinstance(obj, dict) and "value" in obj:
        return obj["value"]
    return obj


def _add_json_section(doc: Document, title: str, data: dict | list | str) -> None:
    """F\u00fcgt einen Abschnitt mit JSON-Daten hinzu."""
    doc.add_heading(title, level=2)
    if isinstance(data, (dict, list)):
        text = json.dumps(data, indent=2, ensure_ascii=False)
    else:
        text = str(data)
    for line in text.splitlines():
        doc.add_paragraph(line)


def _output_path(prefix: str) -> Path:
    out_dir = Path(settings.MEDIA_ROOT) / "reports"
    out_dir.mkdir(parents=True, exist_ok=True)
    fname = f"{prefix}_{uuid.uuid4().hex}.docx"
    return out_dir / fname


def generate_gap_analysis(project: BVProject) -> Path:
    """Erzeugt eine Gap-Analyse f\u00fcr das angegebene Projekt."""
    doc = Document()
    doc.add_heading("Gap-Analyse", level=1)

    if project.classification_json:
        _add_json_section(doc, "Klassifizierung", project.classification_json)

    for anlage in project.anlagen.all():
        if anlage.analysis_json:
            _add_json_section(doc, f"Anlage {anlage.anlage_nr}", anlage.analysis_json)
        notes = []
        if anlage.gap_notiz:
            notes.append(("Intern", anlage.gap_notiz))
        if anlage.gap_summary:
            notes.append(("Extern", anlage.gap_summary))
        if notes:
            doc.add_heading(f"GAP-Notizen Anlage {anlage.anlage_nr}", level=2)
            for label, text in notes:
                doc.add_heading(label, level=3)
                for line in text.splitlines():
                    doc.add_paragraph(line)
        if anlage.anlage_nr == 2:
            results = AnlagenFunktionsMetadaten.objects.filter(
                anlage_datei=anlage
            ).filter(
                Q(gap_summary__isnull=False) & ~Q(gap_summary="")
                | Q(gap_notiz__isnull=False) & ~Q(gap_notiz="")
            ).select_related("funktion", "subquestion")
            if results:
                doc.add_heading(
                    f"Detailnotizen Anlage {anlage.anlage_nr}", level=2
                )
                for r in results:
                    title = r.funktion.name
                    if r.subquestion:
                        title += f" - {r.subquestion.frage_text}"
                    doc.add_heading(title, level=3)
                    if r.gap_notiz:
                        doc.add_paragraph(f"Intern: {r.gap_notiz}")
                    if r.gap_summary:
                        doc.add_paragraph(f"Extern: {r.gap_summary}")

    path = _output_path("gap")
    doc.save(path)
    return path


def generate_management_summary(project: BVProject) -> Path:
    """Erstellt eine Management-Zusammenfassung aus den Analyseergebnissen."""
    doc = Document()
    doc.add_heading("Management Summary", level=1)

    if project.classification_json:
        data = project.classification_json
        cat = _get_value(data.get("kategorie")) if isinstance(data, dict) else None
        begr = _get_value(data.get("begruendung")) if isinstance(data, dict) else None
        doc.add_heading("Klassifizierung", level=2)
        if cat:
            doc.add_paragraph(f"Kategorie: {cat}")
        if begr:
            doc.add_paragraph(f"Begr\u00fcndung: {begr}")
        if not cat and not begr:
            for line in json.dumps(data, indent=2, ensure_ascii=False).splitlines():
                doc.add_paragraph(line)

    for anlage in project.anlagen.all():
        doc.add_heading(f"Anlage {anlage.anlage_nr}", level=2)
        if anlage.manual_comment:
            for line in anlage.manual_comment.splitlines():
                doc.add_paragraph(line)
        if anlage.analysis_json:
            _add_json_section(doc, "Analyse", anlage.analysis_json)

    path = _output_path("summary")
    doc.save(path)
    return path
