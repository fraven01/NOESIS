import re
import string
import logging
from pathlib import Path
from typing import List

from docx import Document

from .models import BVProjectFile, Anlage4Config, Anlage4ParserConfig


logger = logging.getLogger("anlage4_debug")


def _normalize(text: str) -> str:
    """Normiert Text für robuste Vergleiche."""
    text = text.lower()
    text = text.translate(str.maketrans("", "", string.punctuation))
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _phrase_pattern(phrase: str) -> str:
    """Erzeugt ein Regex-Fragment aus einer Phrase."""
    words = _normalize(phrase).split()
    if not words:
        return ""
    return r"\W*".join(re.escape(w) for w in words) + r"\W*"


def _log_phrase_presence(desc: str, phrase: str, block: str) -> None:
    """Gibt Debug-Informationen über das Auftreten einer Phrase aus."""
    logger.debug("--- START DEBUGGING BLOCK ---")
    logger.debug("KONFIGURATION [%s]: '%s'", desc, phrase)
    logger.debug("DOKUMENTEN-BLOCK (ersten 150 Zeichen): '%s...'", block[:150])
    if phrase in block:
        logger.debug("SUCCESS: '%s' wurde im Block gefunden.", desc)
    else:
        logger.debug("FAILURE: '%s' NICHT im Block gefunden.", desc)
    logger.debug("--- END DEBUGGING BLOCK ---")


def parse_anlage4(
    project_file: BVProjectFile, cfg: Anlage4Config | None = None
) -> List[str]:
    """Parst Anlage 4 anhand der Konfiguration."""
    logger.info("parse_anlage4 gestartet für Datei %s", project_file.pk)
    if cfg is None:
        cfg = project_file.anlage4_config or Anlage4Config.objects.first()
    columns = [_normalize(c) for c in (cfg.table_columns if cfg else [])]
    neg_patterns = [re.compile(p, re.I) for p in (cfg.negative_patterns if cfg else [])]
    patterns = [re.compile(p, re.I) for p in (cfg.regex_patterns if cfg else [])]

    logger.debug(
        "Konfiguration: columns=%s, regex_patterns=%s, negative_patterns=%s",
        columns,
        [p.pattern for p in patterns],
        [p.pattern for p in neg_patterns],
    )

    text = project_file.text_content or ""
    logger.debug("Rohtext der Anlage4 (%s Zeichen): %s", len(text), text)

    for pat in neg_patterns:
        logger.debug("Pr\u00fcfe negatives Muster '%s'", pat.pattern)
        match = pat.search(text)
        if match:
            logger.error(
                "Negative pattern erkannt: %s -> %r", pat.pattern, match.group(0)
            )
            return []

    items: List[str] = []
    structure: str | None = None

    path = Path(project_file.upload.path)
    if path.exists() and path.suffix.lower() == ".docx":
        try:
            doc = Document(str(path))
            for table in doc.tables:
                logger.debug("Prüfe Tabelle mit %s Zeilen", len(table.rows))
                found = False
                for row in table.rows:
                    if len(row.cells) < 2:
                        continue
                    key = _normalize(row.cells[0].text.strip())
                    if key in columns:
                        value = row.cells[1].text.strip()
                        if value:
                            if not found:
                                structure = "table detected"
                                found = True
                            logger.debug("Gefundenes Paar %s: %s", key, value)
                            items.append(value)
                if found and items:
                    logger.debug("%s - %s items", structure, len(items))
                    return items
        except Exception as exc:  # pragma: no cover - ungültige Datei
            logger.error("Anlage4Parser Fehler", exc_info=True)
            structure = structure or ("free text found" if text.strip() else "empty document")
    
    for pat in patterns:
        logger.debug("Suche Muster '%s'", pat.pattern)
        matches = pat.findall(text)
        logger.debug("Gefundene Treffer für '%s': %s", pat.pattern, matches)
        for m in matches:
            logger.debug("Treffer hinzugefügt: %s", m)
            items.append(m)

    if structure is None:
        structure = "free text found" if text.strip() else "empty document"
    logger.debug("%s - %s items", structure, len(items))
    for idx, item in enumerate(items):
        logger.debug("Item %s: %s", idx, item)
    logger.info(
        "parse_anlage4 beendet für Datei %s mit %s Auswertungen",
        project_file.pk,
        len(items),
    )
    return items



def parse_anlage4_dual(
    project_file: BVProjectFile,
    cfg: Anlage4ParserConfig | None = None,
) -> List[dict]:
    """Parst Anlage 4 mit Dual-Strategie.

    Wenn keine Konfiguration vorhanden ist, wird ein leerer
    Ergebnis-Liste zurückgegeben und eine Warnung protokolliert.
    """

    logger.info("parse_anlage4_dual gestartet für Datei %s", project_file.pk)
    if cfg is None:
        cfg = project_file.anlage4_parser_config or Anlage4ParserConfig.objects.first()
    if cfg is None:
        logger.warning("Keine Anlage4ParserConfig vorhanden")
        return []

    columns = [_normalize(c) for c in (cfg.table_columns if cfg else [])]
    delimiter_phrase = cfg.delimiter_phrase if cfg else ""
    ges_phrase = cfg.gesellschaften_phrase if cfg else ""
    fach_phrase = cfg.fachbereiche_phrase if cfg else ""
    name_aliases = cfg.name_aliases if cfg else []
    ges_aliases = cfg.gesellschaft_aliases if cfg else []
    fach_aliases = cfg.fachbereich_aliases if cfg else []
    neg_patterns = [re.compile(p, re.I) for p in (cfg.negative_patterns if cfg else [])]

    path = Path(project_file.upload.path)
    if path.exists() and path.suffix.lower() == ".docx" and columns:
        try:
            doc = Document(str(path))
            for table in doc.tables:
                logger.debug("Dual Parser prüft Tabelle mit %s Zeilen", len(table.rows))
                items: list[dict] = []
                for row in table.rows:
                    if len(row.cells) < 2:
                        continue
                    key = _normalize(row.cells[0].text.strip())
                    if key in columns:
                        value = row.cells[1].text.strip()
                        if value:
                            logger.debug("Dual Parser gefundenes Paar %s: %s", key, value)
                            items.append({"name_der_auswertung": value})
                if items:
                    logger.debug("Tabelle erkannt - %s Items", len(items))
                    return items
        except Exception:  # pragma: no cover - ungültige Datei
            logger.error("Anlage4Parser Dual Fehler", exc_info=True)

    text = project_file.text_content or ""
    logger.debug("Rohtext für Dual-Parser (%s Zeichen): %s", len(text), text)

    for pat in neg_patterns:
        logger.debug("Pr\u00fcfe negatives Muster '%s'", pat.pattern)
        if pat.search(text):
            logger.error("Negative pattern erkannt: %s", pat.pattern)
            return []

    block_aliases: list[str] = []
    if delimiter_phrase:
        block_aliases.append(r"\s*" + delimiter_phrase + r"\W*")
    block_aliases.extend(_phrase_pattern(a) for a in name_aliases)
    if not block_aliases:
        logger.warning("Keine delimiter_phrase definiert")
        return []

    pattern_block = re.compile(r"^(?:" + "|".join(block_aliases) + ")", re.I | re.M)
    matches = list(pattern_block.finditer(text))
    logger.debug("Gefundene Blöcke: %s", len(matches))
    if not matches:
        return []

    segs: list[str] = []
    for idx, match in enumerate(matches):
        start = match.end()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        segs.append(text[start:end])

    if ges_phrase or ges_aliases:
        patterns = ([ges_phrase] if ges_phrase else []) + ges_aliases
        escaped = [r"\s*" + _phrase_pattern(p) for p in patterns]
        reg_ges = re.compile(r"^(?:" + "|".join(escaped) + ")", re.I | re.M)
    else:
        reg_ges = None
    if fach_phrase or fach_aliases:
        patterns = ([fach_phrase] if fach_phrase else []) + fach_aliases
        escaped = [r"\s*" + _phrase_pattern(p) for p in patterns]
        reg_fach = re.compile(r"^(?:" + "|".join(escaped) + ")", re.I | re.M)
    else:
        reg_fach = None

    results: list[dict] = []
    for seg in segs:
        entry = {"name_der_auswertung": "", "gesellschaften": "", "fachbereiche": ""}
        anchors_found: list[tuple[str, int, int]] = []

        _log_phrase_presence("gesellschaften_phrase", cfg.gesellschaften_phrase, seg)

        if reg_ges:
            m = reg_ges.search(seg)
            if m:
                anchors_found.append(("gesellschaften", m.start(), m.end()))

        _log_phrase_presence("fachbereiche_phrase", cfg.fachbereiche_phrase, seg)

        if reg_fach:
            m = reg_fach.search(seg)
            if m:
                anchors_found.append(("fachbereiche", m.start(), m.end()))
        anchors_found.sort(key=lambda x: x[1])

        end_pos = len(seg)
        name = seg[: anchors_found[0][1] if anchors_found else end_pos].strip()
        name = re.sub(r'^[\s:.-]+', '', name)
        entry["name_der_auswertung"] = name
        for idx, (field, start, anchor_end) in enumerate(anchors_found):
            next_start = anchors_found[idx + 1][1] if idx + 1 < len(anchors_found) else end_pos
            entry[field] = seg[anchor_end:next_start].strip()

        results.append(entry)
        logger.debug("Parsed Block: %s", entry)

    logger.info(
        "parse_anlage4_dual beendet für Datei %s mit %s Einträgen",
        project_file.pk,
        len(results),
    )
    return results
