from django.contrib.auth.models import User, Group
from django.urls import reverse
from django.test import TestCase
from django.http import QueryDict
from django.db import IntegrityError


from django.apps import apps
from .models import (
    BVProject,
    BVProjectFile,
    Recording,
    Prompt,
    LLMConfig,
    Tile,
    UserTileAccess,
    Anlage1Question,
    Anlage1Config,
    Area,
    Anlage2Function,
    Anlage2Config,
    Anlage2ColumnHeading,
    Anlage2SubQuestion,
    Anlage2FunctionResult,
    Anlage2GlobalPhrase,
    SoftwareKnowledge,
    Gutachten,
)
from .docx_utils import (
    extract_text,
    parse_anlage2_table,
    parse_anlage2_text,
    _normalize_header_text,
)
from pathlib import Path
from tempfile import NamedTemporaryFile
from docx import Document

from django.core.files.uploadedfile import SimpleUploadedFile
from .forms import BVProjectForm, BVProjectUploadForm
from .workflow import set_project_status
from .models import ProjectStatus
from .llm_tasks import (
    classify_system,
    check_anlage1,
    check_anlage2,
    analyse_anlage2,
    check_anlage2_functions,
    worker_verify_feature,
    worker_generate_gutachten,
    worker_run_initial_check,
    get_prompt,
    generate_gutachten,
    parse_anlage1_questions,
    _parse_anlage2,
)
from .views import _verification_to_initial
from .reporting import generate_gap_analysis, generate_management_summary
from unittest.mock import patch, ANY
from django.core.management import call_command
from django.test import override_settings
import json


def create_statuses() -> None:
    if ProjectStatus.objects.exists():
        return
    data = [
        ("NEW", "Neu"),
        ("CLASSIFIED", "Klassifiziert"),
        ("GUTACHTEN_OK", "Gutachten OK"),
        ("GUTACHTEN_FREIGEGEBEN", "Gutachten freigegeben"),
        ("IN_PRUEFUNG_ANLAGE_X", "In Prüfung Anlage X"),
        ("FB_IN_PRUEFUNG", "FB in Prüfung"),
        ("ENDGEPRUEFT", "Endgeprüft"),
    ]
    for idx, (key, name) in enumerate(data, start=1):
        ProjectStatus.objects.create(
            name=name,
            key=key,
            ordering=idx,
            is_default=key == "NEW",
            is_done_status=key == "ENDGEPRUEFT",
        )


def setUpModule():
    create_statuses()


class AdminProjectsTests(TestCase):
    def setUp(self):
        admin_group = Group.objects.create(name="admin")
        self.user = User.objects.create_user("admin", password="pass")
        self.user.groups.add(admin_group)
        self.client.login(username="admin", password="pass")

        self.p1 = BVProject.objects.create(software_typen="A", beschreibung="x")
        self.p2 = BVProject.objects.create(software_typen="B", beschreibung="y")

    def test_delete_selected_projects(self):
        url = reverse("admin_projects")
        resp = self.client.post(url, {"delete_selected": "1", "selected_projects": [self.p1.id]})
        self.assertRedirects(resp, url)
        self.assertFalse(BVProject.objects.filter(id=self.p1.id).exists())
        self.assertTrue(BVProject.objects.filter(id=self.p2.id).exists())

    def test_delete_single_project(self):
        url = reverse("admin_projects")
        resp = self.client.post(url, {"delete_single": str(self.p2.id)})
        self.assertRedirects(resp, url)
        self.assertFalse(BVProject.objects.filter(id=self.p2.id).exists())

    def test_delete_single_requires_post(self):
        url = reverse("admin_project_delete", args=[self.p1.id])
        resp = self.client.get(url)
        self.assertEqual(resp.status_code, 405)
        self.assertTrue(BVProject.objects.filter(id=self.p1.id).exists())


class AdminProjectCleanupTests(TestCase):
    def setUp(self):
        admin_group = Group.objects.create(name="admin")
        self.user = User.objects.create_user("admin2", password="pass")
        self.user.groups.add(admin_group)
        self.client.login(username="admin2", password="pass")

        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        self.file = BVProjectFile.objects.create(
            projekt=self.projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content="Text",
        )

    def test_delete_file(self):
        path = Path(self.file.upload.path)
        url = reverse("admin_project_cleanup", args=[self.projekt.pk])
        resp = self.client.post(url, {"action": "delete_file", "file_id": self.file.id})
        self.assertRedirects(resp, url)
        self.assertFalse(BVProjectFile.objects.filter(id=self.file.id).exists())
        self.assertFalse(path.exists())

    def test_delete_gutachten(self):
        gpath = generate_gutachten(self.projekt.pk, text="foo")
        url = reverse("admin_project_cleanup", args=[self.projekt.pk])
        resp = self.client.post(url, {"action": "delete_gutachten"})
        self.assertRedirects(resp, url)
        self.projekt.refresh_from_db()
        self.assertEqual(self.projekt.gutachten_file.name, "")
        self.assertFalse(gpath.exists())

    def test_delete_classification(self):
        self.projekt.classification_json = {"a": 1}
        self.projekt.save()
        url = reverse("admin_project_cleanup", args=[self.projekt.pk])
        resp = self.client.post(url, {"action": "delete_classification"})
        self.assertRedirects(resp, url)
        self.projekt.refresh_from_db()
        self.assertIsNone(self.projekt.classification_json)


class DocxExtractTests(TestCase):
    def test_extract_text(self):
        doc = Document()
        doc.add_paragraph("Das ist ein Test")
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        try:
            text = extract_text(Path(tmp.name))
        finally:
            Path(tmp.name).unlink(missing_ok=True)
        self.assertIn("Das ist ein Test", text)

    def test_normalize_header_text_variants(self):
        cases = {
            "Technisch vorhanden?": "technisch vorhanden",
            "Technisch vorhanden:" : "technisch vorhanden",
            "Technisch   vorhanden": "technisch vorhanden",
            "Technisch\tvorhanden": "technisch vorhanden",
            " Verf\u00fcgbar?\t": "verf\u00fcgbar",
        }
        for raw, expected in cases.items():
            self.assertEqual(_normalize_header_text(raw), expected)

    def test_parse_anlage2_table(self):
        doc = Document()
        table = doc.add_table(rows=2, cols=5)
        table.cell(0, 0).text = "Funktion"
        table.cell(0, 1).text = "Technisch vorhanden"
        table.cell(0, 2).text = "Einsatz bei Telefónica"
        table.cell(0, 3).text = "Zur LV-Kontrolle"
        table.cell(0, 4).text = "KI-Beteiligung"

        table.cell(1, 0).text = "Login"
        table.cell(1, 1).text = "Ja"
        table.cell(1, 2).text = "Nein"
        table.cell(1, 3).text = "Nein"
        table.cell(1, 4).text = "Ja"

        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        try:
            with patch("core.docx_utils.logging.getLogger") as mock_get_logger:
                mock_logger = mock_get_logger.return_value
                data = parse_anlage2_table(Path(tmp.name))
                expected_raw = [
                    "Funktion",
                    "Technisch vorhanden",
                    "Einsatz bei Telefónica",
                    "Zur LV-Kontrolle",
                    "KI-Beteiligung",
                ]
                expected_norm = [
                    "funktion",
                    "technisch vorhanden",
                    "einsatz bei telefónica",
                    "zur lv-kontrolle",
                    "ki-beteiligung",
                ]
                mock_logger.debug.assert_any_call(
                    f"Tabelle 0: Roh-Header = {expected_raw}, Normiert = {expected_norm}"
                )
        finally:
            Path(tmp.name).unlink(missing_ok=True)

        self.assertEqual(
            data,
            [
                {
                    "funktion": "Login",
                    "technisch_verfuegbar": {"value": True, "note": None},
                    "einsatz_telefonica": {"value": False, "note": None},
                    "zur_lv_kontrolle": {"value": False, "note": None},
                    "ki_beteiligung": {"value": True, "note": None},
                }
            ],
        )

    def test_parse_anlage2_table_multiple_headers(self):
        cfg = Anlage2Config.get_instance()
        Anlage2ColumnHeading.objects.create(
            config=cfg, field_name="technisch_vorhanden", text="Verfügbar?"
        )
        Anlage2ColumnHeading.objects.create(
            config=cfg,
            field_name="technisch_vorhanden",
            text="Steht technisch zur Verfügung?",
        )
        Anlage2ColumnHeading.objects.create(
            config=cfg, field_name="einsatz_bei_telefonica", text="Telefonica Einsatz"
        )
        Anlage2ColumnHeading.objects.create(
            config=cfg, field_name="zur_lv_kontrolle", text="LV Kontrolle"
        )
        Anlage2ColumnHeading.objects.create(
            config=cfg, field_name="ki_beteiligung", text="KI?"
        )
        doc = Document()
        table = doc.add_table(rows=2, cols=5)
        table.cell(0, 0).text = "Funktion"
        table.cell(0, 1).text = "Steht technisch zur Verfügung?"
        table.cell(0, 2).text = "Telefonica Einsatz"
        table.cell(0, 3).text = "LV Kontrolle"
        table.cell(0, 4).text = "KI?"
        table.cell(1, 0).text = "Login"
        table.cell(1, 1).text = "Ja"
        table.cell(1, 2).text = "Nein"
        table.cell(1, 3).text = "Nein"
        table.cell(1, 4).text = "Ja"
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        try:
            data = parse_anlage2_table(Path(tmp.name))
        finally:
            Path(tmp.name).unlink(missing_ok=True)

        self.assertTrue(data[0]["technisch_verfuegbar"]["value"])

    def test_parse_anlage2_table_alias_headers(self):
        cfg = Anlage2Config.get_instance()
        Anlage2ColumnHeading.objects.create(
            config=cfg,
            field_name="technisch_vorhanden",
            text="Steht technisch zur Verfügung?",
        )
        Anlage2ColumnHeading.objects.create(
            config=cfg,
            field_name="einsatz_bei_telefonica",
            text="einsatzweise bei telefónica: soll die funktion verwendet werden?",
        )
        Anlage2ColumnHeading.objects.create(
            config=cfg,
            field_name="zur_lv_kontrolle",
            text="einsatzweise bei telefónica: soll zur überwachung von leistung oder verhalten verwendet werden?",
        )

        doc = Document()
        table = doc.add_table(rows=2, cols=5)
        table.cell(0, 0).text = "Funktion"
        table.cell(0, 1).text = "Steht technisch zur Verfügung?"
        table.cell(0, 2).text = "einsatzweise bei Telefónica: soll die Funktion verwendet werden?"
        table.cell(0, 3).text = "einsatzweise bei Telefónica: soll zur Überwachung von Leistung oder Verhalten verwendet werden?"
        table.cell(0, 4).text = "KI-Beteiligung"

        table.cell(1, 0).text = "Login"
        table.cell(1, 1).text = "Ja"
        table.cell(1, 2).text = "Nein"
        table.cell(1, 3).text = "Nein"
        table.cell(1, 4).text = "Ja"

        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        try:
            data = parse_anlage2_table(Path(tmp.name))
        finally:
            Path(tmp.name).unlink(missing_ok=True)

        self.assertEqual(
            data,
            [
                {
                    "funktion": "Login",
                    "technisch_verfuegbar": {"value": True, "note": None},
                    "einsatz_telefonica": {"value": False, "note": None},
                    "zur_lv_kontrolle": {"value": False, "note": None},
                    "ki_beteiligung": {"value": True, "note": None},
                }
            ],
        )


    def test_parse_anlage2_table_extra_text(self):
        cfg = Anlage2Config.get_instance()
        Anlage2ColumnHeading.objects.create(
            config=cfg,
            field_name="technisch_vorhanden",
            text="Steht technisch zur Verf\u00fcgung?",
        )
        Anlage2ColumnHeading.objects.create(
            config=cfg,
            field_name="einsatz_bei_telefonica",
            text="einsatzweise bei telef\u00f3nica: soll die funktion verwendet werden?",
        )
        Anlage2ColumnHeading.objects.create(
            config=cfg,
            field_name="zur_lv_kontrolle",
            text="einsatzweise bei telef\u00f3nica: soll zur \u00fcberwachung von leistung oder verhalten verwendet werden?",
        )

        doc = Document()
        table = doc.add_table(rows=2, cols=5)
        table.cell(0, 0).text = "Funktion"
        table.cell(0, 1).text = "Steht technisch zur Verf\u00fcgung?\n\nJa/nein"
        table.cell(0, 2).text = (
            "einsatzweise bei Telef\u00f3nica: soll die Funktion verwendet werden?\n\nJa/nein"
        )
        table.cell(0, 3).text = (
            "einsatzweise bei Telef\u00f3nica: soll zur \u00dcberwachung von Leistung oder Verhalten verwendet werden?\n\nJa / nein"
        )
        table.cell(0, 4).text = "KI-Beteiligung"

        table.cell(1, 0).text = "Login"
        table.cell(1, 1).text = "Ja"
        table.cell(1, 2).text = "Nein"
        table.cell(1, 3).text = "Nein"
        table.cell(1, 4).text = "Ja"

        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        try:
            data = parse_anlage2_table(Path(tmp.name))
        finally:
            Path(tmp.name).unlink(missing_ok=True)

        self.assertEqual(
            data,
            [
                {
                    "funktion": "Login",
                    "technisch_verfuegbar": {"value": True, "note": None},
                    "einsatz_telefonica": {"value": False, "note": None},
                    "zur_lv_kontrolle": {"value": False, "note": None},
                    "ki_beteiligung": {"value": True, "note": None},
                }
            ],
        )

    def test_parse_anlage2_table_alias_conflict(self):
        cfg = Anlage2Config.get_instance()
        conflict = "Konflikt"
        Anlage2ColumnHeading.objects.create(
            config=cfg, field_name="technisch_vorhanden", text=conflict
        )
        Anlage2ColumnHeading.objects.create(
            config=cfg, field_name="einsatz_bei_telefonica", text=conflict
        )

        doc = Document()
        table = doc.add_table(rows=2, cols=4)
        table.cell(0, 0).text = "Funktion"
        table.cell(0, 1).text = conflict
        table.cell(0, 2).text = "Einsatz bei Telefónica"
        table.cell(0, 3).text = "Zur LV-Kontrolle"

        table.cell(1, 0).text = "Login"
        table.cell(1, 1).text = "Ja"
        table.cell(1, 2).text = "Nein"
        table.cell(1, 3).text = "Nein"

        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        try:
            with self.assertRaises(ValueError):
                parse_anlage2_table(Path(tmp.name))
        finally:
            Path(tmp.name).unlink(missing_ok=True)

    def test_parse_anlage2_text(self):
        func = Anlage2Function.objects.create(
            name="Login",
            detection_phrases={"name_aliases": ["login"]},
        )
        Anlage2SubQuestion.objects.create(
            funktion=func,
            frage_text="Warum?",
            detection_phrases={"name_aliases": ["warum"]},
        )
        cfg = Anlage2Config.get_instance()
        Anlage2GlobalPhrase.objects.create(
            config=cfg, phrase_type="technisch_verfuegbar_true", phrase_text="tv ja"
        )
        Anlage2GlobalPhrase.objects.create(
            config=cfg, phrase_type="technisch_verfuegbar_false", phrase_text="tv nein"
        )
        Anlage2GlobalPhrase.objects.create(
            config=cfg, phrase_type="ki_beteiligung_false", phrase_text="ki nein"
        )
        text = "Login tv ja ki nein\nWarum tv nein"
        data = parse_anlage2_text(text)
        self.assertEqual(
            data,
            [
                {
                    "funktion": "Login",
                    "technisch_verfuegbar": {"value": True, "note": None},
                    "ki_beteiligung": {"value": False, "note": None},
                },
                {
                    "funktion": "Login: Warum?",
                    "technisch_verfuegbar": {"value": False, "note": None},
                },
            ],
        )


class BVProjectFormTests(TestCase):
    def test_project_form_docx_validation(self):
        data = QueryDict(mutable=True)
        data.update(
            {
                "title": "",
                "beschreibung": "",
            }
        )
        data.setlist("software", ["A"])
        valid = BVProjectForm(data, {"docx_file": SimpleUploadedFile("t.docx", b"d")})
        self.assertTrue(valid.is_valid())
        invalid = BVProjectForm(data, {"docx_file": SimpleUploadedFile("t.txt", b"d")})
        self.assertFalse(invalid.is_valid())

    def test_upload_form_docx_validation(self):
        valid = BVProjectUploadForm(
            {}, {"docx_file": SimpleUploadedFile("t.docx", b"d")}
        )
        self.assertTrue(valid.is_valid())
        invalid = BVProjectUploadForm(
            {}, {"docx_file": SimpleUploadedFile("t.txt", b"d")}
        )
        self.assertFalse(invalid.is_valid())


class BVProjectFileTests(TestCase):
    def test_create_project_with_files(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        for i in range(1, 4):
            f = SimpleUploadedFile(f"f{i}.txt", b"data")
            BVProjectFile.objects.create(
                projekt=projekt,
                anlage_nr=i,
                upload=f,
                text_content="data",
            )
        self.assertEqual(projekt.anlagen.count(), 3)
        self.assertListEqual(
            list(projekt.anlagen.values_list("anlage_nr", flat=True)), [1, 2, 3]
        )


class ProjektFileUploadTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user("user", password="pass")
        self.client.login(username="user", password="pass")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")

    def test_docx_upload_extracts_text(self):
        doc = Document()
        doc.add_paragraph("Docx Inhalt")
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("t.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)

        url = reverse("projekt_file_upload", args=[self.projekt.pk])
        resp = self.client.post(
            url,
            {"anlage_nr": 1, "upload": upload, "manual_comment": ""},
            format="multipart",
        )
        self.assertEqual(resp.status_code, 302)
        file_obj = self.projekt.anlagen.first()
        self.assertIsNotNone(file_obj)
        self.assertIn("Docx Inhalt", file_obj.text_content)


class BVProjectModelTests(TestCase):
    def test_title_auto_set_from_software(self):
        projekt = BVProject.objects.create(software_typen="A, B", beschreibung="x")
        self.assertEqual(projekt.title, "A, B")

    def test_title_preserved_when_set(self):
        projekt = BVProject.objects.create(
            title="X", software_typen="A", beschreibung="x"
        )
        self.assertEqual(projekt.title, "X")


class WorkflowTests(TestCase):
    def test_default_status(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        self.assertEqual(projekt.status.key, "NEW")

    def test_set_project_status(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        set_project_status(projekt, "CLASSIFIED")
        projekt.refresh_from_db()
        self.assertEqual(projekt.status.key, "CLASSIFIED")

    def test_invalid_status(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        with self.assertRaises(ValueError):
            set_project_status(projekt, "XXX")

    def test_set_project_status_new_states(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        for status in [
            "IN_PRUEFUNG_ANLAGE_X",
            "FB_IN_PRUEFUNG",
            "ENDGEPRUEFT",
        ]:
            set_project_status(projekt, status)
            projekt.refresh_from_db()
            self.assertEqual(projekt.status.key, status)

    def test_status_history_created(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        self.assertEqual(projekt.status_history.count(), 1)
        set_project_status(projekt, "CLASSIFIED")
        self.assertEqual(projekt.status_history.count(), 2)


class LLMTasksTests(TestCase):
    maxDiff = None

    def test_classify_system(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        BVProjectFile.objects.create(
            projekt=projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content="Testtext",
        )
        with patch(
            "core.llm_tasks.query_llm",
            return_value='{"kategorie":"X","begruendung":"ok"}',
        ):
            data = classify_system(projekt.pk)
        projekt.refresh_from_db()
        self.assertEqual(projekt.classification_json["kategorie"]["value"], "X")
        self.assertEqual(projekt.status.key, "CLASSIFIED")
        self.assertEqual(data["kategorie"]["value"], "X")

    def test_check_anlage2(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        BVProjectFile.objects.create(
            projekt=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content="Anlagetext",
        )
        func = Anlage2Function.objects.create(name="Login")
        llm_reply = json.dumps({"technisch_verfuegbar": True})
        with patch("core.llm_tasks.query_llm", return_value=llm_reply) as mock_q:
            data = check_anlage2(projekt.pk)
        mock_q.assert_called()
        file_obj = projekt.anlagen.get(anlage_nr=2)
        self.assertTrue(data["functions"][0]["technisch_verfuegbar"])
        self.assertEqual(data["functions"][0]["source"], "llm")
        res = Anlage2FunctionResult.objects.get(projekt=projekt, funktion=func)
        self.assertEqual(res.source, "llm")

    def test_check_anlage2_llm_receives_text(self):
        """Der LLM-Prompt enthält den bekannten Text."""
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        BVProjectFile.objects.create(
            projekt=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content="Testinhalt Anlage2",
        )
        func = Anlage2Function.objects.create(name="Login")
        llm_reply = json.dumps({"technisch_verfuegbar": False})
        with patch("core.llm_tasks.query_llm", return_value=llm_reply) as mock_q:
            data = check_anlage2(projekt.pk)
        self.assertIn("Testinhalt Anlage2", mock_q.call_args_list[0].args[0].text)
        file_obj = projekt.anlagen.get(anlage_nr=2)
        self.assertEqual(data["functions"][0]["funktion"], "Login")

    def test_check_anlage2_prompt_contains_text(self):
        """Der Prompt enth\u00e4lt den gesamten Anlagentext."""
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        BVProjectFile.objects.create(
            projekt=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content="Testinhalt Anlage2",
        )
        func = Anlage2Function.objects.create(name="Login")
        llm_reply = json.dumps({"technisch_verfuegbar": False})
        with patch("core.llm_tasks.query_llm", return_value=llm_reply) as mock_q:
            data = check_anlage2(projekt.pk)
        prompt = mock_q.call_args_list[0].args[0].text
        self.assertIn("Testinhalt Anlage2", prompt)
        file_obj = projekt.anlagen.get(anlage_nr=2)
        self.assertEqual(data["functions"][0]["funktion"], "Login")

    def test_check_anlage2_parser(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        doc = Document()
        table = doc.add_table(rows=2, cols=5)
        table.cell(0, 0).text = "Funktion"
        table.cell(0, 1).text = "Technisch vorhanden"
        table.cell(0, 2).text = "Einsatz bei Telefónica"
        table.cell(0, 3).text = "Zur LV-Kontrolle"
        table.cell(0, 4).text = "KI-Beteiligung"
        table.cell(1, 0).text = "Login"
        table.cell(1, 1).text = "Ja"
        table.cell(1, 2).text = "Nein"
        table.cell(1, 3).text = "Nein"
        table.cell(1, 4).text = "Ja"
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("b.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)
        BVProjectFile.objects.create(
            projekt=projekt,
            anlage_nr=2,
            upload=upload,
            text_content="ignored",
        )
        func = Anlage2Function.objects.create(name="Login")

        with patch("core.llm_tasks.query_llm") as mock_q:
            data = check_anlage2(projekt.pk)
        mock_q.assert_not_called()
        expected = {
            "task": "check_anlage2",
            "functions": [
                {
                    "funktion": "Login",
                    "technisch_verfuegbar": {"value": True, "note": None},
                    "ki_beteiligung": {"value": True, "note": None},
                    "source": "parser",
                }
            ],
        }
        file_obj = projekt.anlagen.get(anlage_nr=2)
        self.assertEqual(data, expected)
        self.assertEqual(file_obj.analysis_json, expected)

    def test_analyse_anlage2(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="b")
        BVProjectFile.objects.create(
            projekt=projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content="Text A1",
        )
        doc = Document()
        table = doc.add_table(rows=2, cols=5)
        table.cell(0, 0).text = "Funktion"
        table.cell(0, 1).text = "Technisch vorhanden"
        table.cell(0, 2).text = "Einsatz bei Telefónica"
        table.cell(0, 3).text = "Zur LV-Kontrolle"
        table.cell(0, 4).text = "KI-Beteiligung"
        table.cell(1, 0).text = "Login"
        table.cell(1, 1).text = "Ja"
        table.cell(1, 2).text = "Nein"
        table.cell(1, 3).text = "Nein"
        table.cell(1, 4).text = "Ja"
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            upload = SimpleUploadedFile("b.docx", fh.read())
        Path(tmp.name).unlink(missing_ok=True)
        BVProjectFile.objects.create(
            projekt=projekt,
            anlage_nr=2,
            upload=upload,
            text_content="- Login",
        )

        data = analyse_anlage2(projekt.pk)
        file_obj = projekt.anlagen.get(anlage_nr=2)
        self.assertEqual(data["missing"]["value"], [])
        self.assertEqual(file_obj.analysis_json["additional"]["value"], [])

    def test_check_anlage1_new_schema(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        BVProjectFile.objects.create(
            projekt=projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content="Text",
        )
        expected = {
            "task": "check_anlage1",
            "version": 1,
            "anlage": 1,
            "companies": {"value": ["ACME"], "editable": True},
            "departments": {"value": ["IT"], "editable": True},
            "it_integration_summary": {"value": "Summe", "editable": True},
            "vendors": {"value": [], "editable": True},
            "question4_raw": {"value": "raw", "editable": False},
            "purpose_summary": {"value": "Zweck", "editable": True},
            "purpose_missing": {"value": False, "editable": True},
            "documentation_links": {"value": [], "editable": True},
            "replaced_systems": {"value": [], "editable": True},
            "legacy_functions": {"value": [], "editable": True},
            "question9_raw": {"value": "", "editable": True},
            "inconsistencies": {"value": [], "editable": True},
            "keywords": {"value": [], "editable": True},
            "plausibility_score": {"value": 0.5, "editable": True},
            "manual_comments": {"value": {}, "editable": True},
        }
        llm_reply = json.dumps({**expected, "questions": {}})
        eval_reply = json.dumps({"status": "ok", "hinweis": "", "vorschlag": ""})
        with patch(
            "core.llm_tasks.query_llm", side_effect=[llm_reply] + [eval_reply] * 9
        ):
            data = check_anlage1(projekt.pk)
        file_obj = projekt.anlagen.get(anlage_nr=1)
        answers = [
            ["ACME"],
            ["IT"],
            "leer",
            "raw",
            "Zweck",
            "leer",
            "leer",
            "leer",
            "leer",
        ]
        nums = [q.num for q in Anlage1Question.objects.order_by("num")]
        expected["questions"] = {
            str(i): {
                "answer": answers[i - 1],
                "status": "ok",
                "hinweis": "",
                "vorschlag": "",
            }
            for i in nums
        }
        self.assertEqual(file_obj.analysis_json, expected)
        self.assertEqual(data, expected)

    def test_check_anlage1_parser(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        text = (
            "Frage 1: Extrahiere alle Unternehmen als Liste.\u00b6A1\u00b6"
            "Frage 2: Extrahiere alle Fachbereiche als Liste.\u00b6A2"
        )
        BVProjectFile.objects.create(
            projekt=projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content=text,
        )
        eval_reply = json.dumps({"status": "ok", "hinweis": "", "vorschlag": ""})
        with patch("core.llm_tasks.query_llm", side_effect=[eval_reply] * 9):
            data = check_anlage1(projekt.pk)
        answers = {
            "1": "A1",
            "2": "A2",
        }
        nums = [q.num for q in Anlage1Question.objects.order_by("num")]
        expected_questions = {
            str(i): {
                "answer": answers.get(str(i), "leer"),
                "status": "ok",
                "hinweis": "",
                "vorschlag": "",
            }
            for i in nums
        }
        file_obj = projekt.anlagen.get(anlage_nr=1)
        self.assertEqual(data["source"], "parser")
        self.assertEqual(data["questions"]["1"]["answer"], "A1")
        self.assertEqual(data["questions"]["2"]["answer"], "A2")
        self.assertEqual(file_obj.analysis_json, data)

    def test_parse_anlage1_questions_extra(self):
        Anlage1Question.objects.create(
            num=10,
            text="Frage 10: Test?",
            enabled=True,
            parser_enabled=True,
            llm_enabled=True,
        )
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        text = (
            "Frage 1: Extrahiere alle Unternehmen als Liste.\u00b6A1\u00b6"
            "Frage 10: Test?\u00b6A10"
        )
        BVProjectFile.objects.create(
            projekt=projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content=text,
        )
        eval_reply = json.dumps({"status": "ok", "hinweis": "", "vorschlag": ""})
        nums = [q.num for q in Anlage1Question.objects.order_by("num")]
        with patch("core.llm_tasks.query_llm", side_effect=[eval_reply] * len(nums)):
            data = check_anlage1(projekt.pk)
        q_data = data["questions"]
        self.assertEqual(q_data["10"]["answer"], "A10")

    def test_parse_anlage1_questions_without_numbers(self):
        """Prüft die Extraktion ohne nummerierte Fragen."""
        # Frage-Texte ohne Präfix "Frage X:" speichern
        q1 = Anlage1Question.objects.get(num=1)
        q2 = Anlage1Question.objects.get(num=2)
        q1.text = q1.text.split(": ", 1)[1]
        q2.text = q2.text.split(": ", 1)[1]
        q1.save(update_fields=["text"])
        q2.save(update_fields=["text"])
        v1 = q1.variants.first()
        v2 = q2.variants.first()
        v1.text = q1.text
        v2.text = q2.text
        v1.save()
        v2.save()

        text = f"{q1.text}\u00b6A1\u00b6{q2.text}\u00b6A2"
        parsed = parse_anlage1_questions(text)
        self.assertEqual(
            parsed,
            {
                "1": {"answer": "A1", "found_num": None},
                "2": {"answer": "A2", "found_num": None},
            },
        )

    def test_parse_anlage1_questions_with_variant(self):
        q1 = Anlage1Question.objects.get(num=1)
        q1.variants.create(text="Alternative Frage 1?")
        text = "Alternative Frage 1?\u00b6A1"
        parsed = parse_anlage1_questions(text)
        self.assertEqual(parsed, {"1": {"answer": "A1", "found_num": "1"}})

    def test_parse_anlage1_questions_with_newlines(self):
        """Extraktion funktioniert trotz Zeilenumbr\u00fcche."""
        text = (
            "Frage 1:\nExtrahiere alle Unternehmen als Liste.\nA1\n"
            "Frage 2:\nExtrahiere alle Fachbereiche als Liste.\nA2"
        )
        parsed = parse_anlage1_questions(text)
        self.assertEqual(
            parsed,
            {
                "1": {"answer": "A1", "found_num": "1"},
                "2": {"answer": "A2", "found_num": "2"},
            },
        )

    def test_parse_anlage1_questions_respects_parser_enabled(self):
        q2 = Anlage1Question.objects.get(num=2)
        q2.parser_enabled = False
        q2.save(update_fields=["parser_enabled"])
        text = "Frage 1: Extrahiere alle Unternehmen als Liste.\u00b6A1"
        parsed = parse_anlage1_questions(text)
        self.assertEqual(parsed, {"1": {"answer": "A1", "found_num": "1"}})

    def test_wrong_question_number_sets_hint(self):
        """Hinweis wird gesetzt, wenn die Nummer nicht passt."""
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        text = "Frage 1.2: Extrahiere alle Unternehmen als Liste.\u00b6A1"
        BVProjectFile.objects.create(
            projekt=projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content=text,
        )
        eval_reply = json.dumps({"status": "ok", "hinweis": "", "vorschlag": ""})
        with patch("core.llm_tasks.query_llm", side_effect=[eval_reply] * 9):
            analysis = check_anlage1(projekt.pk)
        hint = analysis["questions"]["1"]["hinweis"]
        self.assertIn("Frage 1.2 statt 1", hint)

    def test_generate_gutachten_twice_replaces_file(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        first = generate_gutachten(projekt.pk, text="Alt")
        second = generate_gutachten(projekt.pk, text="Neu")
        try:
            self.assertTrue(second.exists())
            self.assertNotEqual(first, second)
            self.assertFalse(first.exists())
        finally:
            second.unlink(missing_ok=True)

    def test_check_anlage1_ignores_disabled_questions(self):
        Anlage1Config.objects.create()  # Standardwerte
        q1 = Anlage1Question.objects.get(num=1)
        q1.llm_enabled = False
        q1.save(update_fields=["llm_enabled"])
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        BVProjectFile.objects.create(
            projekt=projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content="Text",
        )
        eval_reply = json.dumps({"status": "ok", "hinweis": "", "vorschlag": ""})
        enabled_count = Anlage1Question.objects.filter(llm_enabled=True).count()
        with patch(
            "core.llm_tasks.query_llm",
            side_effect=['{"task": "check_anlage1"}'] + [eval_reply] * enabled_count,
        ) as mock_q:
            data = check_anlage1(projekt.pk)
        prompt = mock_q.call_args_list[0].args[0].text
        self.assertNotIn("Frage 1", prompt)
        self.assertIn("1", data["questions"])
        self.assertIsNone(data["questions"]["1"]["status"])

    def test_parse_anlage2_question_list(self):
        text = "Welche Funktionen bietet das System?\u00b6- Login\u00b6- Suche"
        parsed = _parse_anlage2(text)
        self.assertEqual(parsed, ["Login", "Suche"])

    def test_parse_anlage2_table_llm(self):
        text = "Funktion | Beschreibung\u00b6Login | a\u00b6Suche | b"
        with patch(
            "core.llm_tasks.query_llm", return_value='["Login", "Suche"]'
        ) as mock_q:
            parsed = _parse_anlage2(text)
        mock_q.assert_called_once()
        self.assertEqual(parsed, ["Login", "Suche"])


class PromptTests(TestCase):
    def test_get_prompt_returns_default(self):
        self.assertEqual(get_prompt("unknown", "foo"), "foo")

    def test_get_prompt_returns_db_value(self):
        p, _ = Prompt.objects.get_or_create(
            name="classify_system", defaults={"text": "orig"}
        )
        p.text = "DB"
        p.save()
        self.assertEqual(get_prompt("classify_system", "x"), "DB")

    def test_check_anlage1_prompt_text(self):
        p = Prompt.objects.get(name="check_anlage1")
        expected = (
            "System: Du bist ein juristisch-technischer Prüf-Assistent für Systembeschreibungen.\n\n"
            "Frage 1: Extrahiere alle Unternehmen als Liste.\n"
            "Frage 2: Extrahiere alle Fachbereiche als Liste.\n"
            "IT-Landschaft: Fasse den Abschnitt zusammen, der die Einbettung in die IT-Landschaft beschreibt.\n"
            "Frage 3: Liste alle Hersteller und Produktnamen auf.\n"
            "Frage 4: Lege den Textblock als question4_raw ab.\n"
            "Frage 5: Fasse den Zweck des Systems in einem Satz.\n"
            "Frage 6: Extrahiere Web-URLs.\n"
            "Frage 7: Extrahiere ersetzte Systeme.\n"
            "Frage 8: Extrahiere Legacy-Funktionen.\n"
            "Frage 9: Lege den Text als question9_raw ab.\n"
            "Konsistenzprüfung und Stichworte. Gib ein JSON im vorgegebenen Schema zurück.\n\n"
        )
        self.assertEqual(p.text, expected)


class AdminPromptsViewTests(TestCase):
    def setUp(self):
        admin_group = Group.objects.create(name="admin")
        self.user = User.objects.create_user("padmin", password="pass")
        self.user.groups.add(admin_group)
        self.client.login(username="padmin", password="pass")
        self.prompt = Prompt.objects.create(name="p1", text="orig")

    def test_update_prompt(self):
        url = reverse("admin_prompts")
        resp = self.client.post(
            url,
            {
                "pk": self.prompt.id,
                "text": "neu",
                "action": "save",
                "use_system_role": "on",
            },
        )
        self.assertRedirects(resp, url)
        self.prompt.refresh_from_db()
        self.assertEqual(self.prompt.text, "neu")

    def test_delete_prompt(self):
        url = reverse("admin_prompts")
        resp = self.client.post(url, {"pk": self.prompt.id, "action": "delete"})
        self.assertRedirects(resp, url)
        self.assertFalse(Prompt.objects.filter(id=self.prompt.id).exists())


class ReportingTests(TestCase):
    def test_gap_analysis_file_created(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        BVProjectFile.objects.create(
            projekt=projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content="Testtext",
            analysis_json={"ok": {"value": True, "editable": True}},
        )
        path = generate_gap_analysis(projekt)
        try:
            self.assertTrue(path.exists())
        finally:
            path.unlink(missing_ok=True)

    def test_management_summary_includes_comment(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        BVProjectFile.objects.create(
            projekt=projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content="Testtext",
            manual_comment="Hinweis",
        )
        path = generate_management_summary(projekt)
        try:
            doc = Document(path)
            text = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn("Hinweis", text)
        finally:
            path.unlink(missing_ok=True)

    def test_manual_analysis_overrides(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        BVProjectFile.objects.create(
            projekt=projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content="Testtext",
            analysis_json={"foo": {"value": "orig", "editable": True}},
            manual_analysis_json={"foo": "manual"},
        )
        path1 = generate_gap_analysis(projekt)
        try:
            doc = Document(path1)
            text = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn('"foo": "manual"', text)
            self.assertNotIn('"foo": "orig"', text)
        finally:
            path1.unlink(missing_ok=True)

        path2 = generate_management_summary(projekt)
        try:
            doc = Document(path2)
            text = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn('"foo": "manual"', text)
            self.assertNotIn('"foo": "orig"', text)
        finally:
            path2.unlink(missing_ok=True)


class ProjektFileCheckViewTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user("user2", password="pass")
        self.client.login(username="user2", password="pass")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        BVProjectFile.objects.create(
            projekt=self.projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content="Text",
        )

    def test_file_check_endpoint_saves_json(self):
        url = reverse("projekt_file_check", args=[self.projekt.pk, 1])
        expected = {
            "task": "check_anlage1",
        }
        llm_reply = json.dumps(
            {
                "companies": None,
                "departments": None,
                "vendors": None,
                "question4_raw": None,
                "purpose_summary": None,
                "documentation_links": None,
                "replaced_systems": None,
                "legacy_functions": None,
                "question9_raw": None,
            }
        )
        eval_reply = json.dumps({"status": "ok", "hinweis": "", "vorschlag": ""})
        with patch(
            "core.llm_tasks.query_llm", side_effect=[llm_reply] + [eval_reply] * 9
        ):
            resp = self.client.post(url)
        self.assertEqual(resp.status_code, 200)
        file_obj = self.projekt.anlagen.get(anlage_nr=1)
        nums = [q.num for q in Anlage1Question.objects.order_by("num")]
        expected["questions"] = {
            str(i): {"answer": "leer", "status": "ok", "hinweis": "", "vorschlag": ""}
            for i in nums
        }
        self.assertEqual(file_obj.analysis_json, expected)

    def test_file_check_pk_endpoint_saves_json(self):
        file_obj = self.projekt.anlagen.get(anlage_nr=1)
        url = reverse("projekt_file_check_pk", args=[file_obj.pk])
        expected = {"task": "check_anlage1"}
        llm_reply = json.dumps(
            {
                "companies": None,
                "departments": None,
                "vendors": None,
                "question4_raw": None,
                "purpose_summary": None,
                "documentation_links": None,
                "replaced_systems": None,
                "legacy_functions": None,
                "question9_raw": None,
            }
        )
        eval_reply = json.dumps({"status": "ok", "hinweis": "", "vorschlag": ""})
        with patch(
            "core.llm_tasks.query_llm", side_effect=[llm_reply] + [eval_reply] * 9
        ):
            resp = self.client.post(url)
        self.assertEqual(resp.status_code, 200)
        file_obj.refresh_from_db()
        nums = [q.num for q in Anlage1Question.objects.order_by("num")]
        expected["questions"] = {
            str(i): {"answer": "leer", "status": "ok", "hinweis": "", "vorschlag": ""}
            for i in nums
        }
        self.assertEqual(file_obj.analysis_json, expected)


class ProjektFileJSONEditTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user("user3", password="pass")
        self.client.login(username="user3", password="pass")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        self.file = BVProjectFile.objects.create(
            projekt=self.projekt,
            anlage_nr=3,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content="Text",
            analysis_json={"old": {"value": True, "editable": True}},
        )
        self.anlage1 = BVProjectFile.objects.create(
            projekt=self.projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("b.txt", b"data"),
            text_content="Text",
            analysis_json={
                "questions": {
                    "1": {
                        "answer": "foo",
                        "status": None,
                        "hinweis": "",
                        "vorschlag": "",
                    }
                }
            },
        )

    def test_edit_json_updates_and_reports(self):
        url = reverse("projekt_file_edit_json", args=[self.file.pk])
        resp = self.client.post(
            url,
            {
                "analysis_json": '{"new": 1}',
                "manual_analysis_json": '{"manual": 2}',
            },
        )
        self.assertEqual(resp.status_code, 302)
        self.file.refresh_from_db()
        self.assertEqual(self.file.analysis_json["new"], 1)
        self.assertEqual(self.file.manual_analysis_json["manual"], 2)
        path = generate_gap_analysis(self.projekt)
        try:
            doc = Document(path)
            text = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn('"manual": 2', text)
            self.assertNotIn('"old": true', text.lower())
        finally:
            path.unlink(missing_ok=True)

    def test_invalid_json_shows_error(self):
        url = reverse("projekt_file_edit_json", args=[self.file.pk])
        resp = self.client.post(
            url,
            {"analysis_json": "{", "manual_analysis_json": "{}"},
        )
        self.assertEqual(resp.status_code, 200)
        self.file.refresh_from_db()
        self.assertEqual(
            self.file.analysis_json, {"old": {"value": True, "editable": True}}
        )

    def test_question_review_saved(self):
        url = reverse("projekt_file_edit_json", args=[self.anlage1.pk])
        resp = self.client.post(
            url,
            {"q1_ok": "on", "q1_note": "Hinweis"},
        )
        self.assertRedirects(resp, reverse("projekt_detail", args=[self.projekt.pk]))
        self.anlage1.refresh_from_db()
        self.assertTrue(self.anlage1.question_review["1"]["ok"])
        self.assertEqual(self.anlage1.question_review["1"]["note"], "Hinweis")

    def test_question_review_extended_fields_saved(self):
        url = reverse("projekt_file_edit_json", args=[self.anlage1.pk])
        resp = self.client.post(
            url,
            {
                "q1_status": "unvollst\u00e4ndig",
                "q1_hinweis": "Fehlt",
                "q1_vorschlag": "Mehr Infos",
            },
        )
        self.assertRedirects(resp, reverse("projekt_detail", args=[self.projekt.pk]))
        self.anlage1.refresh_from_db()
        data = self.anlage1.question_review["1"]
        self.assertEqual(data["status"], "unvollst\u00e4ndig")
        self.assertEqual(data["hinweis"], "Fehlt")
        self.assertEqual(data["vorschlag"], "Mehr Infos")

    def test_question_review_prefill_from_analysis(self):
        """Initialwerte stammen aus der automatischen Analyse."""
        self.anlage1.question_review = None
        self.anlage1.analysis_json = {
            "questions": {
                "1": {
                    "answer": "A",
                    "status": "ok",
                    "hinweis": "H",
                    "vorschlag": "V",
                }
            }
        }
        self.anlage1.save()

        url = reverse("projekt_file_edit_json", args=[self.anlage1.pk])
        resp = self.client.get(url)
        form = resp.context["form"]
        self.assertEqual(form.initial["q1_status"], "ok")
        self.assertEqual(form.initial["q1_hinweis"], "H")
        self.assertEqual(form.initial["q1_vorschlag"], "V")


class Anlage2ReviewTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user("rev", password="pass")
        self.client.login(username="rev", password="pass")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        self.file = BVProjectFile.objects.create(
            projekt=self.projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("c.txt", b"d"),
            text_content="Text",
            analysis_json={
                "functions": [
                    {
                        "funktion": "Login",
                        "technisch_vorhanden": {"value": True, "note": None},
                        "einsatz_bei_telefonica": {"value": False, "note": None},
                        "zur_lv_kontrolle": {"value": False, "note": None},
                        "ki_beteiligung": {"value": True, "note": None},
                    }
                ]
            },
        )
        self.func = Anlage2Function.objects.create(name="Login")
        self.sub = Anlage2SubQuestion.objects.create(funktion=self.func, frage_text="Warum?")

    def test_get_shows_table(self):
        url = reverse("projekt_file_edit_json", args=[self.file.pk])
        resp = self.client.get(url)
        self.assertContains(resp, "Login")
        self.assertContains(resp, "Warum?")
        self.assertContains(resp, f"name=\"func{self.func.id}_technisch_vorhanden\"")

    def test_post_saves_data(self):
        url = reverse("projekt_file_edit_json", args=[self.file.pk])
        resp = self.client.post(
            url,
            {
                f"func{self.func.id}_technisch_vorhanden": "on",
                f"sub{self.sub.id}_ki_beteiligung": "on",
            },
        )
        self.assertRedirects(resp, reverse("projekt_detail", args=[self.projekt.pk]))
        self.file.refresh_from_db()
        data = self.file.manual_analysis_json["functions"][str(self.func.id)]
        self.assertTrue(data["technisch_vorhanden"])
        self.assertTrue(data["subquestions"][str(self.sub.id)]["ki_beteiligung"])

    def test_prefill_from_analysis(self):
        """Die Formulardaten verwenden Analysewerte als Vorgabe."""
        self.file.manual_analysis_json = None
        self.file.analysis_json = {
            "functions": [
                {
                    "funktion": "Login",
                    "technisch_verfuegbar": {"value": True, "note": None},
                    "einsatz_telefonica": {"value": True, "note": None},
                    "zur_lv_kontrolle": {"value": True, "note": None},
                }
            ]
        }
        self.file.save()

        url = reverse("projekt_file_edit_json", args=[self.file.pk])
        resp = self.client.get(url)
        field = f"func{self.func.id}_technisch_vorhanden"
        self.assertTrue(resp.context["form"].initial[field])


class WorkerGenerateGutachtenTests(TestCase):
    def setUp(self):
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        BVProjectFile.objects.create(
            projekt=self.projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content="Text",
        )
        self.knowledge = SoftwareKnowledge.objects.create(
            projekt=self.projekt,
            software_name="A",
            is_known_by_llm=True,
            description="",
        )

    def test_worker_creates_file(self):
        with patch("core.llm_tasks.query_llm", return_value="Text"):
            path = worker_generate_gutachten(self.projekt.pk, self.knowledge.pk)
        self.projekt.refresh_from_db()
        self.assertTrue(self.projekt.gutachten_file.name)
        self.assertEqual(self.projekt.status.key, "GUTACHTEN_OK")
        self.assertEqual(Gutachten.objects.filter(software_knowledge=self.knowledge).count(), 1)
        Path(path).unlink(missing_ok=True)


class GutachtenEditDeleteTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user("euser", password="pass")
        self.client.login(username="euser", password="pass")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        doc = Document()
        doc.add_paragraph("Alt")
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        tmp.close()
        with open(tmp.name, "rb") as fh:
            self.projekt.gutachten_file.save(
                "g.docx", SimpleUploadedFile("g.docx", fh.read())
            )
        Path(tmp.name).unlink(missing_ok=True)
        self.knowledge = SoftwareKnowledge.objects.create(
            projekt=self.projekt,
            software_name="A",
            is_known_by_llm=True,
        )
        self.gutachten = Gutachten.objects.create(software_knowledge=self.knowledge, text="Alt")

    def test_view_shows_content(self):
        url = reverse("gutachten_view", args=[self.gutachten.pk])
        resp = self.client.get(url)
        self.assertContains(resp, "Alt")

    def test_edit_updates_text(self):
        url = reverse("gutachten_edit", args=[self.gutachten.pk])
        resp = self.client.post(url, {"text": "Neu"})
        self.assertRedirects(resp, reverse("gutachten_view", args=[self.gutachten.pk]))
        self.gutachten.refresh_from_db()
        self.assertEqual(self.gutachten.text, "Neu")

    def test_delete_removes_file(self):
        url = reverse("gutachten_delete", args=[self.gutachten.pk])
        resp = self.client.post(url)
        self.assertRedirects(resp, reverse("projekt_detail", args=[self.projekt.pk]))
        self.assertFalse(Gutachten.objects.filter(pk=self.gutachten.pk).exists())


class ProjektFileCheckResultTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user("vuser", password="pass")
        self.client.login(username="vuser", password="pass")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        self.file = BVProjectFile.objects.create(
            projekt=self.projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content="Text",
        )
        self.file2 = BVProjectFile.objects.create(
            projekt=self.projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("b.txt", b"data"),
            text_content="Text2",
        )

    def test_get_runs_check_and_redirects_to_edit(self):
        url = reverse("projekt_file_check_view", args=[self.file.pk])
        expected = {"task": "check_anlage1"}
        llm_reply = json.dumps(
            {
                "companies": None,
                "departments": None,
                "vendors": None,
                "question4_raw": None,
                "purpose_summary": None,
                "documentation_links": None,
                "replaced_systems": None,
                "legacy_functions": None,
                "question9_raw": None,
            }
        )
        eval_reply = json.dumps({"status": "ok", "hinweis": "", "vorschlag": ""})
        with patch(
            "core.llm_tasks.query_llm", side_effect=[llm_reply] + [eval_reply] * 9
        ):
            resp = self.client.get(url)
        self.assertRedirects(resp, reverse("projekt_file_edit_json", args=[self.file.pk]))
        self.file.refresh_from_db()
        nums = [q.num for q in Anlage1Question.objects.order_by("num")]
        expected["questions"] = {
            str(i): {"answer": "leer", "status": "ok", "hinweis": "", "vorschlag": ""}
            for i in nums
        }
        self.assertEqual(self.file.analysis_json, expected)

    def test_post_triggers_check_and_redirects(self):
        url = reverse("projekt_file_check_view", args=[self.file.pk])
        with patch("core.views.check_anlage1") as mock_func:
            mock_func.return_value = {"task": "check_anlage1"}
            resp = self.client.post(url)
        self.assertRedirects(resp, reverse("projekt_file_edit_json", args=[self.file.pk]))
        mock_func.assert_called_with(self.projekt.pk, model_name=None)

    def test_anlage2_uses_parser_by_default(self):
        url = reverse("projekt_file_check_view", args=[self.file2.pk])
        with patch("core.views.analyse_anlage2") as mock_func:
            mock_func.return_value = {"task": "analyse_anlage2"}
            resp = self.client.get(url)
        self.assertRedirects(resp, reverse("projekt_file_edit_json", args=[self.file2.pk]))
        mock_func.assert_called_with(self.projekt.pk, model_name=None)

    def test_llm_param_triggers_full_check(self):
        url = reverse("projekt_file_check_view", args=[self.file2.pk]) + "?llm=1"
        with patch("core.views.check_anlage2") as mock_func:
            mock_func.return_value = {"task": "check_anlage2"}
            resp = self.client.get(url)
        self.assertRedirects(resp, reverse("projekt_file_edit_json", args=[self.file2.pk]))
        mock_func.assert_called_with(self.projekt.pk, model_name=None)


class LLMConfigTests(TestCase):
    @override_settings(GOOGLE_API_KEY="x")
    @patch("google.generativeai.list_models")
    @patch("google.generativeai.configure")
    def test_ready_populates_models(self, mock_conf, mock_list):
        mock_list.return_value = [
            type("M", (), {"name": "m1"})(),
            type("M", (), {"name": "m2"})(),
        ]
        apps.get_app_config("core").ready()
        cfg = LLMConfig.objects.first()
        self.assertIsNotNone(cfg)
        self.assertEqual(cfg.available_models, ["m1", "m2"])
        self.assertTrue(cfg.models_changed)

    @override_settings(GOOGLE_API_KEY="x")
    @patch("google.generativeai.list_models")
    @patch("google.generativeai.configure")
    def test_ready_updates_models(self, mock_conf, mock_list):
        LLMConfig.objects.create(available_models=["old"])
        mock_list.return_value = [type("M", (), {"name": "new"})()]
        apps.get_app_config("core").ready()
        cfg = LLMConfig.objects.first()
        self.assertEqual(cfg.available_models, ["new"])
        self.assertTrue(cfg.models_changed)



class Anlage2ConfigSingletonTests(TestCase):
    def test_single_instance_enforced(self):
        first = Anlage2Config.get_instance()
        from django.db import transaction
        with self.assertRaises(IntegrityError):
            with transaction.atomic():
                Anlage2Config.objects.create()
        self.assertEqual(Anlage2Config.objects.count(), 1)


class AdminModelsViewTests(TestCase):
    def setUp(self):
        admin_group = Group.objects.create(name="admin")
        self.user = User.objects.create_user("amodel", password="pass")
        self.user.groups.add(admin_group)
        self.client.login(username="amodel", password="pass")
        self.cfg = LLMConfig.objects.create(
            default_model="a",
            gutachten_model="a",
            anlagen_model="a",
            available_models=["a", "b"],
        )

    def test_update_models(self):
        url = reverse("admin_models")
        resp = self.client.post(
            url,
            {
                "default_model": "b",
                "gutachten_model": "b",
                "anlagen_model": "b",
            },
        )
        self.assertRedirects(resp, url)
        self.cfg.refresh_from_db()
        self.assertEqual(self.cfg.default_model, "b")
        self.assertEqual(self.cfg.gutachten_model, "b")
        self.assertEqual(self.cfg.anlagen_model, "b")


class Anlage1EmailTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user("emailer", password="pass")
        self.client.login(username="emailer", password="pass")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        self.file = BVProjectFile.objects.create(
            projekt=self.projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"data"),
            question_review={"1": {"vorschlag": "Text"}},
        )

    def test_generate_email(self):
        url = reverse("anlage1_generate_email", args=[self.file.pk])
        with patch("core.views.query_llm", return_value="Mail"):
            resp = self.client.post(url)
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.json()["text"], "Mail")


class TileVisibilityTests(TestCase):
    def setUp(self):
        admin_group = Group.objects.create(name="admin")
        self.user = User.objects.create_user("tileuser", password="pass")
        self.user.groups.add(admin_group)
        work = Area.objects.get_or_create(slug="work", defaults={"name": "Work"})[0]
        self.personal = Area.objects.get_or_create(
            slug="personal", defaults={"name": "Personal"}
        )[0]
        self.talkdiary = Tile.objects.get_or_create(
            slug="talkdiary",
            defaults={
                "name": "TalkDiary",
                "bereich": self.personal,
                "url_name": "talkdiary_personal",
            },
        )[0]
        self.projekt = Tile.objects.get_or_create(
            slug="projektverwaltung",
            defaults={
                "name": "Projektverwaltung",
                "bereich": work,
                "url_name": "projekt_list",
            },
        )[0]
        self.cfg = LLMConfig.objects.first() or LLMConfig.objects.create(
            models_changed=False
        )
        self.client.login(username="tileuser", password="pass")

    def test_personal_without_access(self):
        resp = self.client.get(reverse("personal"))
        self.assertNotContains(resp, "TalkDiary")

    def test_personal_with_access(self):
        UserTileAccess.objects.create(user=self.user, tile=self.talkdiary)
        resp = self.client.get(reverse("personal"))
        self.assertContains(resp, "TalkDiary")

    def test_personal_with_image(self):
        UserTileAccess.objects.create(user=self.user, tile=self.talkdiary)
        self.talkdiary.image.save(
            "img.png",
            SimpleUploadedFile("img.png", b"data"),
            save=True,
        )
        resp = self.client.get(reverse("personal"))
        self.assertContains(resp, "<img", html=False)

    def test_work_with_projekt_access(self):
        UserTileAccess.objects.create(user=self.user, tile=self.projekt)
        resp = self.client.get(reverse("work"))
        self.assertContains(resp, "Projektverwaltung")
        self.assertNotContains(resp, "TalkDiary")

    def test_flag_reset_on_get(self):
        self.cfg.models_changed = True
        self.cfg.save(update_fields=["models_changed"])
        url = reverse("admin_models")
        self.client.get(url)
        self.cfg.refresh_from_db()
        self.assertFalse(self.cfg.models_changed)


class TileAccessTests(TestCase):
    def setUp(self):
        work = Area.objects.get_or_create(slug="work", defaults={"name": "Work"})[0]
        personal = Area.objects.get_or_create(
            slug="personal", defaults={"name": "Personal"}
        )[0]
        self.talkdiary = Tile.objects.get_or_create(
            slug="talkdiary",
            defaults={
                "name": "TalkDiary",
                "bereich": personal,
                "url_name": "talkdiary_personal",
            },
        )[0]
        self.projekt = Tile.objects.get_or_create(
            slug="projektverwaltung",
            defaults={
                "name": "Projektverwaltung",
                "bereich": work,
                "url_name": "projekt_list",
            },
        )[0]

    def _login(self, name: str) -> User:
        """Erzeugt einen Benutzer und loggt ihn ein."""
        user = User.objects.create_user(name, password="pass")
        self.client.login(username=name, password="pass")
        return user

    def test_talkdiary_access_denied_without_tile(self):
        self._login("nodiac")
        resp = self.client.get(reverse("personal"))
        self.assertNotContains(resp, "TalkDiary")
        resp = self.client.get(reverse("talkdiary_personal"))
        self.assertEqual(resp.status_code, 403)

    def test_talkdiary_access_allowed_with_tile(self):
        user = self._login("withdia")
        UserTileAccess.objects.create(user=user, tile=self.talkdiary)
        resp = self.client.get(reverse("personal"))
        self.assertContains(resp, "TalkDiary")
        resp = self.client.get(reverse("talkdiary_personal"))
        self.assertEqual(resp.status_code, 200)

    def test_projekt_access_denied_without_tile(self):
        self._login("noproj")
        resp = self.client.get(reverse("work"))
        self.assertNotContains(resp, "Projektverwaltung")
        resp = self.client.get(reverse("projekt_list"))
        self.assertEqual(resp.status_code, 403)

    def test_projekt_access_allowed_with_tile(self):
        user = self._login("withproj")
        UserTileAccess.objects.create(user=user, tile=self.projekt)
        resp = self.client.get(reverse("work"))
        self.assertContains(resp, "Projektverwaltung")
        resp = self.client.get(reverse("projekt_list"))
        self.assertEqual(resp.status_code, 200)


class LLMConfigNoticeMiddlewareTests(TestCase):
    def setUp(self):
        admin_group = Group.objects.create(name="admin")
        self.user = User.objects.create_user("llmadmin", password="pass")
        self.user.groups.add(admin_group)
        self.client.login(username="llmadmin", password="pass")
        LLMConfig.objects.create(models_changed=True)

    def test_message_shown(self):
        resp = self.client.get(reverse("home"))
        msgs = [m.message for m in resp.context["messages"]]
        self.assertTrue(any("LLM-Einstellungen" in m for m in msgs))


class HomeRedirectTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user("redir", password="pass")
        personal = Area.objects.get_or_create(
            slug="personal", defaults={"name": "Personal"}
        )[0]
        tile = Tile.objects.get_or_create(
            slug="talkdiary",
            defaults={
                "name": "TalkDiary",
                "bereich": personal,
                "url_name": "talkdiary_personal",
            },
        )[0]
        UserTileAccess.objects.create(user=self.user, tile=tile)
        self.client.login(username="redir", password="pass")

    def test_redirect_personal(self):
        resp = self.client.get(reverse("home"))
        self.assertRedirects(resp, reverse("personal"))


class AreaImageTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user("areauser", password="pass")
        self.client.login(username="areauser", password="pass")

    def test_home_without_images(self):
        Area.objects.get_or_create(slug="work", defaults={"name": "Work"})
        Area.objects.get_or_create(slug="personal", defaults={"name": "Personal"})
        resp = self.client.get(reverse("home"))
        self.assertNotContains(resp, 'alt="Work"', html=False)
        self.assertNotContains(resp, 'alt="Personal"', html=False)

    def test_home_with_images(self):
        work, _ = Area.objects.get_or_create(slug="work", defaults={"name": "Work"})
        personal, _ = Area.objects.get_or_create(
            slug="personal", defaults={"name": "Personal"}
        )
        work.image.save("w.png", SimpleUploadedFile("w.png", b"d"), save=True)
        personal.image.save("p.png", SimpleUploadedFile("p.png", b"d"), save=True)
        resp = self.client.get(reverse("home"))
        self.assertContains(resp, f'alt="{work.name}"', html=False)
        self.assertContains(resp, f'alt="{personal.name}"', html=False)


class RecordingDeleteTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user("recuser", password="pass")
        self.client.login(username="recuser", password="pass")
        self.personal = Area.objects.get_or_create(
            slug="personal", defaults={"name": "Personal"}
        )[0]
        self.tile = Tile.objects.get_or_create(
            slug="talkdiary",
            defaults={
                "name": "TalkDiary",
                "bereich": self.personal,
                "url_name": "talkdiary_personal",
            },
        )[0]
        UserTileAccess.objects.create(user=self.user, tile=self.tile)
        audio = SimpleUploadedFile("a.wav", b"data")
        transcript = SimpleUploadedFile("a.md", b"text")
        self.rec = Recording.objects.create(
            user=self.user,
            bereich=self.personal,
            audio_file=audio,
            transcript_file=transcript,
        )
        self.audio_path = Path(self.rec.audio_file.path)
        self.trans_path = Path(self.rec.transcript_file.path)

    def test_delete_own_recording(self):
        url = reverse("recording_delete", args=[self.rec.pk])
        resp = self.client.post(url)
        self.assertRedirects(resp, reverse("talkdiary_personal"))
        self.assertFalse(Recording.objects.filter(pk=self.rec.pk).exists())
        self.assertFalse(self.audio_path.exists())
        self.assertFalse(self.trans_path.exists())

    def test_delete_requires_post(self):
        url = reverse("recording_delete", args=[self.rec.pk])
        resp = self.client.get(url)
        self.assertEqual(resp.status_code, 405)
        self.assertTrue(Recording.objects.filter(pk=self.rec.pk).exists())

    def test_delete_other_user_recording(self):
        other = User.objects.create_user("other", password="pass")
        rec = Recording.objects.create(
            user=other,
            bereich=self.personal,
            audio_file=SimpleUploadedFile("b.wav", b"d"),
            transcript_file=SimpleUploadedFile("b.md", b"t"),
        )
        url = reverse("recording_delete", args=[rec.pk])
        resp = self.client.post(url)
        self.assertEqual(resp.status_code, 404)
        self.assertTrue(Recording.objects.filter(pk=rec.pk).exists())


class AdminAnlage1ViewTests(TestCase):
    def setUp(self):
        admin_group = Group.objects.create(name="admin")
        self.user = User.objects.create_user("a1admin", password="pass")
        self.user.groups.add(admin_group)
        self.client.login(username="a1admin", password="pass")

    def test_delete_question(self):
        url = reverse("admin_anlage1")
        questions = list(Anlage1Question.objects.all())
        q = questions[0]
        data = {}
        for question in questions:
            if question.id == q.id:
                data[f"delete{question.id}"] = "on"
            if question.parser_enabled:
                data[f"parser_enabled{question.id}"] = "on"
            if question.llm_enabled:
                data[f"llm_enabled{question.id}"] = "on"
            data[f"text{question.id}"] = question.text
        resp = self.client.post(url, data)
        self.assertRedirects(resp, url)
        self.assertFalse(Anlage1Question.objects.filter(id=q.id).exists())
        self.assertEqual(Anlage1Question.objects.count(), len(questions) - 1)

    def _build_post_data(self, *, new=False, parser=True, llm=True):
        """Hilfsfunktion zum Erstellen der POST-Daten."""
        data = {}
        for q in Anlage1Question.objects.all():
            if q.parser_enabled:
                data[f"parser_enabled{q.id}"] = "on"
            if q.llm_enabled:
                data[f"llm_enabled{q.id}"] = "on"
            data[f"text{q.id}"] = q.text
        if new:
            data["new_text"] = "Neue Frage?"
            if parser:
                data["new_parser_enabled"] = "on"
            if llm:
                data["new_llm_enabled"] = "on"
        return data

    def test_add_new_question_with_flags(self):
        url = reverse("admin_anlage1")
        count = Anlage1Question.objects.count()
        resp = self.client.post(
            url, self._build_post_data(new=True, parser=True, llm=False)
        )
        self.assertRedirects(resp, url)
        self.assertEqual(Anlage1Question.objects.count(), count + 1)
        q = Anlage1Question.objects.order_by("-num").first()
        self.assertEqual(q.text, "Neue Frage?")
        self.assertTrue(q.parser_enabled)
        self.assertFalse(q.llm_enabled)

    def test_add_new_question_unchecked(self):
        url = reverse("admin_anlage1")
        count = Anlage1Question.objects.count()
        resp = self.client.post(
            url, self._build_post_data(new=True, parser=False, llm=False)
        )
        self.assertRedirects(resp, url)
        self.assertEqual(Anlage1Question.objects.count(), count + 1)
        q = Anlage1Question.objects.order_by("-num").first()
        self.assertFalse(q.parser_enabled)
        self.assertFalse(q.llm_enabled)


class ModelSelectionTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user("modeluser", password="pass")
        self.client.login(username="modeluser", password="pass")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        BVProjectFile.objects.create(
            projekt=self.projekt,
            anlage_nr=1,
            upload=SimpleUploadedFile("a.txt", b"data"),
            text_content="Text",
        )
        LLMConfig.objects.create(
            default_model="d",
            gutachten_model="g",
            anlagen_model="a",
        )

    def test_projekt_check_uses_category(self):
        url = reverse("projekt_check", args=[self.projekt.pk])
        with patch("core.views.query_llm", return_value="ok") as mock_q:
            resp = self.client.post(url, {"model_category": "gutachten"})
        self.assertEqual(resp.status_code, 200)
        mock_q.assert_called_with(ANY, {}, model_name="g", model_type="default")

    def test_file_check_uses_category(self):
        url = reverse("projekt_file_check", args=[self.projekt.pk, 1])
        with patch("core.views.check_anlage1") as mock_func:
            mock_func.return_value = {"task": "check_anlage1"}
            resp = self.client.post(url, {"model_category": "anlagen"})
        self.assertEqual(resp.status_code, 200)
        mock_func.assert_called_with(self.projekt.pk, model_name="a")

    def test_forms_show_categories(self):
        edit_url = reverse("projekt_edit", args=[self.projekt.pk])
        resp = self.client.get(edit_url)
        self.assertContains(resp, "Standard")
        self.assertContains(resp, "Gutachten")
        self.assertContains(resp, "Anlagen")

        view_url = reverse(
            "projekt_file_check_view", args=[self.projekt.anlagen.first().pk]
        )
        with patch("core.views.check_anlage1") as mock_func:
            mock_func.return_value = {"task": "check_anlage1"}
            resp = self.client.get(view_url)
        self.assertRedirects(
            resp,
            reverse("projekt_file_edit_json", args=[self.projekt.anlagen.first().pk]),
        )


    def test_functions_check_uses_model(self):
        url = reverse("projekt_functions_check", args=[self.projekt.pk])
        with patch("core.views.check_anlage2_functions") as mock_func:
            mock_func.return_value = []
            resp = self.client.post(url, {"model": "mf"})
        self.assertEqual(resp.status_code, 200)
        mock_func.assert_called_with(self.projekt.pk, model_name="mf")


class CommandModelTests(TestCase):
    def test_command_passes_model(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        BVProjectFile.objects.create(
            projekt=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"d"),
            text_content="Text",
        )
        with patch("core.management.commands.check_anlage2.check_anlage2") as mock_func:
            mock_func.return_value = {"ok": True}
            call_command("check_anlage2", str(projekt.pk), "--model", "m3")
        mock_func.assert_called_with(projekt.pk, model_name="m3")

    def test_analyse_command_passes_model(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        BVProjectFile.objects.create(
            projekt=projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("b.txt", b"d"),
            text_content="- Login",
        )
        with patch(
            "core.management.commands.analyse_anlage2.analyse_anlage2"
        ) as mock_func:
            mock_func.return_value = {"missing": [], "additional": []}
            call_command("analyse_anlage2", str(projekt.pk), "--model", "m4")
        mock_func.assert_called_with(projekt.pk, model_name="m4")



class Anlage2FunctionTests(TestCase):
    def test_check_anlage2_functions_creates_result(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        func = Anlage2Function.objects.create(name="Login")
        llm_reply = json.dumps({
            "technisch_verfuegbar": True,
            "ki_beteiligung": False,
        })
        with patch("core.llm_tasks.query_llm", return_value=llm_reply):
            data = check_anlage2_functions(projekt.pk)
        res = Anlage2FunctionResult.objects.get(projekt=projekt, funktion=func)
        self.assertTrue(res.technisch_verfuegbar)
        self.assertIs(res.ki_beteiligung, False)
        self.assertTrue(data[0]["technisch_verfuegbar"])
        self.assertEqual(list(data[0].keys()), ["technisch_verfuegbar", "ki_beteiligung", "source", "funktion"])


class CommandFunctionsTests(TestCase):
    def test_functions_command_passes_model(self):
        projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        Anlage2Function.objects.create(name="Login")
        with patch(
            "core.management.commands.check_anlage2_functions.check_anlage2_functions"
        ) as mock_func:
            mock_func.return_value = []
            call_command("check_anlage2_functions", str(projekt.pk), "--model", "m5")
        mock_func.assert_called_with(projekt.pk, model_name="m5")







class ProjektDetailAdminButtonTests(TestCase):
    def setUp(self):
        admin_group = Group.objects.create(name="admin")
        self.admin = User.objects.create_user("padmin", password="pass")
        self.admin.groups.add(admin_group)
        self.user = User.objects.create_user("puser", password="pass")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")

    def test_admin_user_sees_link(self):
        self.client.login(username="padmin", password="pass")
        resp = self.client.get(reverse("projekt_detail", args=[self.projekt.pk]))
        self.assertContains(resp, reverse("admin_projects"))

    def test_regular_user_hides_link(self):
        self.client.login(username="puser", password="pass")
        resp = self.client.get(reverse("projekt_detail", args=[self.projekt.pk]))
        self.assertNotContains(resp, reverse("admin_projects"))


class FunctionImportExportTests(TestCase):
    def setUp(self):
        admin_group = Group.objects.create(name="admin")
        self.user = User.objects.create_user("adminie", password="pass")
        self.user.groups.add(admin_group)
        self.client.login(username="adminie", password="pass")

    def test_export_returns_json(self):
        func = Anlage2Function.objects.create(name="Login")
        Anlage2SubQuestion.objects.create(funktion=func, frage_text="Warum?")
        url = reverse("anlage2_function_export")
        resp = self.client.get(url)
        self.assertEqual(resp.status_code, 200)
        data = json.loads(resp.content)
        self.assertEqual(data[0]["name"], "Login")
        self.assertEqual(data[0]["subquestions"][0]["frage_text"], "Warum?")

    def test_import_creates_functions(self):
        payload = json.dumps([
            {"name": "Login", "subquestions": ["Frage"]}
        ])
        file = SimpleUploadedFile("func.json", payload.encode("utf-8"))
        url = reverse("anlage2_function_import")
        resp = self.client.post(
            url,
            {"json_file": file, "clear_first": "on"},
            format="multipart",
        )
        self.assertRedirects(resp, reverse("anlage2_function_list"))
        self.assertTrue(Anlage2Function.objects.filter(name="Login").exists())

    def test_import_accepts_german_keys(self):
        payload = json.dumps([
            {
                "funktion": "Anwesenheit",
                "unterfragen": [
                    {"frage": "Testfrage"}
                ],
            }
        ])
        file = SimpleUploadedFile("func_de.json", payload.encode("utf-8"))
        url = reverse("anlage2_function_import")
        resp = self.client.post(
            url,
            {"json_file": file, "clear_first": "on"},
            format="multipart",
        )
        self.assertRedirects(resp, reverse("anlage2_function_list"))
        self.assertTrue(Anlage2Function.objects.filter(name="Anwesenheit").exists())
        self.assertEqual(
            Anlage2SubQuestion.objects.filter(funktion__name="Anwesenheit").count(),
            1,
        )


class GutachtenLLMCheckTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user("gcheck", password="pass")
        self.client.login(username="gcheck", password="pass")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        self.knowledge = SoftwareKnowledge.objects.create(
            projekt=self.projekt,
            software_name="A",
            is_known_by_llm=True,
        )
        self.gutachten = Gutachten.objects.create(software_knowledge=self.knowledge, text="Test")

    def test_endpoint_updates_note(self):
        url = reverse("gutachten_llm_check", args=[self.gutachten.pk])
        with patch("core.views.check_gutachten_functions") as mock_func:
            mock_func.return_value = "Hinweis"
            resp = self.client.post(url)
        self.assertRedirects(resp, reverse("gutachten_view", args=[self.gutachten.pk]))
        self.projekt.refresh_from_db()
        self.assertEqual(self.projekt.gutachten_function_note, "Hinweis")


class FeatureVerificationTests(TestCase):
    def setUp(self):
        self.projekt = BVProject.objects.create(
            software_typen="Word, Excel",
            beschreibung="x",
        )
        BVProjectFile.objects.create(
            projekt=self.projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"data"),
        )
        self.func = Anlage2Function.objects.create(name="Export")
        self.sub = Anlage2SubQuestion.objects.create(
            funktion=self.func,
            frage_text="Warum?",
        )

    def test_any_yes_returns_true(self):
        with patch(
            "core.llm_tasks.query_llm",
            side_effect=["Ja", "Nein", "Begruendung", "Nein"],
        ) as mock_q:
            result = worker_verify_feature(self.projekt.pk, "function", self.func.pk)
        self.assertEqual(
            result,
            {
                "technisch_verfuegbar": True,
                "ki_begruendung": "Begruendung",
                "ki_beteiligt": False,
                "ki_beteiligt_begruendung": "",
            },
        )
        self.assertEqual(mock_q.call_count, 4)
        pf = BVProjectFile.objects.get(projekt=self.projekt, anlage_nr=2)
        self.assertEqual(
            pf.verification_json["Export"],
            {
                "technisch_verfuegbar": True,
                "ki_begruendung": "Begruendung",
                "ki_beteiligt": False,
                "ki_beteiligt_begruendung": "",
            },
        )

    def test_all_no_returns_false(self):
        with patch(
            "core.llm_tasks.query_llm",
            side_effect=["Nein", "Nein"],
        ):
            result = worker_verify_feature(self.projekt.pk, "subquestion", self.sub.pk)
        self.assertEqual(
            result,
            {
                "technisch_verfuegbar": False,
                "ki_begruendung": "",
                "ki_beteiligt": None,
                "ki_beteiligt_begruendung": "",
            },
        )

    def test_mixed_returns_none(self):
        with patch(
            "core.llm_tasks.query_llm",
            side_effect=["Unsicher", "Nein"],
        ):
            result = worker_verify_feature(self.projekt.pk, "function", self.func.pk)
        self.assertIsNone(result["technisch_verfuegbar"])
        self.assertEqual(result["ki_begruendung"], "")
        self.assertIsNone(result["ki_beteiligt"])
        self.assertEqual(result["ki_beteiligt_begruendung"], "")


class InitialCheckTests(TestCase):
    def setUp(self):
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")

    def test_known_software_stores_description(self):
        with patch(
            "core.llm_tasks.query_llm",
            side_effect=["Ja", "Beschreibung"],
        ) as mock_q:
            sk = SoftwareKnowledge.objects.create(
                projekt=self.projekt,
                software_name="A",
            )
            result = worker_run_initial_check(sk.pk)
        self.assertTrue(result["is_known_by_llm"])
        self.assertEqual(result["description"], "Beschreibung")
        self.assertEqual(mock_q.call_count, 2)
        sk.refresh_from_db()
        self.assertTrue(sk.is_known_by_llm)
        self.assertEqual(sk.description, "Beschreibung")

    def test_unknown_sets_flags(self):
        with patch("core.llm_tasks.query_llm", return_value="Nein") as mock_q:
            sk = SoftwareKnowledge.objects.create(
                projekt=self.projekt,
                software_name="A",
            )
            result = worker_run_initial_check(sk.pk)
        self.assertFalse(result["is_known_by_llm"])
        self.assertEqual(result["description"], "")
        self.assertEqual(mock_q.call_count, 1)
        sk.refresh_from_db()
        self.assertFalse(sk.is_known_by_llm)
        self.assertEqual(sk.description, "")



class EditKIJustificationTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user("justi", password="pass")
        self.client.login(username="justi", password="pass")
        self.projekt = BVProject.objects.create(software_typen="A", beschreibung="x")
        self.file = BVProjectFile.objects.create(
            projekt=self.projekt,
            anlage_nr=2,
            upload=SimpleUploadedFile("a.txt", b"data"),

            verification_json={
                "Export": {"technisch_verfuegbar": True, "ki_begruendung": "Alt"}
            },
        )
        self.func = Anlage2Function.objects.create(name="Export")

    def test_get_returns_form(self):
        url = reverse(
            "edit_ki_justification",
            args=[self.file.pk],
        ) + f"?function={self.func.pk}"
        resp = self.client.get(url)
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, "Alt")

    def test_post_updates_value(self):
        url = reverse("edit_ki_justification", args=[self.file.pk])
        resp = self.client.post(
            url,
            {"function": self.func.pk, "ki_begruendung": "Neu"},
        )
        self.assertRedirects(resp, reverse("projekt_file_edit_json", args=[self.file.pk]))
        self.file.refresh_from_db()
        self.assertEqual(
            self.file.verification_json["Export"]["ki_begruendung"],
            "Neu",
        )


class VerificationToInitialTests(TestCase):
    def setUp(self):
        self.project = BVProject.objects.create(software_typen="A", beschreibung="x")
        BVProjectFile.objects.create(
            projekt=self.project,
            anlage_nr=2,
            upload=SimpleUploadedFile("v.txt", b"data"),
        )
        self.func = Anlage2Function.objects.create(name="Export")
        self.sub = Anlage2SubQuestion.objects.create(
            funktion=self.func,
            frage_text="Warum?",
        )

    def test_conversion_reads_ai_fields(self):
        data = {
            "Export": {
                "technisch_verfuegbar": True,
                "ki_beteiligt": True,
                "ki_beteiligt_begruendung": "Grund",
            },
            "Export: Warum?": {
                "technisch_verfuegbar": False,
                "ki_beteiligt": False,
                "ki_beteiligt_begruendung": "Nein",
            },
        }

        result = _verification_to_initial(data)
        func_data = result["functions"][str(self.func.id)]
        self.assertTrue(func_data["technisch_vorhanden"])
        self.assertTrue(func_data["ki_beteiligt"])
        self.assertEqual(func_data["ki_beteiligt_begruendung"], "Grund")

        sub_data = func_data["subquestions"][str(self.sub.id)]
        self.assertFalse(sub_data["technisch_vorhanden"])
        self.assertFalse(sub_data["ki_beteiligt"])
        self.assertEqual(sub_data["ki_beteiligt_begruendung"], "Nein")







