from pathlib import Path
from docx import Document
import logging
import re
import string
import zipfile

from .models import (
    Anlage2Config,
    Anlage2ColumnHeading,
    Anlage2Function,
    Anlage2SubQuestion,
)

# Zuordnung der Standardspalten zu ihren Feldnamen
HEADER_FIELDS = {
    "technisch vorhanden": "technisch_vorhanden",  # <-- Der kanonische TEXT-Schlüssel, der gesucht wird
    "einsatz bei telefónica": "einsatz_bei_telefonica",
    "zur lv-kontrolle": "zur_lv_kontrolle",
    "ki-beteiligung": "ki_beteiligung",
}

# Aliasüberschriften werden ausschließlich über das Modell
# ``Anlage2ColumnHeading`` gepflegt


def extract_text(path: Path) -> str:
    """Extrahiert den gesamten Text einer DOCX-Datei."""
    doc = Document(str(path))
    text = "\n".join(p.text for p in doc.paragraphs)
    debug_logger = logging.getLogger("anlage2_debug")
    debug_logger.debug("Rohtext aus %s: %r", path, text)
    return text


def get_docx_page_count(path: Path) -> int:
    """Ermittelt die Seitenzahl eines DOCX-Dokuments.

    Es werden sowohl manuelle Seitenumbrüche (``<w:br w:type="page"/>``)
    als auch Abschnittswechsel gezählt, sofern sie einen neuen
    Seitenumbruch bewirken. Ein einfaches Dokument besitzt somit immer
    mindestens eine Seite.
    """

    doc = Document(str(path))
    root = doc.part.element

    # Zähle explizite Seitenumbrüche
    page_breaks = len(root.xpath('.//w:br[@w:type="page"]'))

    # Abschnittswechsel können ebenfalls neue Seiten erzeugen. Das
    # abschließende ``sectPr`` beschreibt lediglich das letzte
    # Abschnittslayout und zählt daher nicht als Umbruch.
    section_breaks = 0
    sect_prs = root.xpath(".//w:sectPr")
    for sp in sect_prs[1:]:
        type_el = sp.xpath("./w:type")
        val = (
            type_el[0].get(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
            )
            if type_el
            else None
        )
        if val != "continuous":
            section_breaks += 1

    return 1 + page_breaks + section_breaks


def get_pdf_page_count(path: Path) -> int:
    """Ermittelt die Seitenzahl eines PDF-Dokuments."""

    import fitz  # PyMuPDF

    with fitz.open(str(path)) as pdf:
        return pdf.page_count


def extract_images(path: Path) -> list[bytes]:
    """Extrahiert alle eingebetteten Bilder aus einer DOCX-Datei.

    Die Funktion liest das DOCX-Archiv als ZIP und gibt die rohen
    Bilddaten aus dem Unterordner ``word/media`` zurück.
    """

    logger = logging.getLogger(__name__)
    images: list[bytes] = []
    try:
        with zipfile.ZipFile(path) as zf:
            for name in zf.namelist():
                if name.startswith("word/media/"):
                    with zf.open(name) as fh:
                        images.append(fh.read())
    except Exception as exc:  # pragma: no cover - ungültige Datei
        logger.error("Fehler beim Extrahieren der Bilder aus %s: %s", path, exc)
    return images


def _parse_cell_value(text: str) -> dict[str, object]:
    """Parst eine Tabellenzelle mit optionaler Zusatzinfo.

    Gibt ein Dictionary mit ``value`` (bool oder ``None``) und ``note``
    (optionaler Zusatztext) zurück.
    """
    text = text.strip()
    match = re.match(r"^(ja|nein)\b(.*)$", text, flags=re.IGNORECASE)
    if match:
        value = match.group(1).lower() == "ja"
        rest = match.group(2).strip()
        rest = re.sub(r"^[,:;\-\s]+", "", rest)
        if rest.startswith("(") and rest.endswith(")"):
            rest = rest[1:-1].strip()
        note = rest or None
        return {"value": value, "note": note}

    note: str | None = None
    m = re.search(r"\(([^)]*)\)", text)
    if m:
        note = m.group(1).strip() or None
        text = (text[: m.start()] + text[m.end() :]).strip()

    lower = text.lower()
    if lower.startswith("ja"):
        value = True
    elif lower.startswith("nein"):
        value = False
    else:
        value = None
    return {"value": value, "note": note}


def _normalize_header_text(text: str) -> str:
    """Bereinigt eine Tabellenüberschrift für den Vergleich."""
    # Zuerst doppelt-escapte Newlines durch ein einfaches Leerzeichen ersetzen
    # Wichtig: Dies muss passieren, bevor "echte" Newlines ersetzt werden,
    # falls beides vorkommt, oder wenn die Quelle schon so escapet.
    text = text.replace("\\n", " ")
    # Ersetzt '\\n' durch ein Leerzeichen

    text = text.replace(
        "\n", " "
    )  # Behält die alte Zeile für den Fall von "echten" Newlines
    text = re.sub(r"ja\s*/\s*nein", "", text, flags=re.I)
    text = text.strip().lower()
    # Fragezeichen und Doppelpunkte am Ende entfernen
    text = re.sub(r"[?:]+$", "", text).strip()
    # Mehrere Leerzeichen und Tabulatoren vereinheitlichen
    text = re.sub(r"[ \t]+", " ", text)
    return text


def _normalize_function_name(name: str) -> str:
    """Bereinigt Funktionsnamen für Vergleiche."""
    text = name.replace("\n", " ").strip().lower()
    return re.sub(r"[ \t]+", " ", text)




def _build_header_map(cfg: Anlage2Config | None) -> dict[str, str]:
    """Erzeugt ein Mapping aller bekannten Header auf ihre kanonische Form.

    Aliasüberschriften werden ausschließlich aus ``Anlage2ColumnHeading``
    geladen.
    """

    logger = logging.getLogger(__name__)

    mapping: dict[str, str] = {"funktion": "funktion"}

    def _add_mapping(key: str, canonical: str) -> None:
        """Fügt einen Mapping-Eintrag hinzu und prüft auf Konflikte."""
        if key in mapping and mapping[key] != canonical:
            msg = (
                f"Mehrdeutige Überschrift '{key}' "
                f"für '{mapping[key]}' und '{canonical}'"
            )
            logger.warning(msg)
            raise ValueError(msg)
        mapping[key] = canonical

    if cfg:
        logger.debug(
            "Konfiguriere Alias-Überschriften: %s",
            [str(h) for h in cfg.headers.all()],
        )
    else:
        logger.debug("Keine Anlage2Config gefunden")

    for canonical, field in HEADER_FIELDS.items():
        _add_mapping(_normalize_header_text(canonical), canonical)
        if cfg:
            for h in cfg.headers.filter(field_name=field):
                logger.debug(f"Roh-Alias (repr): {repr(h.text)}")
                logger.debug(
                    f"Normalisierter Alias: '{_normalize_header_text(h.text)}'"
                )
                _add_mapping(_normalize_header_text(h.text), canonical)

    logger.debug(
        "Alias-Überschriften: %s",
        [str(h) for h in cfg.headers.all()] if cfg else [],
    )

    return mapping


def parse_anlage2_table(path: Path) -> list[dict[str, object]]:
    """Liest und parst eine Anlage‑2‑Tabelle aus einer DOCX-Datei.

    Die Funktion versucht, sowohl einfache Tabellen als auch die in Anlage 2
    vorkommende hierarchische Struktur zu interpretieren. Dabei werden zwei
    Zeilentypen unterschieden:

    - **Hauptfunktions-Zeilen** enthalten einen Funktionsnamen in der ersten
      Spalte, die zweite Funktionsspalte ist leer.
    - **Unterfragen-Zeilen** besitzen einen Text in der zweiten
      Funktionsspalte. Der Text der ersten Spalte wird ignoriert, da er in der
      Regel einen generischen Hinweis enthält. Der vollständige Funktionsname
      ergibt sich aus dem zuletzt gefundenen Hauptfunktionsnamen gefolgt von
      einem Doppelpunkt und dem Unterfragentext.

    Die Rückgabe ist eine Liste von Dictionaries mit dem Schlüssel
    ``funktion`` sowie den in ``HEADER_FIELDS`` definierten Spalten. Jede
    Spalte enthält wiederum ein Dictionary mit den Schlüsseln ``value`` und
    ``note``.
    """
    logger = logging.getLogger(__name__)
    debug_logger = logging.getLogger("anlage2_debug")
    logger.debug(f"Starte parse_anlage2_table mit Pfad: {path}")
    debug_logger.info("parse_anlage2_table gestartet: %s", path)

    try:
        doc = Document(str(path))
        logger.debug(f"Dokument erfolgreich geladen: {path}")
    except Exception as e:  # pragma: no cover - ungültige Datei
        logger.error(f"Fehler beim Laden der Datei {path}: {e}")
        return []

    cfg = Anlage2Config.get_instance()
    logger.debug("Aktive Anlage2Config: %s", cfg)
    header_map = _build_header_map(cfg)
    logger.debug("Erzeugtes Header-Mapping: %s", header_map)

    results: list[dict[str, object]] = []
    found: list[str] = []
    skipped = 0
    if not doc.tables:
        logger.debug("Keine Tabellen im Dokument gefunden")
        debug_logger.debug("Keine Tabellen im Dokument gefunden")
    for table_idx, table in enumerate(doc.tables):
        headers_raw = [cell.text for cell in table.rows[0].cells]
        headers = [
            header_map.get(_normalize_header_text(h), _normalize_header_text(h))
            for h in headers_raw
        ]
        logger.debug(
            f"Tabelle {table_idx}: Roh-Header = {headers_raw}, Normiert = {headers}"
        )
        try:
            func_indices = [i for i, h in enumerate(headers) if h == "funktion"]
            if not func_indices:
                raise ValueError("funktion")
            idx_func = func_indices[0]
            idx_tech = headers.index("technisch vorhanden")
            idx_tel = headers.index("einsatz bei telefónica")
            idx_lv = headers.index("zur lv-kontrolle")
        except ValueError as ve:
            logger.debug(f"Tabelle {table_idx}: Erwartete Spalten nicht gefunden: {ve}")
            continue
        idx_ki = (
            headers.index("ki-beteiligung") if "ki-beteiligung" in headers else None
        )

        col_indices = {
            "technisch_verfuegbar": idx_tech,
            "einsatz_telefonica": idx_tel,
            "zur_lv_kontrolle": idx_lv,
        }
        if idx_ki is not None:
            col_indices["ki_beteiligung"] = idx_ki

        current_main_function_name = None

        for row_idx, row in enumerate(table.rows[1:], start=1):
            main_col_text = row.cells[idx_func].text.strip()
            sub_col_text = (
                row.cells[idx_func + 1].text.strip()
                if len(row.cells) > idx_func + 1
                else ""
            )

            logger.debug(
                "Tabelle %s, Zeile %s: main='%s', sub='%s'",
                table_idx,
                row_idx,
                main_col_text,
                sub_col_text,
            )

            row_data: dict[str, object] | None = None
            is_sub = False

            if main_col_text and "Wenn die Funktion technisch" not in main_col_text:
                current_main_function_name = main_col_text
                debug_logger.debug(
                    "Hauptfunktion erkannt: %s", current_main_function_name
                )
                row_data = {"funktion": current_main_function_name}
            elif sub_col_text and current_main_function_name:
                full_name = f"{current_main_function_name}: {sub_col_text}"
                debug_logger.debug("Unterfrage erkannt: %s", full_name)
                row_data = {"funktion": full_name}
                is_sub = True

            if row_data is None:
                logger.debug(
                    "Zeile %s: Keine verarbeitbare Funktion gefunden, übersprungen",
                    row_idx,
                )
                skipped += 1
                continue

            for col_name, idx in col_indices.items():
                if idx is not None:
                    if is_sub and col_name == "technisch_verfuegbar":
                        continue
                    row_data[col_name] = _parse_cell_value(row.cells[idx].text)

            debug_logger.debug("Verarbeite Zeile %s: %s", row_idx, row_data)
            found.append(row_data["funktion"])
            logger.debug(
                "Zeile %s: Funktion '%s' Daten %s",
                row_idx,
                row_data["funktion"],
                {k: row_data[k] for k in col_indices},
            )

            results.append(row_data)

        if results:
            logger.debug(f"Tabelle {table_idx}: Ergebnisse gefunden, beende Suche.")
            break

    logger.debug(f"Endgültige Ergebnisse: {results}")
    if found:
        debug_logger.info("Gefundene Funktionen: %s", ", ".join(found))
    if skipped:
        debug_logger.info("Übersprungene Zeilen: %s", skipped)
    debug_logger.info("parse_anlage2_table beendet: %s", path)
    return results


