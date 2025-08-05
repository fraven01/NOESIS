import logging
import re
from pathlib import Path
from typing import Dict

from docx import Document

from .models import BVProjectFile, Anlage3ParserRule

logger = logging.getLogger("anlage3_detail")
result_logger = logging.getLogger("anlage3_result")


_DEFAULT_ALIAS_MAP: Dict[str, list[str]] = {
    "name": ["name der auswertung", "name", "bezeichnung"],
    "beschreibung": ["beschreibung", "kurzbeschreibung"],
    "zeitraum": ["zeitraum", "auswertungszeitraum"],
    "art": ["art der auswertung", "auswertungsart", "typ"],
}


def _get_alias_map() -> Dict[str, list[str]]:
    rules = Anlage3ParserRule.objects.all()
    if rules:
        alias_map: Dict[str, list[str]] = {}
        for r in rules:
            alias_map[r.field_name] = list(r.aliases)
        return alias_map
    return _DEFAULT_ALIAS_MAP


def _normalize(text: str) -> str:
    text = text.lower().strip()
    text = re.sub(r"[\s_:]+", " ", text)
    return text


def parse_anlage3(project_file: BVProjectFile) -> Dict[str, str]:
    """Liest Metadaten aus einer DOCX-Datei der Anlage 3."""
    path = Path(project_file.upload.path)
    if not path.exists() or path.suffix.lower() != ".docx":
        logger.debug("Anlage3 Parser: Datei nicht vorhanden oder kein DOCX")
        return {}

    try:
        doc = Document(str(path))
    except Exception as exc:  # pragma: no cover - ungültige Datei
        logger.error("Anlage3 Parser Fehler beim Laden: %s", exc)
        return {}

    alias_map = _get_alias_map()
    result: Dict[str, str] = {"name": "", "beschreibung": "", "zeitraum": "", "art": ""}

    def handle_pair(key: str, value: str) -> None:
        norm = _normalize(key)
        for field, aliases in alias_map.items():
            for alias in aliases:
                if norm.startswith(_normalize(alias)):
                    if not result[field]:
                        result[field] = value.strip()
                    return

    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            handle_pair(row.cells[0].text, row.cells[1].text)

    for para in doc.paragraphs:
        text = para.text.strip()
        if ":" in text:
            before, after = text.split(":", 1)
            handle_pair(before, after)

    logger.debug("Anlage3 Parser Ergebnis: %s", result)
    result_logger.debug("Anlage3 Endergebnis: %s", result)
    return result
